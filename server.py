#!/usr/bin/env python3
"""
微信文件监控 + 作业提交仪表盘 后端服务
监控 WeChat FileStorage 目录，匹配学生名单，追踪作业提交状态
"""

import os
import sys
import json
import io
import time
import shutil
import threading
import argparse
import subprocess
import tempfile
import atexit
import re
import hashlib
import traceback
import queue
import uuid
import urllib.request
import urllib.error
import secrets
import socket
import ipaddress
import hmac
from http.cookies import SimpleCookie
from datetime import datetime, date, timedelta
from pathlib import Path
from http.server import HTTPServer, SimpleHTTPRequestHandler
from socketserver import ThreadingMixIn

# 文档解析库（可选，缺少时降级）
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

from urllib.parse import urlparse, parse_qs, quote

try:
    import ai_classifier
    HAS_AI_CLASSIFIER = True
except Exception as _ai_import_error:
    ai_classifier = None
    HAS_AI_CLASSIFIER = False
    print(f"[WARN] 分类大脑模块不可用，已降级到原有规则：{_ai_import_error}")

class ThreadingHTTPServer(ThreadingMixIn, HTTPServer):
    """多线程 HTTP 服务器，每个请求在独立线程中处理"""
    daemon_threads = True
    allow_reuse_address = True

# ---------------------------------------------------------------------------
# 数据路径
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
CONFIG_PATH = DATA_DIR / "config.json"
STUDENTS_PATH = DATA_DIR / "students.json"
SUBMISSIONS_PATH = DATA_DIR / "submissions.json"
LOCK_PATH = DATA_DIR / "server.lock"
WATCHER_STATE_FILE = DATA_DIR / "watcher_state.json"
AI_RULES_PATH = DATA_DIR / "ai_rules.json"

# ANSI 终端颜色（PowerShell 7+ / Windows Terminal 支持，旧版 PowerShell 5.1 会显示转义码）
class C:
    RESET = "\033[0m"
    GREEN = "\033[92m"
    RED = "\033[91m"
    YELLOW = "\033[93m"
    GRAY = "\033[90m"

_SPINNER_FRAMES = "⠋⠙⠹⠸⠼⠴⠦⠧⠇⠏"
_ASCII_SPINNER_FRAMES = "|/-\\"

def _load_watcher_state():
    """读取持久化的 known_files 状态。文件不存在 → 空状态。JSON 损坏 → 自动备份后重置。"""
    if not WATCHER_STATE_FILE.exists():
        return {"known_files": {}}
    try:
        with open(WATCHER_STATE_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            if "known_files" not in data or not isinstance(data["known_files"], dict):
                return {"known_files": {}}
            return data
    except (json.JSONDecodeError, OSError) as e:
        print(f"[WARN] watcher_state.json 损坏，已备份为 .bak: {e}", file=sys.stderr)
        try:
            WATCHER_STATE_FILE.rename(WATCHER_STATE_FILE.with_suffix(".json.bak"))
        except OSError:
            pass
        return {"known_files": {}}

def _save_watcher_state(state):
    """原子写入 watcher 状态：先写 .tmp 再 os.replace，防止半写损坏。"""
    try:
        WATCHER_STATE_FILE.parent.mkdir(parents=True, exist_ok=True)
        tmp = WATCHER_STATE_FILE.with_suffix(".json.tmp")
        with open(tmp, "w", encoding="utf-8") as f:
            json.dump(state, f, ensure_ascii=False, indent=2)
        os.replace(tmp, WATCHER_STATE_FILE)
        return True
    except OSError as e:
        print(f"[ERROR] 保存 watcher_state.json 失败: {e}", file=sys.stderr)
        return False

# 默认班级路径（将在 resolve_class_config() 中根据配置动态解析）
_DEFAULT_CLASS_NAME = "课程班级"
_DEFAULT_CLASS_FOLDER = str(Path.home() / "Desktop" / _DEFAULT_CLASS_NAME)
ORGANIZED_DIR = Path(_DEFAULT_CLASS_FOLDER) / "已收作业"
CONVERT_TEMP_DIR = Path(tempfile.gettempdir()) / "wechat-tracker-convert"
PREVIEW_CACHE_DIR = DATA_DIR / "preview_cache"  # Word→PDF 预览缓存，持久化到 data/
APP_VERSION = "0.1.0"
UPDATE_REPOSITORY = "Trip1eY/Assignment_Dashboard"
VERSION_MANIFEST = BASE_DIR / "manifest.json"


def _version_key(value):
    """将常见的 x.y.z 版本号转换为可比较的元组。"""
    numbers = re.findall(r"\d+", str(value or ""))
    return tuple(int(item) for item in numbers[:4]) + (0,) * (4 - len(numbers[:4]))


def current_app_version():
    """返回当前安装版本；更新包安装后以 manifest.json 为准。"""
    try:
        if VERSION_MANIFEST.exists():
            manifest = json.loads(VERSION_MANIFEST.read_text(encoding="utf-8"))
            version = str(manifest.get("version") or "").strip()
            if version:
                return version
    except (OSError, json.JSONDecodeError, AttributeError):
        pass
    return APP_VERSION

THEME_PRESETS = [
    {
        "id": "clean-blue",
        "name": "清爽蓝",
        "description": "默认现代 SaaS 风格",
        "style": "modern",
        "colors": {
            "primary": "#2563eb",
            "accent": "#10b981",
            "background": "#f8fafc",
            "surface": "#ffffff",
            "text": "#0f172a",
            "border": "#dbe3ef",
        },
    },
    {
        "id": "teal",
        "name": "松石绿",
        "description": "清新、低压力",
        "style": "modern",
        "colors": {
            "primary": "#0f766e",
            "accent": "#14b8a6",
            "background": "#f3fbf9",
            "surface": "#ffffff",
            "text": "#102a27",
            "border": "#cde7e2",
        },
    },
    {
        "id": "graphite",
        "name": "石墨灰",
        "description": "克制的工作台感",
        "style": "modern",
        "colors": {
            "primary": "#475569",
            "accent": "#0ea5e9",
            "background": "#f6f7f9",
            "surface": "#ffffff",
            "text": "#111827",
            "border": "#d4dae3",
        },
    },
    {
        "id": "apricot",
        "name": "暖杏色",
        "description": "柔和、低攻击性",
        "style": "modern",
        "colors": {
            "primary": "#d97706",
            "accent": "#059669",
            "background": "#fffaf3",
            "surface": "#ffffff",
            "text": "#1f2937",
            "border": "#f0dcc4",
        },
    },
    {
        "id": "berry",
        "name": "莓果红",
        "description": "更有识别度",
        "style": "modern",
        "colors": {
            "primary": "#be123c",
            "accent": "#7c3aed",
            "background": "#fff7fa",
            "surface": "#ffffff",
            "text": "#1f1720",
            "border": "#f2ccd8",
        },
    },
    {
        "id": "high-contrast",
        "name": "高对比",
        "description": "可读性优先",
        "style": "modern",
        "colors": {
            "primary": "#111827",
            "accent": "#f59e0b",
            "background": "#ffffff",
            "surface": "#ffffff",
            "text": "#000000",
            "border": "#111827",
        },
    },
    {
        "id": "mecha-core",
        "name": "机甲核心",
        "description": "装甲 HUD、能量轨、机械面板",
        "style": "mecha",
        "colors": {
            "primary": "#f97316",
            "accent": "#22d3ee",
            "background": "#08090b",
            "surface": "#111318",
            "text": "#f8fafc",
            "border": "#334155",
        },
    },
    {
        "id": "mecha-heart-awakened",
        "name": "机甲之心·觉醒",
        "description": "悬挂机甲、核心点火、装甲切页转场",
        "style": "mecha-heart",
        "colors": {
            "primary": "#ef233c",
            "accent": "#36e2ff",
            "background": "#050609",
            "surface": "#10151d",
            "text": "#f8fbff",
            "border": "#354454",
        },
    },
    {
        "id": "minimal-line",
        "name": "极简线框",
        "description": "雾白线框、冷调留白",
        "style": "minimal",
        "colors": {
            "primary": "#26364f",
            "accent": "#6d8fb8",
            "background": "#f6f8fb",
            "surface": "#ffffff",
            "text": "#132033",
            "border": "#d8e0ea",
        },
    },
    {
        "id": "spiral-nebula",
        "name": "螺旋星云",
        "description": "深邃 shader、旋涡流光、高级舞台感",
        "style": "shader",
        "colors": {
            "primary": "#8b5cf6",
            "accent": "#38f8d4",
            "background": "#05040b",
            "surface": "#11101c",
            "text": "#f8fbff",
            "border": "#3b3b68",
        },
    },
    {
        "id": "paper-shader",
        "name": "纸感流光",
        "description": "简约纸张、柔和噪声、轻盈高级感",
        "style": "paper",
        "colors": {
            "primary": "#1f1d1a",
            "accent": "#b98252",
            "background": "#f6eddc",
            "surface": "#fff8ea",
            "text": "#1f1d1a",
            "border": "#d8cbb7",
        },
    },
]

DEFAULT_UI_THEME = {
    "active": "clean-blue",
    "style": "modern",
    "apply_to_classic": True,
    "publish_with_update": True,
    "force_publish_theme": False,
    "custom": THEME_PRESETS[0]["colors"].copy(),
}

# 模块级全局（供重启/关闭 API 和 main() 共享）
_http_server = None       # ThreadingHTTPServer 实例，在 main() 里赋值
_app_start_time = time.time()  # 进程启动时间，用于 server-status uptime 计算

# 清理上次遗留的临时转换文件
if CONVERT_TEMP_DIR.exists():
    try:
        shutil.rmtree(str(CONVERT_TEMP_DIR))
    except Exception:
        pass
CONVERT_TEMP_DIR.mkdir(parents=True, exist_ok=True)
_json_locks = {}
watcher = None
_warmup_thread = None  # 后台预热线程引用
_warmup_lock = threading.Lock()

# Preview conversion is deliberately serialized. Word COM is much more stable
# when one worker owns the conversion lifecycle instead of each HTTP request
# starting a new Word instance.
_preview_queue = queue.Queue()
_preview_jobs = {}
_preview_jobs_lock = threading.Lock()
_preview_worker = None
_preview_worker_lock = threading.Lock()
_PREVIEW_CACHE_MAX_BYTES = 1024 * 1024 * 1024
_PREVIEW_CACHE_MAX_AGE_SECONDS = 30 * 24 * 60 * 60

def _preview_cache_key(src):
    stat = src.stat()
    path_key = hashlib.md5(str(src).encode("utf-8")).hexdigest()[:12]
    return f"{path_key}_{stat.st_mtime_ns}_{stat.st_size}"

def _preview_cache_path(src):
    PREVIEW_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    return PREVIEW_CACHE_DIR / f"{_preview_cache_key(src)}.pdf"

def _find_libreoffice():
    candidates = [shutil.which("soffice"), shutil.which("libreoffice")]
    candidates.extend([
        os.environ.get("PROGRAMFILES", "") + r"\LibreOffice\program\soffice.exe",
        os.environ.get("PROGRAMFILES(X86)", "") + r"\LibreOffice\program\soffice.exe",
    ])
    for value in candidates:
        if value and Path(value).exists():
            return str(Path(value))
    return ""

def _convert_preview_source(src, target):
    """Convert one document, preferring Word and falling back to LibreOffice."""
    try:
        with _WordComContext() as ctx:
            doc = ctx.open(src)
            doc.ExportAsFixedFormat(str(target), 17)
        if target.exists() and target.stat().st_size > 0:
            return "word", ""
    except Exception as exc:
        print(f"[Preview] Word conversion failed for {src.name}: {exc}")

    soffice = _find_libreoffice()
    if soffice:
        temp_dir = target.parent / f"lo_{uuid.uuid4().hex}"
        temp_dir.mkdir(parents=True, exist_ok=True)
        try:
            flags = getattr(subprocess, "CREATE_NO_WINDOW", 0)
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(temp_dir), str(src)],
                capture_output=True, text=True, timeout=60, creationflags=flags,
            )
            generated = temp_dir / f"{src.stem}.pdf"
            if result.returncode == 0 and generated.exists() and generated.stat().st_size > 0:
                shutil.move(str(generated), str(target))
                return "libreoffice", ""
            return "", (result.stderr or result.stdout or "LibreOffice conversion failed")[:300]
        except Exception as exc:
            return "", str(exc)[:300]
        finally:
            shutil.rmtree(str(temp_dir), ignore_errors=True)
    return "", "未检测到 Microsoft Word 或 LibreOffice"

def _cleanup_preview_cache():
    if not PREVIEW_CACHE_DIR.exists():
        return
    now = time.time()
    files = [p for p in PREVIEW_CACHE_DIR.glob("*.pdf") if p.is_file()]
    for item in files[:]:
        try:
            if now - item.stat().st_atime > _PREVIEW_CACHE_MAX_AGE_SECONDS:
                item.unlink()
                files.remove(item)
        except OSError:
            pass
    total = sum((p.stat().st_size for p in files if p.exists()), 0)
    if total <= _PREVIEW_CACHE_MAX_BYTES:
        return
    for item in sorted(files, key=lambda p: p.stat().st_atime):
        if total <= _PREVIEW_CACHE_MAX_BYTES:
            break
        try:
            size = item.stat().st_size
            item.unlink()
            total -= size
        except OSError:
            pass

def _preview_worker_loop():
    while True:
        job_id, src = _preview_queue.get()
        try:
            with _preview_jobs_lock:
                _preview_jobs[job_id]["status"] = "running"
            cache_pdf = _preview_cache_path(src)
            if cache_pdf.exists() and cache_pdf.stat().st_size > 0:
                engine, error = "cache", ""
            else:
                tmp_pdf = CONVERT_TEMP_DIR / f"preview_{uuid.uuid4().hex}.pdf"
                engine, error = _convert_preview_source(src, tmp_pdf)
                if engine and tmp_pdf.exists():
                    shutil.copy2(str(tmp_pdf), str(cache_pdf))
                try:
                    tmp_pdf.unlink()
                except OSError:
                    pass
            with _preview_jobs_lock:
                job = _preview_jobs[job_id]
                job.update({"status": "ready" if cache_pdf.exists() else "error",
                            "engine": engine, "error": error,
                            "cache_url": str(cache_pdf)})
            _cleanup_preview_cache()
        except Exception as exc:
            with _preview_jobs_lock:
                _preview_jobs[job_id].update({"status": "error", "error": str(exc)[:300]})
        finally:
            _preview_queue.task_done()

def _ensure_preview_worker():
    global _preview_worker
    with _preview_worker_lock:
        if _preview_worker and _preview_worker.is_alive():
            return
        _preview_worker = threading.Thread(target=_preview_worker_loop, name="preview-worker", daemon=True)
        _preview_worker.start()

def _queue_preview_job(src):
    _ensure_preview_worker()
    try:
        cache_pdf = _preview_cache_path(src)
    except OSError as exc:
        return "", {"status": "error", "error": str(exc)}
    if cache_pdf.exists() and cache_pdf.stat().st_size > 0:
        return "", {"status": "ready", "cache_url": str(cache_pdf), "engine": "cache"}
    with _preview_jobs_lock:
        for job_id, job in _preview_jobs.items():
            if job.get("path") == str(src) and job.get("status") in ("queued", "running"):
                return job_id, dict(job)
        job_id = uuid.uuid4().hex
        job = {"job_id": job_id, "path": str(src), "status": "queued"}
        _preview_jobs[job_id] = job
    _preview_queue.put((job_id, src))
    return job_id, dict(job)

def warm_preview_cache_async():
    """[已禁用] 后台线程：预转换 .docx/.doc 文件到 PDF 缓存（Word COM 会弹窗）"""
    global _warmup_thread
    with _warmup_lock:
        if _warmup_thread and _warmup_thread.is_alive():
            return  # 已有预热任务在运行
        t = threading.Thread(target=_warm_preview_cache_worker, daemon=True)
        t.start()
        _warmup_thread = t

def _warm_preview_cache_worker():
    """实际执行预热的工作线程（复用 Word COM 实例，避免重复 Disp/Quit 导致 RPC 错误）"""
    try:
        # 扫描已收作业目录下的所有 .docx/.doc 文件
        doc_files = []
        scan_roots = []
        # 班级已收作业目录
        if ORGANIZED_DIR.exists():
            scan_roots.append(ORGANIZED_DIR)
        # scan-existing 的扫描目录（多目录队列）
        cfg = load_config_raw()
        for d in cfg.get("scan_dirs", []):
            resolved = _safe_resolve_path(d)
            if resolved:
                scan_roots.append(resolved)
        for root in scan_roots:
            if not root.exists():
                continue
            for ext in (".docx", ".doc"):
                for f in root.rglob(f"*{ext}"):
                    if f.is_file():
                        doc_files.append(f)
        if not doc_files:
            return
        print(f"[WarmUp] 发现 {len(doc_files)} 个 Word 文档，开始后台预热...")
        PREVIEW_CACHE_DIR.mkdir(parents=True, exist_ok=True)

        # 复用单个 Word COM 实例，避免频繁 Disp/Quit 导致 RPC 错误
        import pythoncom
        import win32com.client
        pythoncom.CoInitialize()
        word = None
        converted = 0
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            for src in doc_files:
                try:
                    src_mtime = int(src.stat().st_mtime)
                    cache_key = hashlib.md5(str(src).encode()).hexdigest()[:12] + f"_{src_mtime}"
                    cache_pdf = PREVIEW_CACHE_DIR / f"{cache_key}.pdf"
                    if cache_pdf.exists() and cache_pdf.stat().st_size > 0:
                        continue  # 已缓存

                    tmp_pdf = CONVERT_TEMP_DIR / (src.stem + "_warmup.pdf")
                    ok = False
                    ext = src.suffix.lower()

                    if ext in (".docx", ".doc"):
                        for attempt in range(2):  # 最多重试 1 次
                            doc = None
                            try:
                                doc = word.Documents.Open(str(src), ReadOnly=True)
                                doc.ExportAsFixedFormat(str(tmp_pdf), 17)
                                ok = tmp_pdf.exists() and tmp_pdf.stat().st_size > 0
                                break  # 成功，跳出重试循环
                            except Exception as e:
                                if attempt == 0:
                                    print(f"[WarmUp] Word COM 失败 {src.name}: {e}，1s后重试...")
                                    time.sleep(1)
                                    # 重试前尝试恢复 Word
                                    try:
                                        if doc: doc.Close(SaveChanges=False)
                                    except: pass
                                else:
                                    print(f"[WarmUp] Word COM 重试仍失败 {src.name}: {e}")
                            finally:
                                try:
                                    if doc: doc.Close(SaveChanges=False)
                                except: pass

                    if ok:
                        shutil.copy2(str(tmp_pdf), str(cache_pdf))
                        converted += 1
                        # 清理旧缓存
                        hash_prefix = hashlib.md5(str(src).encode()).hexdigest()[:12]
                        for old in PREVIEW_CACHE_DIR.glob(f"{hash_prefix}_*.pdf"):
                            if old != cache_pdf:
                                try: old.unlink()
                                except: pass
                        # 文件间小延迟，避免 Word COM 过载
                        time.sleep(0.3)

                    try:
                        if tmp_pdf.exists():
                            tmp_pdf.unlink()
                    except: pass

                except Exception as e:
                    print(f"[WarmUp] 跳过 {getattr(src, 'name', src)}: {e}")

        finally:
            try:
                if word: word.Quit()
            except: pass
            pythoncom.CoUninitialize()

        if converted > 0:
            print(f"[WarmUp] 预转换完成，新增 {converted} 个缓存")
    except Exception as e:
        print(f"[WarmUp] 预热失败: {e}")

# 智能检测实验目录（延迟初始化，通过 resolve_class_config() 获取准确路径）
EXPERIMENT_BASE = None  # 将在 resolve_class_config() 中设置

# 实验次数排序映射（确保按逻辑顺序而非字母序排列）
_EXPERIMENT_ORDER = {
    "第一次": 1, "第二次": 2, "第三次": 3, "第四次": 4, "第五次": 5,
    "第六次": 6, "第七次": 7, "第八次": 8, "第九次": 9, "第十次": 10,
    "实验一": 11, "实验二": 12, "实验三": 13, "实验四": 14, "实验五": 15,
    "Lab1": 21, "Lab2": 22, "Lab3": 23, "Lab4": 24, "Lab5": 25,
}

# 中文数字统一映射（全局一份，避免多处重复定义）
_CN_NUM_MAP = {
    '一': '第一次', '二': '第二次', '三': '第三次', '四': '第四次',
    '五': '第五次', '六': '第六次', '七': '第七次', '八': '第八次',
    '九': '第九次', '十': '第十次',
}

def parse_experiment_number(text):
    """从文本中提取实验次数，返回如'第一次'/'第二次'，未找到返回空字符串"""
    import re as _re
    patterns = [
        (r'实验\s*([一二三四五六七八九十]+)', lambda m: _CN_NUM_MAP.get(m.group(1), '')),
        (r'第\s*([一二三四五六七八九十\d]+)\s*次', lambda m: f'第{m.group(1)}次'),
        (r'([一二三四五六七八九十]+)次实验', lambda m: _CN_NUM_MAP.get(m.group(1), '')),
        (r'Lab\s*(\d+)', lambda m: f'第{m.group(1)}次'),
    ]
    for pat, fmt in patterns:
        m = _re.search(pat, text)
        if m:
            result = fmt(m) if callable(fmt) else fmt
            if result:
                return result
    return ''

def is_file_in_correct_path(organized_to, subject_group, experiment):
    """检查 organized_to 路径是否匹配 科目/实验次数/ 结构。
    - 无路径或无 subject_group → 不拦截（让 student/filename 兜底）
    - 否则：subject_group（或其同义词之一）在路径中即视为学科匹配
    - experiment 校验：仅在路径中能识别出**错误实验次数**时拒；
      路径无实验次数（学生名直接挂在科目下）→ True（合法状态，student 字段会兜底）
    """
    if not organized_to or not subject_group:
        return True
    path_lower = organized_to.replace("\\", "/").lower()

    # 学科匹配：同义词任一在路径中
    syns = _get_subject_synonyms(subject_group)
    if not any(s.lower() in path_lower for s in syns):
        return False

    # experiment 校验：仅当路径含其他"第X次"且与目标不一致时拒
    if not experiment:
        return True
    # 路径中有"第X次"但和目标 experiment 不一致 → 拒
    import re as _re_path
    other_exp = _re_path.findall(r"第[一二三四五]次", organized_to)
    if other_exp and experiment not in other_exp:
        return False
    # 路径含"/N/"数字实验次数且与目标不一致 → 拒
    num_map = {"第一次": "1", "第二次": "2", "第三次": "3", "第四次": "4", "第五次": "5"}
    exp_num = num_map.get(experiment)
    if exp_num:
        # 寻找路径中的 /数字/ 段（避免误匹配如 /杨焱煜1/）
        for seg in path_lower.split("/"):
            if seg in num_map.values() and seg != exp_num:
                return False
    return True

def build_submission_record(file_info, student_name, assignment_id, assignment_name,
                            match_score=0, status="matched", organized_to=None, extra=None):
    """构建标准的 submission 记录字典"""
    record = {
        "file": file_info,
        "student": student_name,
        "match_score": match_score,
        "detected_at": datetime.now().isoformat(),
        "assignment_id": assignment_id,
        "assignment_name": assignment_name,
        "status": status,
    }
    if organized_to:
        record["organized_to"] = organized_to
    if extra:
        record.update(extra)
    return record

# 微信文件目录
WECHAT_FILES_BASE = Path.home() / "Documents" / "WeChat Files"
# 微信 4.x 新路径（C:\Users\xxx\xwechat_files\）
XWECHAT_BASE = Path.home() / "xwechat_files"

# ---------------------------------------------------------------------------
# 配置/数据管理
# ---------------------------------------------------------------------------

def load_json(path, default=None):
    try:
        if path.exists():
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception as e:
        print(f"[WARN] Failed to load {path}: {e}")
    if default is not None:
        return default
    return {}

def _json_lock(path):
    key = str(path.resolve())
    if key not in _json_locks:
        _json_locks[key] = threading.Lock()
    return _json_locks[key]

def save_json(path, data):
    path.parent.mkdir(parents=True, exist_ok=True)
    with _json_lock(path):
        fd, tmp_name = tempfile.mkstemp(prefix=path.name + ".", suffix=".tmp", dir=str(path.parent))
        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
                f.write("\n")
            os.replace(tmp_name, path)
        except Exception:
            try:
                os.unlink(tmp_name)
            except OSError:
                pass
            raise

def load_ai_rules():
    if not HAS_AI_CLASSIFIER:
        return {"schema_version": 1, "profile": {}, "subjects": {}, "types": {}}
    return ai_classifier.load_rule_pack(AI_RULES_PATH)

def save_ai_rules(payload):
    if not HAS_AI_CLASSIFIER:
        raise RuntimeError("分类大脑模块不可用")
    return ai_classifier.save_rule_pack(AI_RULES_PATH, payload)

def ai_settings(cfg=None):
    cfg = cfg or load_config_raw()
    defaults = ai_classifier.default_settings() if HAS_AI_CLASSIFIER else {
        "mode": "off", "sensitivity": 0.70,
        "sensitivity_preset": "balanced", "active_semester": "",
    }
    value = cfg.get("ai_classifier")
    if isinstance(value, dict):
        defaults.update(value)
    if defaults.get("mode") not in ("off", "rules"):
        defaults["mode"] = "rules"
    try:
        defaults["sensitivity"] = min(0.95, max(0.50, float(defaults.get("sensitivity", 0.70))))
    except (TypeError, ValueError):
        defaults["sensitivity"] = 0.70
    return defaults

def ai_brain_payload():
    cfg = load_config_raw()
    rules = load_ai_rules()
    assignments = cfg.get("assignments", [])
    feedback = cfg.get("match_feedback", {}) or {}
    submissions = load_submissions()
    pending = submissions.get(PENDING_ARCHIVE_BUCKET, [])
    unknown = 0
    conflicts = 0
    for record in pending if isinstance(pending, list) else []:
        stage = (record.get("classification") or {}).get("stage") or record.get("status", "")
        if stage in ("unknown_subject", "unmatched"):
            unknown += 1
        if stage == "subject_conflict":
            conflicts += 1
    subjects = rules.get("subjects", {})
    return {
        "ok": True,
        "available": HAS_AI_CLASSIFIER,
        "settings": ai_settings(cfg),
        "profile": rules.get("profile", {}),
        "stats": {
            "subjects": len(subjects),
            "active_subjects": sum(1 for item in subjects.values() if item.get("active", True)),
            "confirmed_aliases": sum(len(item.get("confirmed_aliases", [])) for item in subjects.values()),
            "keywords": sum(len(item.get("keywords", [])) for item in subjects.values()),
            "pending": len(pending) if isinstance(pending, list) else 0,
            "unknown": unknown,
            "conflicts": conflicts,
            "corrections": len(feedback.get("subject_corrections", [])),
            "rejections": len(feedback.get("rejected", [])),
        },
        "recent_feedback": list(reversed(feedback.get("subject_corrections", [])[-20:])),
        "rejected": list(reversed(feedback.get("rejected", [])[-20:])),
        "assignments": assignments,
    }

def default_config():
    return {
        "class_name": _DEFAULT_CLASS_NAME,          # 班级名称（如 "课程班级"）
        "class_folder": _DEFAULT_CLASS_FOLDER,      # 班级根目录（如 Desktop/课程班级）
        "wechat_accounts": [],       # 自动检测
        "watch_enabled": True,
        "auto_organize": True,
        "organized_dir": str(ORGANIZED_DIR),
        "experiment_enabled": False,
        "experiment_dir": str(Path(_DEFAULT_CLASS_FOLDER) / "实验"),
        "assignments": [
            {"id": "a1", "name": "第一次作业", "subject": "课程作业", "keywords": ["第一次", "作业"], "due": "", "active": True},
            {"id": "a2", "name": "课程报告", "subject": "课程报告", "keywords": ["课程报告", "报告"], "due": "", "active": True},
            {"id": "a3", "name": "项目作业", "subject": "项目作业", "keywords": ["项目", "项目作业"], "due": "", "active": True},
            {"id": "a4", "name": "课程论文", "subject": "课程论文", "keywords": ["课程论文", "论文"], "due": "", "active": True},
        ],
        "poll_interval": 5,          # 轮询秒数
        "file_keywords": ["作业", "报告", "论文", "实验", "习题", "课设"],
        "scan_dirs": [str(EXPERIMENT_BASE)],  # "扫描已有文件"的根目录列表
        "templates": [],  # 模板文件列表
        "ignored_subjects": [],  # 用户已删除的科目，防止 load_config 从目录重建时复活
        "ignored_assignments": [],  # 用户已删除的作业，防止目录自动扫描重建时复活
        "ui_theme": DEFAULT_UI_THEME.copy(),
        "default_frontend": "classic",
        "lan_access_enabled": False,
        "lan_access_token": "",
        "ai_classifier": {
            "mode": "rules",
            "sensitivity": 0.70,
            "sensitivity_preset": "balanced",
            "active_semester": "",
        },
    }

def normalize_ui_theme(value=None):
    """Return a safe ui_theme object with known preset ids and color keys."""
    preset_map = {p["id"]: p for p in THEME_PRESETS}
    preset_ids = set(preset_map)
    style_ids = {"modern", "mecha", "mecha-heart", "minimal", "shader", "paper"}
    theme = DEFAULT_UI_THEME.copy()
    theme["custom"] = DEFAULT_UI_THEME["custom"].copy()
    if isinstance(value, dict):
        theme.update({k: v for k, v in value.items() if k != "custom"})
        custom = value.get("custom")
        if isinstance(custom, dict):
            clean_custom = {}
            for key in ("primary", "accent", "background", "surface", "text", "border"):
                val = str(custom.get(key, "")).strip()
                if re.fullmatch(r"#[0-9A-Fa-f]{6}", val):
                    clean_custom[key] = val
            theme["custom"].update(clean_custom)
    if theme.get("active") not in preset_ids and theme.get("active") != "custom":
        theme["active"] = DEFAULT_UI_THEME["active"]
    if theme["active"] == "custom":
        if theme.get("style") not in style_ids:
            theme["style"] = DEFAULT_UI_THEME["style"]
    else:
        theme["style"] = preset_map.get(theme["active"], {}).get("style", DEFAULT_UI_THEME["style"])
    for key in ("apply_to_classic", "publish_with_update", "force_publish_theme"):
        theme[key] = bool(theme.get(key))
    return theme

def theme_payload(cfg=None):
    cfg = cfg or load_config_raw()
    theme = normalize_ui_theme(cfg.get("ui_theme"))
    return {
        "active": theme["active"],
        "style": theme["style"],
        "presets": THEME_PRESETS,
        "custom": theme["custom"],
        "apply_to_classic": theme["apply_to_classic"],
        "publish_with_update": theme["publish_with_update"],
        "force_publish_theme": theme["force_publish_theme"],
    }

def configured_file_types(cfg=None):
    """Return normalized file extensions allowed by the current config."""
    default_exts = [".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png", ".zip", ".rar", ".7z"]
    try:
        values = (cfg or load_config_raw()).get("file_types", default_exts)
    except Exception:
        values = default_exts
    result = set()
    for ext in values or default_exts:
        ext = str(ext).strip().lower()
        if not ext:
            continue
        if not ext.startswith("."):
            ext = "." + ext
        result.add(ext)
    return result or set(default_exts)

def load_config_raw():
    default = default_config()
    cfg = load_json(CONFIG_PATH, default)
    for k, v in default.items():
        if k not in cfg:
            cfg[k] = v
    return _clean_config_paths(cfg)

LAN_SESSION_COOKIE = "assignment_dashboard_lan"

def _is_loopback_ip(value):
    try:
        address = ipaddress.ip_address(str(value).split("%", 1)[0])
        if getattr(address, "ipv4_mapped", None):
            address = address.ipv4_mapped
        return address.is_loopback
    except ValueError:
        return False

def _lan_access_token(cfg=None, create=False):
    cfg = cfg or load_config_raw()
    token = str(cfg.get("lan_access_token") or "").strip()
    if create and not token:
        token = secrets.token_urlsafe(12)
        cfg["lan_access_token"] = token
        save_config(cfg)
    return token

def _local_ipv4_addresses():
    found = set()
    try:
        for info in socket.getaddrinfo(socket.gethostname(), None, socket.AF_INET):
            ip = info[4][0]
            address = ipaddress.ip_address(ip)
            if not address.is_loopback and not address.is_link_local:
                found.add(ip)
    except OSError:
        pass
    return sorted(found)

def network_access_payload(cfg=None, port=18765, include_token=False, is_local=True):
    cfg = cfg or load_config_raw()
    enabled = bool(cfg.get("lan_access_enabled", False))
    payload = {
        "enabled": enabled,
        "is_local_request": bool(is_local),
        "bind_host": "0.0.0.0" if enabled else "127.0.0.1",
        "port": int(port),
        "lan_urls": [f"http://{ip}:{port}" for ip in _local_ipv4_addresses()] if enabled else [],
    }
    if include_token and is_local:
        payload["access_token"] = str(cfg.get("lan_access_token") or "")
    return payload

def resolve_class_config():
    """解析班级配置，返回 (class_name, class_folder_path, experiment_base_path)
    
    优先从 config.json 读取，支持不同班级/老师自定义。
    同时初始化全局 EXPERIMENT_BASE 变量。
    """
    global EXPERIMENT_BASE
    cfg = load_config_raw()
    class_name = cfg.get("class_name", _DEFAULT_CLASS_NAME)
    class_folder = Path(cfg.get("class_folder", _DEFAULT_CLASS_FOLDER))
    organized_dir = Path(cfg.get("organized_dir", str(class_folder / "已收作业")))
    
    # 公示/实验目录：优先从 config 读取，旧 config 无此字段时回退
    experiment_dir_str = cfg.get("experiment_dir", "")
    if experiment_dir_str:
        experiment_base = Path(experiment_dir_str)
    else:
        experiment_base = class_folder / "实验"
        print(f"[INFO] experiment_dir 未配置，回退使用: {experiment_base}")
    EXPERIMENT_BASE = experiment_base
    EXPERIMENT_BASE.mkdir(parents=True, exist_ok=True)
    
    return class_name, class_folder, organized_dir, experiment_base

def _assignment_ignore_keys(assignment):
    keys = []
    aid = str((assignment or {}).get("id", "")).strip()
    if aid:
        keys.append("id:" + aid)
    subject = str((assignment or {}).get("subject_group") or (assignment or {}).get("subject") or "").strip()
    experiment = str((assignment or {}).get("experiment", "")).strip()
    if subject or experiment:
        keys.append("pair:" + subject + "||" + experiment)
    return keys

def _norm_assignment_text(value):
    return str(value or "").strip()

def _assignment_matches_delete_payload(assignment, payload):
    assignment = assignment or {}
    payload = payload or {}
    aid = _norm_assignment_text(payload.get("id"))
    if aid and _norm_assignment_text(assignment.get("id")) == aid:
        return True

    subject = _canonical_subject_name(payload.get("subject_group") or payload.get("subject"), [assignment])
    experiment = _norm_assignment_text(payload.get("experiment"))
    name = _norm_assignment_text(payload.get("name"))
    assignment_subject = _norm_assignment_text(assignment.get("subject_group") or assignment.get("subject"))
    assignment_experiment = _norm_assignment_text(assignment.get("experiment"))
    assignment_name = _norm_assignment_text(assignment.get("name"))

    if subject and assignment_subject != subject:
        return False
    if subject and experiment and assignment_experiment == experiment:
        return True
    if subject and name and (assignment_name == name or assignment_experiment == name):
        return True
    if not subject and name and assignment_name == name:
        return True
    return False

def _next_manual_assignment_id(assignments):
    existing = []
    for assignment in assignments or []:
        aid = str((assignment or {}).get("id", ""))
        if aid.startswith("m") and aid[1:].isdigit():
            existing.append(int(aid[1:]))
    return f"m{max(existing, default=0) + 1}"

def _canonical_subject_name(subject, assignments=None):
    text = str(subject or "").strip()
    if not text:
        return ""
    existing = []
    for assignment in assignments or []:
        value = str((assignment or {}).get("subject_group") or (assignment or {}).get("subject") or "").strip()
        if value and value not in existing:
            existing.append(value)
    for value in existing:
        if value == text:
            return value
    for canonical, synonyms in _SUBJECT_SYNONYMS.items():
        if text == canonical or text in synonyms:
            return canonical
    lowered = text.lower()
    for canonical, synonyms in _SUBJECT_SYNONYMS.items():
        if lowered == canonical.lower() or any(lowered == str(s).lower() for s in synonyms):
            return canonical
    return text

def _assignment_is_ignored(assignment, ignored):
    return any(key in ignored for key in _assignment_ignore_keys(assignment))

def load_config(refresh_assignments=True):
    cfg = load_config_raw()
    # 确保 EXPERIMENT_BASE 已初始化
    if EXPERIMENT_BASE is None:
        resolve_class_config()
    if not refresh_assignments:
        return cfg
    # 公示/实验目录只用于回填文件与展示归档，不再默认反向生成作业。
    # 否则历史目录、临时验证目录或系统公示副本会突然变成前端科目卡片。
    if not cfg.get("auto_generate_assignments_from_experiment", False):
        return cfg
    # 从目录重建作业列表时，保留手动设置的字段
    auto_assignments = build_assignments_from_dir()
    if auto_assignments:
        old_assignments = list(cfg.get("assignments", []))
        ignored_subjects = set(cfg.get("ignored_subjects", []))
        ignored_assignments = set(cfg.get("ignored_assignments", []))
        # 按 (subject_group, experiment) 建立索引，用于去重自动生成的重复条目
        # 优先用已有手工 id：自动扫描的同名条目会复用旧 id 而非新增
        sg_exp_to_old = {}
        for oa in old_assignments:
            key = (oa.get("subject_group", ""), oa.get("experiment", ""))
            # 手工配置（experiment 字段非空）的优先级最高
            if oa.get("experiment") and key not in sg_exp_to_old:
                sg_exp_to_old[key] = oa.get("id", "")

        merged = []
        auto_ids = set()
        for a in auto_assignments:
            # 跳过已被用户显式删除的科目
            sg = a.get("subject_group", "")
            if sg and sg in ignored_subjects:
                continue
            if _assignment_is_ignored(a, ignored_assignments):
                continue
            # 去重：自动扫描的 (subject_group, experiment) 与手工配置重复 → 复用旧 id
            key = (a.get("subject_group", ""), a.get("experiment", ""))
            if key in sg_exp_to_old and sg_exp_to_old[key] and sg_exp_to_old[key] != a["id"]:
                old_id = sg_exp_to_old[key]
                # 把已知的 due/notes/completed/active 从旧记录同步过来
                old = next((oa for oa in old_assignments if oa.get("id") == old_id), {})
                a["id"] = old_id
                carry_fields = ["due", "notes", "completed", "active"]
                if str(old_id).startswith("m"):
                    carry_fields.extend(["name", "keywords"])
                for k in carry_fields:
                    if k in old:
                        a[k] = old[k]
                # 同步 id 后，从索引里移除（防止同一组被多次复用）
                del sg_exp_to_old[key]
            # 用 list 查找同 id 旧记录（保留手动设置的 due/notes/completed/active）
            old = next((oa for oa in old_assignments if oa.get("id") == a["id"]), {})
            # auto 只覆盖 4 个用户字段，不覆盖 name/subject_group/experiment/keywords
            for k in ("due", "notes", "completed", "active"):
                if k in old:
                    a[k] = old[k]
            merged.append(a)
            auto_ids.add(a["id"])
        # 保留手动添加的作业（m\d+ prefix，不在目录自动检测中）
        for old_a in old_assignments:
            if old_a.get("id") not in auto_ids and not _assignment_is_ignored(old_a, ignored_assignments):
                merged.append(old_a)
        cfg["assignments"] = merged
    return cfg

def save_config(cfg):
    save_json(CONFIG_PATH, _clean_config_paths(cfg))

def _clean_path_list(values):
    """清理配置中的路径列表，移除 None、空字符串和字符串 'None'。"""
    cleaned = []
    seen = set()
    if not isinstance(values, list):
        values = [values] if values else []
    for value in values:
        if value is None:
            continue
        text = str(value).strip()
        if not text or text.lower() == "none":
            continue
        key = text.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(text)
    return cleaned

def _clean_config_paths(cfg):
    """Normalize path-bearing config lists in-place and return cfg."""
    cfg["wechat_accounts"] = _clean_path_list(cfg.get("wechat_accounts", []))
    scan_dirs = _clean_path_list(cfg.get("scan_dirs", []))
    legacy_scan_dir = cfg.get("scan_dir")
    if legacy_scan_dir:
        scan_dirs.extend(_clean_path_list([legacy_scan_dir]))
    cfg["scan_dirs"] = _clean_path_list(scan_dirs)
    if "scan_dir" in cfg:
        cfg.pop("scan_dir", None)
    cfg["ui_theme"] = normalize_ui_theme(cfg.get("ui_theme"))
    cfg["ignored_subjects"] = [str(v).strip() for v in cfg.get("ignored_subjects", []) if str(v).strip()]
    cfg["ignored_assignments"] = [str(v).strip() for v in cfg.get("ignored_assignments", []) if str(v).strip()]
    if cfg.get("default_frontend") not in ("classic", "modern"):
        cfg["default_frontend"] = "classic"
    return cfg

def _safe_resolve_path(path_value):
    """Resolve a user-supplied path without throwing for bad input."""
    if not path_value:
        return None
    try:
        return Path(path_value).expanduser().resolve()
    except (OSError, RuntimeError, ValueError):
        return None

def _path_exists(path_value):
    resolved = _safe_resolve_path(path_value)
    return resolved.exists() if resolved else False

def _is_path_inside(path, root):
    try:
        path.relative_to(root)
        return True
    except ValueError:
        return False

def _safe_delete_experiment_dir(cfg):
    """Delete the configured public/experiment directory with guardrails."""
    target = _safe_resolve_path(cfg.get("experiment_dir"))
    class_root = _safe_resolve_path(cfg.get("class_folder") or _DEFAULT_CLASS_FOLDER)
    if not target:
        return False, "公示文件夹路径无效"
    if not target.exists():
        return True, "公示文件夹不存在，无需删除"
    if not target.is_dir():
        return False, "公示文件夹路径不是文件夹"
    if not class_root or not _is_path_inside(target, class_root):
        return False, "为避免误删，只允许删除班级文件夹内的公示文件夹"
    if target == class_root or target.parent == target:
        return False, "拒绝删除班级根目录或系统根目录"
    if target.name not in {"实验", "公示", "公示文件夹"} and "公示" not in target.name:
        return False, "为避免误删，只允许删除名称明确的公示/实验文件夹"
    shutil.rmtree(str(target))
    return True, "公示文件夹已删除"

def _system_output_roots(cfg=None):
    """Directories written by this app; realtime watcher must not process them."""
    cfg = cfg or load_config_raw()
    roots = [
        cfg.get("organized_dir"),
        cfg.get("experiment_dir"),
        str(ORGANIZED_DIR),
    ]
    if EXPERIMENT_BASE:
        roots.append(str(EXPERIMENT_BASE))
    return [r for r in (_safe_resolve_path(root) for root in roots) if r]

def _is_under_system_output(path_value, cfg=None):
    resolved = _safe_resolve_path(path_value)
    if not resolved:
        return False
    return any(_is_path_inside(resolved, root) for root in _system_output_roots(cfg))

def _allowed_file_roots():
    """Directories that may be exposed by file preview/open APIs."""
    cfg = load_config_raw()
    roots = [
        cfg.get("organized_dir"),
        cfg.get("experiment_dir"),
        CONVERT_TEMP_DIR,
        PREVIEW_CACHE_DIR,
    ]
    roots.extend(cfg.get("scan_dirs", []) or [])
    if cfg.get("scan_dir"):
        roots.append(cfg.get("scan_dir"))
    roots.extend(get_effective_watch_dirs(cfg))

    allowed = []
    seen = set()
    for root in roots:
        resolved = _safe_resolve_path(root)
        if not resolved or not resolved.exists() or not resolved.is_dir():
            continue
        key = str(resolved).lower()
        if key in seen:
            continue
        seen.add(key)
        allowed.append(resolved)
    return allowed

def is_allowed_file_path(path_value, require_file=True):
    """Return resolved path if it is under an approved file root."""
    resolved = _safe_resolve_path(path_value)
    if not resolved or not resolved.exists():
        return None
    if require_file and not resolved.is_file():
        return None
    if not require_file and not (resolved.is_file() or resolved.is_dir()):
        return None
    for root in _allowed_file_roots():
        if _is_path_inside(resolved, root):
            return resolved
    return None

def load_students():
    data = load_json(STUDENTS_PATH, [])
    if not isinstance(data, list):
        data = []
    return data

def save_students(data):
    save_json(STUDENTS_PATH, data)

def load_submissions():
    data = load_json(SUBMISSIONS_PATH, {})
    if not isinstance(data, dict):
        data = {}
    # 确保所有值都是 list
    for k, v in list(data.items()):
        if not isinstance(v, list):
            data[k] = []
    return data

def save_submissions(data):
    save_json(SUBMISSIONS_PATH, data)

def file_size(path):
    try:
        return path.stat().st_size if path.exists() else 0
    except OSError:
        return 0

def data_health():
    cfg = load_config(refresh_assignments=False)
    students = load_students()
    submissions = load_submissions()
    assignments = cfg.get("assignments", [])
    record_count = sum(len(records) for records in submissions.values() if isinstance(records, list))
    watch_dirs = get_effective_watch_dirs(cfg)
    watcher_running = False
    watcher_known = 0
    if watcher:
        try:
            watcher_running = watcher.running
            watcher_known = len(watcher.known_files)
        except Exception:
            pass
    class_name, class_folder, _, _ = resolve_class_config()
    return {
        "ok": True,
        "version": current_app_version(),
        "class_name": class_name,
        "class_folder": str(class_folder),
        "base_dir": str(BASE_DIR),
        "data_dir": str(DATA_DIR),
        "students": len(students),
        "assignments": len(assignments),
        "submission_records": record_count,
        "files": {
            "config_json": file_size(CONFIG_PATH),
            "students_json": file_size(STUDENTS_PATH),
            "submissions_json": file_size(SUBMISSIONS_PATH),
        },
        "watching": watcher_running,
        "known_files": watcher_known,
        "watch_dirs": [
            {"path": d, "exists": _path_exists(d)}
            for d in watch_dirs
        ],
        "scan_dirs": [
            {"path": d, "exists": _path_exists(d)}
            for d in cfg.get("scan_dirs", [])
        ],
    }

def process_is_running(pid):
    if not pid:
        return False
    if sys.platform == "win32":
        try:
            result = subprocess.run(
                ["tasklist", "/FI", f"PID eq {pid}"],
                capture_output=True,
                text=True,
                timeout=3,
            )
            return str(pid) in result.stdout
        except Exception:
            return False
    try:
        os.kill(pid, 0)
        return True
    except OSError:
        return False

def acquire_server_lock(port):
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    if LOCK_PATH.exists():
        try:
            old = json.loads(LOCK_PATH.read_text(encoding="utf-8"))
            old_pid = int(old.get("pid", 0))
        except Exception:
            old_pid = 0
        if old_pid and process_is_running(old_pid):
            print(f"\n[ERROR] 作业追踪器已经在运行（PID {old_pid}）。")
            print(f"请先打开 http://localhost:{port} 检查，或关闭旧命令行窗口后再启动。\n")
            return False
    LOCK_PATH.write_text(json.dumps({
        "pid": os.getpid(),
        "port": port,
        "started_at": datetime.now().isoformat(timespec="seconds"),
    }, ensure_ascii=False, indent=2), encoding="utf-8")
    atexit.register(release_server_lock)
    return True

def release_server_lock():
    try:
        if LOCK_PATH.exists():
            data = json.loads(LOCK_PATH.read_text(encoding="utf-8"))
            if int(data.get("pid", 0)) == os.getpid():
                LOCK_PATH.unlink()
    except Exception:
        pass

# ---------------------------------------------------------------------------
# 微信目录自动发现
# ---------------------------------------------------------------------------

def discover_wechat_accounts():
    """自动发现微信文件目录（支持新旧版本路径）"""
    accounts = []
    # 旧版路径: Documents\WeChat Files\wxid_xxx\FileStorage\File
    for base in (WECHAT_FILES_BASE, XWECHAT_BASE):
        if not base.exists():
            continue
        for d in base.iterdir():
            if not d.is_dir():
                continue
            # 旧版: d/FileStorage/File
            file_dir = d / "FileStorage" / "File"
            if file_dir.exists():
                accounts.append(str(d))
            # 新版 4.x: d/msg/file (如 xwechat_files\wxid_xxx_port\msg\file)
            msg_file_dir = d / "msg" / "file"
            if msg_file_dir.exists():
                accounts.append(str(d))
    return accounts

def get_watch_dirs():
    """获取需要监控的所有微信目录（兼容新旧版本路径）"""
    cfg = load_config_raw()
    accounts = _clean_path_list(cfg.get("wechat_accounts", []))
    if not accounts:
        accounts = discover_wechat_accounts()
        if accounts:
            cfg["wechat_accounts"] = accounts
            save_config(cfg)

    dirs = []
    now = datetime.now()
    for acct in accounts:
        acct_path = Path(acct)
        # 新版 4.x 路径: xwechat_files\wxid_port\msg\file\YYYY-MM
        new_style = acct_path / "msg" / "file" / now.strftime("%Y-%m")
        if new_style.exists():
            dirs.append(str(new_style))
        # 新版根目录也加入（处理旧月份）
        new_root = acct_path / "msg" / "file"
        if new_root.exists() and str(new_root) not in dirs:
            dirs.append(str(new_root))

        # 旧版路径: Documents\WeChat Files\wxid\FileStorage\File\YYYY-MM
        old_style = acct_path / "FileStorage" / "File" / now.strftime("%Y-%m")
        if old_style.exists():
            dirs.append(str(old_style))
        old_root = acct_path / "FileStorage" / "File"
        if old_root.exists() and str(old_root) not in dirs:
            dirs.append(str(old_root))

    return dirs

def _path_key(path_value):
    path_text = str(path_value or "").strip()
    if not path_text:
        return ""
    try:
        return str(Path(path_text).expanduser().resolve()).lower()
    except (OSError, RuntimeError, ValueError):
        return os.path.abspath(os.path.expandvars(os.path.expanduser(path_text))).lower()

def _current_wechat_month_dirs(accounts=None, now=None):
    """Return existing WeChat YYYY-MM directories for the current month."""
    accounts = _clean_path_list(accounts or [])
    now = now or datetime.now()
    month = now.strftime("%Y-%m")
    dirs = []
    seen = set()
    for acct in accounts:
        acct_path = Path(acct)
        candidates = [
            acct_path / "msg" / "file" / month,
            acct_path / "FileStorage" / "File" / month,
        ]
        for candidate in candidates:
            if not candidate.exists():
                continue
            key = _path_key(candidate)
            if key in seen:
                continue
            seen.add(key)
            dirs.append(str(candidate))
    return dirs

def sync_current_wechat_scan_dirs(cfg=None):
    """Persist newly-created monthly WeChat folders into scan_dirs."""
    cfg = cfg or load_config_raw()
    changed = False
    accounts = _clean_path_list(cfg.get("wechat_accounts", []))
    if not accounts:
        accounts = discover_wechat_accounts()
        if accounts:
            cfg["wechat_accounts"] = accounts
            changed = True

    scan_dirs = _clean_path_list(cfg.get("scan_dirs", []))
    known = {_path_key(d) for d in scan_dirs}
    for path_text in _current_wechat_month_dirs(accounts):
        key = _path_key(path_text)
        if key and key not in known:
            scan_dirs.append(path_text)
            known.add(key)
            changed = True

    if changed or scan_dirs != cfg.get("scan_dirs", []):
        cfg["scan_dirs"] = scan_dirs
        save_config(cfg)
    return cfg

def scan_dirs_payload(cfg=None):
    cfg = sync_current_wechat_scan_dirs(cfg)
    auto_dirs = get_watch_dirs()
    auto_keys = {_path_key(d) for d in auto_dirs}
    rows = []
    seen = set()

    def add_row(path_value, source):
        path_text = str(path_value or "").strip()
        key = _path_key(path_text)
        if not path_text or not key or key in seen:
            return
        seen.add(key)
        rows.append({
            "path": path_text,
            "exists": _path_exists(path_text),
            "source": "wechat-auto" if key in auto_keys or source == "wechat-auto" else "manual",
        })

    for d in cfg.get("scan_dirs", []) or []:
        add_row(d, "manual")
    for d in auto_dirs:
        add_row(d, "wechat-auto")
    if cfg.get("scan_dir"):
        add_row(cfg.get("scan_dir"), "manual")
    return rows

def get_effective_watch_dirs(cfg=None):
    """获取实际实时扫描目录：微信自动目录 + 用户配置的额外扫描目录。

    系统输出目录（已收作业、公示/实验目录）只用于展示和手动回填，不能进入实时 watcher。
    否则 organize_file 写出的归档副本会被下一轮轮询再次当作 [NEW] 处理。
    """
    if cfg is None:
        cfg = load_config_raw()

    dirs = []
    seen = set()
    output_roots = _system_output_roots(cfg)

    def add_dir(path_value):
        if not path_value:
            return
        path_text = str(path_value).strip()
        if not path_text or path_text.lower() == "none":
            return
        resolved = _safe_resolve_path(path_text)
        if resolved and any(_is_path_inside(resolved, root) for root in output_roots):
            return
        key = _path_key(path_text)
        if key in seen:
            return
        seen.add(key)
        dirs.append(path_text)

    for d in get_watch_dirs():
        add_dir(d)

    for d in cfg.get("scan_dirs", []) or []:
        add_dir(d)
    if cfg.get("scan_dir"):
        add_dir(cfg.get("scan_dir"))

    return dirs

# ---------------------------------------------------------------------------
# 文件扫描引擎
# ---------------------------------------------------------------------------

def scan_new_files(known_files):
    """扫描实时监听目录，返回新增的文件列表"""
    new_files = []
    cfg = load_config_raw()
    watch_dirs = get_effective_watch_dirs(cfg)
    keywords = cfg.get("file_keywords", ["作业", "报告", "论文"])
    doc_exts = configured_file_types(cfg)
    known_keys = {_path_key(p) for p in (known_files or [])}
    seen_this_scan = set()
    for wdir in watch_dirs:
        wpath = Path(wdir)
        if not wpath.exists():
            continue
        for f in wpath.rglob("*"):
            if f.is_file():
                fstr = str(f)
                fkey = _path_key(fstr)
                if fkey and fkey not in known_keys and fkey not in seen_this_scan:
                    # 防御：即使系统输出目录意外进入 watch_dirs，也不要处理系统自己归档/公示的副本。
                    if _is_under_system_output(f, cfg):
                        continue
                    # 跳过临时文件和缩略图
                    if f.name.startswith("~") or f.name.startswith("."):
                        continue
                    if f.suffix.lower() in (".tmp", ".crdownload", ".!ut"):
                        continue
                    # 只关心配置启用的文件类型
                    if f.suffix.lower() not in doc_exts:
                        continue
                    # 关键词过滤：只识别含"作业/报告/论文/实验/习题"等关键词的文件
                    if keywords and not any(kw in f.name for kw in keywords):
                        continue
                    seen_this_scan.add(fkey)
                    new_files.append({
                        "path": fstr,
                        "name": f.name,
                        "size": f.stat().st_size,
                        "mtime": f.stat().st_mtime,
                        "suffix": f.suffix.lower(),
                        "dir": str(f.parent),
                    })
    return new_files

# ---------------------------------------------------------------------------
# 文件匹配引擎
# ---------------------------------------------------------------------------

def match_file_to_student(filename, students):
    """根据文件名匹配学生，返回 (student_name, score)"""
    name_no_ext = Path(filename).stem.lower()
    best_match = None
    best_score = 0

    for s in students:
        name = s.get("name", "")
        if not name:
            continue

        # 直接包含匹配
        if name in filename or name in name_no_ext:
            score = len(name)  # 名字越长越精确
            if score > best_score:
                best_score = score
                best_match = name

        # 学号匹配
        sid = s.get("student_id", "")
        if sid and sid in filename:
            score = len(sid) + 10  # 学号匹配优先
            if score > best_score:
                best_score = score
                best_match = name

        # 拼音匹配（名字首字母）
        pinyin = s.get("pinyin", "")
        if pinyin and pinyin.lower() in name_no_ext:
            score = len(pinyin) + 5
            if score > best_score:
                best_score = score
                best_match = name

    return best_match, best_score

# ---------------------------------------------------------------------------
# 作业匹配引擎
# ---------------------------------------------------------------------------

def _fuzzy_contains(part, text):
    """检查 part 的所有字符是否按顺序出现在 text 中（中文缩略名模糊匹配）
    例如 "自控" 可匹配 "自动控制原理"（自→动→控→制，自和控顺序出现）"""
    if part in text:
        return True
    idx = 0
    for ch in part:
        idx = text.find(ch, idx)
        if idx == -1:
            return False
        idx += 1
    return True

def _detect_subjects_in_filename(filename):
    """从 filename 提取所有可能的学科 token（用 _SUBJECT_SYNONYMS 正向匹配）
    返回 set，如 {'自控'} 或 {'数电', '自控'}（如果文件名同时含两个学科）
    """
    fname = filename.lower() if filename else ""
    detected = set()
    for sg, syns in _SUBJECT_SYNONYMS.items():
        for syn in syns:
            if syn.lower() in fname:
                detected.add(sg)
                break
    return detected

def match_file_to_assignment(filename, assignments):
    """根据文件名匹配作业，返回 (assignment_id, score)。

    核心逻辑（修复 Bug1 + Bug2 的根治方案）：
    1. 先提取 filename 中的学科 token（用 _SUBJECT_SYNONYMS 正向匹配）
    2. 学科硬约束：如果 filename 能识别出学科（detected_subjects 非空），
       作业 a 的 subject_group 必须命中 detected_subjects，否则淘汰
       （防止"何智健自动控制原理报告"被误匹配到数电作业）
    3. 学科加分：filename 含 subject_group 同义词 → +30
    4. 作业名拆词、关键词、experiment 依次加分
    """
    if not filename or not assignments:
        return None, 0

    import re as _re
    fname = filename.lower()
    detected_subjects = _detect_subjects_in_filename(filename)

    best_id, best_score = None, 0
    for a in assignments:
        if not a.get("active", True):
            continue

        a_sg = a.get("subject_group", "")
        a_sg_lower = a_sg.lower()

        # 学科硬约束：如果 filename 能识别出学科 token，作业 a 的 subject_group 必须命中
        if detected_subjects and a_sg and a_sg not in detected_subjects:
            continue  # 明确不属本作业，淘汰

        # 学科硬约束 2：如果 filename 完全无法识别学科（detected_subjects 为空），
        # 但 a_sg 也不在 filename 里 → 文件不属于本作业
        # （防止"机械制图第一次作业"被其他作业名中的"第一次"单独误匹配）
        subject_in_fname = False
        if a_sg:
            syns = _get_subject_synonyms(a_sg)
            subject_in_fname = any(s.lower() in fname for s in syns)
        if not detected_subjects and a_sg and not subject_in_fname:
            # detected_subjects 为空 + subject_group 不在 filename → 拒绝
            continue

        score = 0

        # 作业名拆词匹配（保留原逻辑）
        name = a.get("name", "").lower()
        raw_parts = _re.findall(r'[\u4e00-\u9fff]+|[a-zA-Z]+|\d+', name)
        parts = [p for p in raw_parts if len(p) >= 2]
        exact_match = 0
        fuzzy_match = 0
        for p in parts:
            if p in fname:
                exact_match += 1
            elif _fuzzy_contains(p, fname):
                fuzzy_match += 1
        if exact_match or fuzzy_match:
            score = exact_match * 10 + fuzzy_match * 7

        # 学科加分：filename 含 subject_group 的同义词之一 → +30（防止张冠李戴）
        if a_sg and detected_subjects and a_sg in detected_subjects:
            score += 30
        elif a_sg and len(a_sg) >= 2 and a_sg_lower in fname:
            # 兼容旧逻辑：直接简称命中
            score += 30

        # 关键词加分（保留旧规则）
        for kw in a.get("keywords", []):
            kw_lower = kw.lower()
            if len(kw_lower) <= 1:
                continue
            if _re.match(r'^ex\d+$', kw_lower):
                continue
            if kw_lower in fname:
                score += 15

        # 负反馈词：用户在“非本作业”中指出的误导词条，命中则降权。
        for term in a.get("negative_keywords", []):
            term_lower = str(term or "").strip().lower()
            if len(term_lower) >= 2 and term_lower in fname:
                score -= 35

        # experiment 字段直接匹配加分
        exp = a.get("experiment", "")
        if exp and exp in filename:
            score += 20

        if score > best_score:
            best_score = score
            best_id = a["id"]
    return best_id, best_score

PENDING_ARCHIVE_BUCKET = "pending_archive"
_GENERIC_SUBJECT_GROUPS = {"课程作业", "课程报告", "项目作业", "课程论文", "实验报告", "课程设计", "小组作业"}

def classify_file_subject(file_info, assignments, cfg=None, source_kind="wechat"):
    """Stage one: identify a subject without using WeChat directory names."""
    cfg = cfg or load_config_raw()
    filename = str((file_info or {}).get("name", ""))

    if HAS_AI_CLASSIFIER and ai_settings(cfg).get("mode") == "rules":
        try:
            specific_assignments = [
                item for item in assignments
                if str(item.get("subject_group") or item.get("subject") or "").strip() not in _GENERIC_SUBJECT_GROUPS
            ]
            specific_synonyms = {
                name: aliases for name, aliases in _SUBJECT_SYNONYMS.items()
                if name not in _GENERIC_SUBJECT_GROUPS
            }
            result = ai_classifier.classify_subject(
                filename,
                assignments=specific_assignments,
                rules=load_ai_rules(),
                feedback=cfg.get("match_feedback", {}),
                subject_synonyms=specific_synonyms,
                students=load_students(),
                sensitivity=ai_settings(cfg).get("sensitivity", 0.70),
            )
            if result.get("status") in ("subject_matched", "subject_conflict", "subject_suggested"):
                return result
        except Exception as exc:
            print(f"[WARN] 分类大脑规则失败，使用原有规则：{exc}")

    detected = _detect_subjects_in_filename(filename) - _GENERIC_SUBJECT_GROUPS
    if len(detected) > 1:
        return {"status": "subject_conflict", "subject_group": "", "score": 0,
                "evidence": ["文件名同时命中多个科目"]}
    if len(detected) == 1:
        subject = next(iter(detected))
        return {"status": "subject_matched", "subject_group": subject, "score": 95,
                "evidence": [f"文件名命中科目：{subject}"]}

    for item in reversed((cfg.get("match_feedback", {}) or {}).get("subject_corrections", [])):
        token = str(item.get("token", "")).strip().lower()
        subject = str(item.get("to_subject", "")).strip()
        if len(token) >= 2 and subject and token in filename.lower():
            return {"status": "subject_matched", "subject_group": subject, "score": 70,
                    "evidence": [f"人工反馈匹配科目：{subject}"]}

    # Only public-folder backfill is allowed to use its relative directory.
    if source_kind == "public_backfill" and EXPERIMENT_BASE:
        try:
            rel = Path(file_info.get("path", "")).resolve().relative_to(Path(EXPERIMENT_BASE).resolve())
            known = {str(a.get("subject_group", "")).strip() for a in assignments}
            if rel.parts and rel.parts[0] in known:
                return {"status": "subject_matched", "subject_group": rel.parts[0], "score": 80,
                        "evidence": [f"公示目录命中科目：{rel.parts[0]}"]}
        except (OSError, ValueError):
            pass
    return {"status": "unmatched", "subject_group": "", "score": 0,
            "evidence": ["未识别到科目"]}

def classify_assignment_in_subject(filename, subject_group, assignments):
    """Stage two: compare only assignments in the confirmed subject."""
    candidates = []
    for assignment in assignments:
        if assignment.get("subject_group", "") != subject_group or not assignment.get("active", True):
            continue
        aid, score = match_file_to_assignment(filename, [assignment])
        candidates.append({"assignment_id": assignment.get("id", ""),
                           "name": assignment.get("name", ""),
                           "experiment": assignment.get("experiment", ""),
                           "score": score if aid else 0})
    candidates.sort(key=lambda item: item["score"], reverse=True)
    best = candidates[0] if candidates else None
    runner_score = candidates[1]["score"] if len(candidates) > 1 else -999
    if best and best["score"] >= 75 and best["score"] - runner_score >= 15:
        return {"status": "matched", "assignment_id": best["assignment_id"],
                "score": best["score"], "candidates": candidates[:3],
                "evidence": [f"科目内作业匹配：{best['name'] or best['experiment']}" ]}
    reason = "未发现明确作业信息" if not best or best["score"] < 50 else "作业候选分数接近，等待确认"
    return {"status": "assignment_pending", "assignment_id": "", "score": best["score"] if best else 0,
            "candidates": candidates[:3], "evidence": [reason]}

def _can_archive_record(record):
    return record.get("status") in ("matched", "manual_matched") and bool(record.get("assignment_id"))

def _record_matches_assignment(r, a):
    """统一判断一条 submission 记录是否属于作业 a（列表/详情共用入口）。

    优先级（修复 Bug1 列表/详情口径不一致 + Bug2 全称识别）：
      0. 学生名缺失 → False
      1. r["assignment_id"] == a["id"] 且学科 token 不冲突
         - 如果 filename 能识别出学科 token，必须命中 a_sg（否则脏数据）
         - 如果 filename 无法识别出学科 token（detected_subjects 为空），
           但 a_sg 也不在 filename 中 → 视为可疑，需 filename 走 match_file_to_assignment
      2. organized_to 路径匹配（is_file_in_correct_path，含同义词）
      3. filename 走 match_file_to_assignment 重新匹配（学科硬约束已修）
      4. 否则 False
    """
    a_id = a.get("id", "")
    if not a_id:
        return False

    sn = r.get("student", "")
    if not sn:
        return False

    a_sg = a.get("subject_group", "")
    a_exp = a.get("experiment", "")
    fname = (r.get("file") or {}).get("name", "")
    org_to = r.get("organized_to") or ""

    # 0. filename 学科 token 检测（用于脏数据过滤）
    detected_subjects = _detect_subjects_in_filename(fname) if fname else set()

    # 1. 桶匹配 + 学科不冲突
    if r.get("assignment_id") == a_id:
        if detected_subjects:
            # filename 能识别出学科，必须命中 a_sg，否则视为脏数据
            if a_sg and a_sg not in detected_subjects:
                pass  # 学科冲突，继续走后续校验
            else:
                return True
        else:
            # filename 无法识别学科（极端情况），但桶 a_id 命中 → 仍信任
            # （历史已有正确归类的数据，避免误伤）
            return True

    # 2. 路径匹配
    if org_to and is_file_in_correct_path(org_to, a_sg, a_exp):
        return True

    # 3. filename 匹配（match_file_to_assignment 已含学科硬约束 + 同义词加分）
    if fname:
        aid_match, _ = match_file_to_assignment(fname, [a])
        if aid_match == a_id:
            return True

    return False

# ---------------------------------------------------------------------------
# 文件整理
# ---------------------------------------------------------------------------

# 学科同义词表（唯一真值表 - Single Source of Truth）
# 格式: subject_group → [所有可能的写法：全称 / 简称 / 英文 / 俗名]
# key 是规范的 subject_group（用于 a["subject_group"] 和目录名），value 中所有元素视为同一学科
_SUBJECT_SYNONYMS = {
    "课程作业": ["课程作业", "作业", "阶段作业", "平时作业"],
    "课程报告": ["课程报告", "报告", "学习报告", "调研报告"],
    "项目作业": ["项目作业", "项目", "项目报告", "小组项目"],
    "课程论文": ["课程论文", "论文", "期末论文"],
    "实验报告": ["实验报告", "实验", "实训报告", "实践报告"],
    "课程设计": ["课程设计", "课设", "设计报告", "综合设计"],
    "小组作业": ["小组作业", "小组报告", "小组展示", "汇报"],
    # 电气/工科课程示例词保留为兼容识别，不作为默认安装主语境。
    "数字电子技术": ["数字电子技术", "数电", "数字电路", "数字逻辑", "数字电子"],
    "程序设计": ["程序设计", "程序设计基础", "C语言", "Python", "编程"],
    "信号与系统": ["信号与系统", "信号", "通信原理"],
    "自动控制原理": ["自动控制原理", "自控", "自控原理", "自动控制", "控制原理"],
    "单片机": ["单片机", "单片机原理", "微机原理", "嵌入式", "微控制器"],
    "电机学": ["电机学", "电机", "电动机", "电机拖动"],
    "电气导论": ["电气导论", "电气工程导论", "导论"],
}

# 兼容 alias：旧 SUBJECT_RULES 调用方不变
# 旧格式: (subject, keywords)，这里用同义词列表代替
SUBJECT_RULES = [(sg, syns) for sg, syns in _SUBJECT_SYNONYMS.items()]

# 学科全称 → 简称映射（保留旧 _SUBJECT_SHORT_MAP 兼容）
_SUBJECT_SHORT_MAP = {
    full: sg for sg, syns in _SUBJECT_SYNONYMS.items() for full in syns if full != sg
}

def auto_detect_assignments():
    """扫描实验目录，自动生成按科目分组的作业列表"""
    subjects = {}
    if not EXPERIMENT_BASE.exists():
        return []
    
    # 扫描顶层科目目录
    for subject_dir in sorted(EXPERIMENT_BASE.iterdir()):
        if not subject_dir.is_dir():
            continue
        subject_name = subject_dir.name
        # 跳过非实验目录
        if "报告" in subject_name or "文档" in subject_name or "zip" in subject_name.lower():
            continue
        
        experiments = []
        for exp_dir in subject_dir.iterdir():
            if not exp_dir.is_dir():
                continue
            if "报告" in exp_dir.name or "zip" in exp_dir.name.lower():
                continue
            exp_name = exp_dir.name  # "第一次", "第二次", "第三次"
            experiments.append(exp_name)
        # 按实验次数逻辑顺序排列
        experiments.sort(key=lambda x: _EXPERIMENT_ORDER.get(x, 99))
        
        if experiments:
            subjects[subject_name] = experiments
    
    return subjects

def build_assignments_from_dir():
    """从目录结构构建 assignments 列表"""
    subjects = auto_detect_assignments()
    if not subjects:
        # 回退到默认（不要递归调用 load_config！）
        return [{"id": "a1", "name": "默认作业", "subject_group": "默认", "experiment": "", "due": "", "notes": "", "active": True, "keywords": []}]

    # 实验次数映射：中文 → ex格式
    EXP_NUMBERS = {
        "第一次": ("一", "ex1"),
        "第二次": ("二", "ex2"),
        "第三次": ("三", "ex3"),
        "第四次": ("四", "ex4"),
        "第五次": ("五", "ex5"),
    }
    
    # 从现有 config 找最大 a\d+ id，避免与已有数据撞车
    cfg_raw = load_config_raw()
    existing_a = [int(a["id"][1:]) for a in cfg_raw.get("assignments", []) if a.get("id","").startswith("a") and a["id"][1:].isdigit()]
    next_a = max(existing_a, default=0) + 1
    
    assignments = []
    for subject, exps in subjects.items():
        for exp in exps:
            nums = EXP_NUMBERS.get(exp, (exp.replace("第","").replace("次",""),))
            keywords = [subject] + list(nums)
            assignments.append({
                "id": f"a{next_a}",
                "name": f"{subject}{exp}",
                "subject_group": subject,
                "experiment": exp,
                "due": "",
                "notes": "",
                "active": True,
                "completed": False,
                "keywords": keywords,
            })
            next_a += 1
    return assignments

def detect_subject(filename):
    """根据文件名检测科目"""
    for subject, keywords in SUBJECT_RULES:
        for kw in keywords:
            if kw.lower() in filename.lower():
                return subject
    return "其他"

def get_subject_experiment_for_file(file_name, assignment_id=None):
    """根据文件名和作业ID获取 (subject_name, experiment_no) 用于目录归类
    优先使用 assignment_id 精确匹配，确保文件存到正确的科目/次数目录"""
    cfg = load_config_raw()
    assignments = cfg.get("assignments", build_assignments_from_dir())
    
    # 优先：如果传入了 assignment_id，直接从 assignments 中获取精确的科目和实验次数
    if assignment_id:
        for a in assignments:
            if a.get("id") == assignment_id:
                sg = a.get("subject_group", "其他")
                exp = a.get("experiment", "")
                if sg and exp:
                    return sg, exp
                # 如果 assignment 中没有 experiment，尝试从文件名提取
                break
    
    # 从文件名中识别科目
    subject = detect_subject(file_name)

    # 从文件名中识别实验次数（使用统一的公共函数）
    experiment = parse_experiment_number(file_name)
    
    # 如果从 assignment_id 已经获取了 subject_group，但 experiment 为空，用文件名提取的补充
    if assignment_id:
        for a in assignments:
            if a.get("id") == assignment_id:
                sg = a.get("subject_group", subject)
                exp = experiment or a.get("experiment", "")
                return sg, exp

    return subject, experiment


def _sync_experiment_dir(subject_group: str, experiment: str):
    """新建作业时同步在公示/实验目录创建空目录（幂等）"""
    if not load_config_raw().get("experiment_enabled", False):
        return
    if not subject_group or not EXPERIMENT_BASE:
        return
    try:
        target = (EXPERIMENT_BASE / subject_group / experiment) if experiment else (EXPERIMENT_BASE / subject_group)
        target.mkdir(parents=True, exist_ok=True)
    except Exception as e:
        print(f"[WARN] 同步公示目录失败 {EXPERIMENT_BASE}/{subject_group}/{experiment}: {e}")


def _is_under_path(path: Path, root: Path) -> bool:
    try:
        path.resolve().relative_to(root.resolve())
        return True
    except Exception:
        return False


def _collapse_repeated_student_prefix(filename: str, student_name: str) -> str:
    """把 张三_张三_文件.docx 规范成 张三_文件.docx。"""
    if not student_name:
        return filename
    p = Path(filename)
    prefix = f"{student_name}_"
    stem = p.stem
    count = 0
    while stem.startswith(prefix):
        count += 1
        stem = stem[len(prefix):]
    if count <= 1:
        return filename
    return f"{student_name}_{stem}{p.suffix}"


def _submission_compare_name(filename: str, student_name: str) -> str:
    """提交去重用文件名：折叠重复学生名前缀，并忽略一个系统生成的学生名前缀。"""
    text = str(filename or "").strip()
    if not text:
        return ""
    collapsed = _collapse_repeated_student_prefix(text, student_name)
    if not student_name:
        return collapsed.casefold()
    p = Path(collapsed)
    prefix = f"{student_name}_"
    stem = p.stem
    if stem.startswith(prefix):
        stem = stem[len(prefix):]
    return f"{stem}{p.suffix}".casefold()


def _submission_dedupe_key_from_parts(student_name, assignment_id, filename, size):
    """统一提交去重键：学生 + 作业 + 规范化文件名 + 文件大小。"""
    student = str(student_name or "").strip()
    aid = str(assignment_id or "").strip()
    name_key = _submission_compare_name(filename, student)
    try:
        size_key = int(size or 0)
    except (TypeError, ValueError):
        size_key = str(size or "")
    return (student, aid, name_key, size_key)


def _submission_record_dedupe_key(record, assignment_id=None):
    """从 submission 记录生成统一去重键。"""
    if not isinstance(record, dict):
        return None
    file_info = record.get("file") or {}
    return _submission_dedupe_key_from_parts(
        record.get("student"),
        assignment_id or record.get("assignment_id"),
        file_info.get("name", ""),
        file_info.get("size", 0),
    )


def _submission_record_score(record, cfg=None):
    """重复记录保留优先级：优先保留微信源文件，其次保留信息更完整的记录。"""
    if not isinstance(record, dict):
        return -1
    cfg = cfg or load_config_raw()
    file_path = (record.get("file") or {}).get("path", "")
    score = 0
    if file_path and not _is_under_system_output(file_path, cfg):
        score += 100
    if record.get("status") in ("matched", "manual", "confirmed", "existing"):
        score += 20
    if record.get("organized_to"):
        score += 5
    if (record.get("file") or {}).get("mtime"):
        score += 3
    if record.get("detected_at"):
        score += 1
    return score


def _prefer_submission_record(candidate, current, cfg=None):
    """同一去重键下是否用 candidate 替换 current。"""
    return _submission_record_score(candidate, cfg) > _submission_record_score(current, cfg)


def dedupe_submission_records(submissions, cfg=None):
    """原地清理 submissions 中的重复提交记录，返回移除数量。"""
    if not isinstance(submissions, dict):
        return 0
    cfg = cfg or load_config_raw()
    removed = 0
    for aid, records in list(submissions.items()):
        if not isinstance(records, list):
            continue
        seen = {}
        kept = []
        for record in records:
            key = _submission_record_dedupe_key(record, aid)
            if not key or not key[0] or not key[2]:
                kept.append(record)
                continue
            existing_index = seen.get(key)
            if existing_index is None:
                seen[key] = len(kept)
                kept.append(record)
                continue
            existing = kept[existing_index]
            if _prefer_submission_record(record, existing, cfg):
                kept[existing_index] = record
            removed += 1
        if removed and len(kept) != len(records):
            submissions[aid] = kept
    return removed


def _append_submission_delete_log(cfg, record, assignment_id, reason):
    """记录人工删除提交的理由，最多保留最近 500 条。"""
    file_info = (record or {}).get("file") or {}
    assignments = cfg.get("assignments", []) if isinstance(cfg, dict) else []
    assignment = next((a for a in assignments if a.get("id") == assignment_id), {})
    assignment_label = " · ".join([
        str(assignment.get("subject_group") or assignment.get("subject") or "").strip(),
        str(assignment.get("experiment") or assignment.get("name") or "").strip(),
    ]).strip(" ·") or assignment_id
    logs = cfg.setdefault("submission_delete_log", [])
    if not isinstance(logs, list):
        logs = []
    logs.append({
        "file_name": file_info.get("name", ""),
        "file_path": file_info.get("path", ""),
        "student": (record or {}).get("student"),
        "assignment_id": assignment_id,
        "assignment_label": assignment_label,
        "reason": str(reason or "手动删除").strip() or "手动删除",
        "deleted_at": datetime.now().isoformat(),
    })
    cfg["submission_delete_log"] = logs[-500:]


def _experiment_dest_name(filename: str, student_name: str) -> str:
    """生成公示目录文件名：最多保留一个学生名前缀。"""
    collapsed = _collapse_repeated_student_prefix(filename, student_name)
    if student_name and not collapsed.startswith(f"{student_name}_"):
        return f"{student_name}_{collapsed}"
    return collapsed


def _copy_to_experiment(src: Path, subject: str, experiment: str, dest_name: str) -> bool:
    """复制文件到公示/实验目录（学生文件平铺），异常返回 False 不阻断"""
    if not load_config_raw().get("experiment_enabled", False):
        return True
    if not EXPERIMENT_BASE:
        return False
    try:
        # 公示目录回填时源文件本来就在公示目录内，不能再复制回自身，否则会生成重复前缀文件。
        if _is_under_path(src, EXPERIMENT_BASE):
            return True
        sub = (EXPERIMENT_BASE / subject / experiment) if experiment else (EXPERIMENT_BASE / subject)
        sub.mkdir(parents=True, exist_ok=True)
        dest = sub / dest_name
        if not dest.exists() or dest.stat().st_size != src.stat().st_size:
            shutil.copy2(str(src), str(dest))
        return True
    except Exception as e:
        print(f"[WARN] 公示复制失败 {src} -> {EXPERIMENT_BASE}: {e}")
        return False


def _auto_backfill_from_experiment():
    """server 启动时自动把公示目录未入库文件回填到已收作业 + 同步到公示"""
    if not load_config_raw().get("experiment_enabled", False):
        print("[INFO] 公示文件夹未启用，跳过启动回填")
        return
    if not EXPERIMENT_BASE or not EXPERIMENT_BASE.exists():
        return
    start = time.time()
    print(f"[INFO] 启动回填扫描: {EXPERIMENT_BASE}")
    result = scan_existing_directory(str(EXPERIMENT_BASE), filename_student=True)
    elapsed = time.time() - start
    matched = result.get("matched", 0)
    print(f"[INFO] 启动回填完成: {matched} 个文件, 耗时 {elapsed:.2f}s")


def organize_file(file_info, student_name, assignment_id, sync_experiment=True):
    """将文件复制到整理目录，按 科目/第X次/学生姓名 结构存放"""
    cfg = load_config()
    if not cfg.get("auto_organize"):
        return None

    dest_base = Path(cfg.get("organized_dir", str(ORGANIZED_DIR)))
    subject, experiment = get_subject_experiment_for_file(file_info["name"], assignment_id)
    
    if experiment:
        sub_dir = dest_base / subject / experiment / student_name
    else:
        sub_dir = dest_base / subject / student_name
    sub_dir.mkdir(parents=True, exist_ok=True)

    src = Path(file_info["path"])
    dest = sub_dir / src.name

    # 如果目标已存在且内容相同（大小一致），直接返回，不产生副本
    if dest.exists():
        try:
            if dest.stat().st_size == src.stat().st_size:
                # 快速检查：大小相同大概率是同一文件，跳过复制
                return str(dest)
        except OSError:
            pass
        # 大小不同才加时间戳避免覆盖
        stem = dest.stem
        ext = dest.suffix
        dest = sub_dir / f"{stem}_{datetime.now().strftime('%H%M%S')}{ext}"

    try:
        shutil.copy2(src, dest)
        if sync_experiment:
            # 双写到公示/实验目录（学生文件平铺，文件名加学生名前缀，已含前缀时去重）
            exp_dest_name = _experiment_dest_name(src.name, student_name)
            ok = _copy_to_experiment(src, subject, experiment, exp_dest_name)
            if not ok:
                print(f"[WARN] 公示目录写入失败（不阻断）: {src.name} -> {EXPERIMENT_BASE}")
        return str(dest)
    except Exception as e:
        print(f"[ERROR] Copy failed: {src} -> {dest}: {e}")
        return None

# ---------------------------------------------------------------------------
# 通用 ZIP 打包工具
# ---------------------------------------------------------------------------

def _create_zip(entries, compress=True):
    """通用 ZIP 打包函数
    
    Args:
        entries: [(file_path, arcname), ...] 或 {arcname: file_path}
                 每个条目为 (源文件路径, ZIP内部路径)
        compress: 是否启用压缩（默认 True）
    
    Returns:
        bytes: ZIP 文件的二进制数据
    
    Example:
        data = _create_zip([
            ("/path/to/file1.docx", "科目/第一次/学生_文件.docx"),
            ("/path/to/file2.pdf", "科目/第二次/学生_文件.pdf"),
        ])
    """
    import zipfile
    import io
    
    if isinstance(entries, dict):
        entries = [(path, arcname) for arcname, path in entries.items()]
    
    buf = io.BytesIO()
    compression = zipfile.ZIP_DEFLATED if compress else zipfile.ZIP_STORED
    with zipfile.ZipFile(buf, "w", compression) as zf:
        for file_path, arcname in entries:
            fp = Path(file_path)
            if fp.exists() and fp.is_file():
                zf.write(str(fp), arcname)
    return buf.getvalue()

# ---------------------------------------------------------------------------
# 已有文件批量扫描
# ---------------------------------------------------------------------------

def scan_existing_directory(base_path, filename_student=False):
    """扫描已有目录，匹配学生和作业，去重后导入提交记录
    
    当 filename_student=True 时（公示目录回填），额外尝试从文件名前缀提取学生名
    （如 杨焱煜_report.docx → student="杨焱煜"）
    """
    cfg = load_config()
    students = load_students()
    assignments = cfg.get("assignments", build_assignments_from_dir())
    report_exts = configured_file_types(cfg)
    if not assignments:
        assignments = build_assignments_from_dir()
    submissions = load_submissions()

    scanned = 0
    matched = 0
    deleted = 0
    skipped_no_student = 0
    base = Path(base_path)
    if not base.exists():
        return {"error": f"目录不存在: {base_path}", "scanned": 0, "matched": 0, "deleted": 0}

    # 第一遍：收集所有文件
    all_files = []
    for f in base.rglob("*"):
        if not f.is_file():
            continue
        # 黑名单：任意层级"其他"子目录跳过
        if any(p.name == "其他" for p in f.parents):
            continue
        if f.suffix.lower() not in report_exts:
            continue
        if f.name.startswith("~") or f.name.startswith("."):
            continue
        scanned += 1
        all_files.append(f)

    # 第二遍：匹配学生+作业（用目录路径辅助），同时去重
    # key = (student_name, assignment_id, filename_lower)
    seen = {}
    for f in all_files:
        student_name, score = match_file_to_student(f.name, students)

        # 文件名匹配不到学生时，尝试从目录名匹配（如 电机/第二次/杨焱煜/xxx.docx）
        if not student_name:
            try:
                rel = f.relative_to(base)
                for part in rel.parts:
                    for s in students:
                        if s["name"] == part:
                            student_name = part
                            score = len(part) + 5
                            break
                    if student_name:
                        break
            except:
                pass

        # 公示目录回填：从文件名前缀提取学生名（如 杨焱煜_report.docx）
        if not student_name and filename_student:
            try:
                prefix = f.stem.split("_")[0] if "_" in f.stem else ""
                if prefix:
                    for s in students:
                        if s["name"] == prefix:
                            student_name = prefix
                            score = len(prefix) + 5
                            break
            except:
                pass

        if not student_name:
            skipped_no_student += 1
            # 不跳过，仍记录为未匹配，让用户可以在未归类列表中看到
            student_name = ""  # 空 = 未匹配
            score = 0

        # 根据目录路径推断科目和实验
        assignment_id = None
        subject_dir = ""
        exp_dir = ""
        try:
            rel = f.relative_to(base)
            parts = rel.parts
            exp_nums = ("第一次","第二次","第三次","第四次","第五次")
            if len(parts) >= 2:
                exp_dir = parts[-2] if parts[-2] in exp_nums else ""
                subject_dir = parts[-3] if len(parts) >= 3 and exp_dir else ""
                if not subject_dir and len(parts) >= 3:
                    exp_dir2 = parts[-3] if parts[-3] in exp_nums else ""
                    subject_dir2 = parts[-4] if len(parts) >= 4 and exp_dir2 else ""
                    if subject_dir2 and exp_dir2:
                        subject_dir, exp_dir = subject_dir2, exp_dir2
                if subject_dir and exp_dir:
                    for a in assignments:
                        if a.get("subject_group") == subject_dir and a.get("experiment") == exp_dir:
                            assignment_id = a["id"]
                            break
                # PDF实验一/二/三 映射
                if not assignment_id:
                    pdf_map = {"一":"第一次","二":"第二次","三":"第三次"}
                    for pdf_ch, exp_name in pdf_map.items():
                        if f"PDF实验{pdf_ch}" in str(rel):
                            subj = parts[-4] if len(parts) >= 4 else ""
                            for a in assignments:
                                if a.get("subject_group") == subj and a.get("experiment") == exp_name:
                                    assignment_id = a["id"]
                                    break
                            break
        except:
            pass

        if not assignment_id:
            subject_result = classify_file_subject({"path": str(f), "name": f.name}, assignments, cfg,
                                                   source_kind="public_backfill" if filename_student else "manual")
            if subject_result.get("status") == "subject_matched":
                assignment_result = classify_assignment_in_subject(f.name, subject_result.get("subject_group", ""), assignments)
                assignment_id = assignment_result.get("assignment_id", "")
            else:
                assignment_result = {"status": subject_result.get("status", "unmatched"), "score": 0,
                                     "candidates": [], "evidence": subject_result.get("evidence", [])}
        else:
            assignment = next((a for a in assignments if a.get("id") == assignment_id), {})
            subject_result = {"status": "subject_matched", "subject_group": assignment.get("subject_group", ""),
                              "score": 100, "evidence": ["目录匹配科目和作业"]}
            assignment_result = {"status": "matched", "assignment_id": assignment_id, "score": 100,
                                 "candidates": [], "evidence": ["目录匹配具体作业"]}
        status = assignment_result.get("status", "unmatched")
        bucket_id = assignment_id if assignment_id else PENDING_ARCHIVE_BUCKET

        # 去重: 同学生+同作业+同名文件 → 保留较新的。
        # 公示目录回填时把重复学生名前缀折叠，避免 张三_张三_文件 与 张三_文件 被当作不同文件。
        canonical_name = _collapse_repeated_student_prefix(f.name, student_name) if filename_student else f.name
        dup_key = (
            _submission_dedupe_key_from_parts(student_name, assignment_id or PENDING_ARCHIVE_BUCKET, f.name, f.stat().st_size)
            if filename_student
            else (student_name, assignment_id, canonical_name.lower(), f.stat().st_size)
        )
        if dup_key in seen:
            existing_f = seen[dup_key]
            existing_canonical = _collapse_repeated_student_prefix(existing_f.name, student_name) if filename_student else existing_f.name
            prefer_current = (
                f.name == canonical_name and existing_f.name != existing_canonical
            ) or (
                f.name == canonical_name and existing_f.name == existing_canonical and f.stat().st_mtime > existing_f.stat().st_mtime
            ) or (
                f.name != canonical_name and existing_f.name != existing_canonical and f.stat().st_mtime > existing_f.stat().st_mtime
            )
            if prefer_current:
                try:
                    existing_f.unlink()
                    deleted += 1
                except:
                    pass
                seen[dup_key] = f
            else:
                try:
                    f.unlink()
                    deleted += 1
                except:
                    pass
            continue

        seen[dup_key] = f

        # 检查已在 submissions 中的去重
        existing_records = submissions.get(bucket_id, [])
        new_record_key = _submission_dedupe_key_from_parts(student_name, assignment_id or PENDING_ARCHIVE_BUCKET, f.name, f.stat().st_size)
        if any(
            _submission_record_dedupe_key(r, assignment_id) == new_record_key
            for r in existing_records
        ):
            continue

        matched += 1
        status = "existing" if status == "matched" and student_name else status
        file_info = {
            "path": str(f),
            "name": f.name,
            "size": f.stat().st_size,
            "suffix": f.suffix.lower(),
            "dir": str(f.parent),
        }
        record = {
            "file": file_info,
            "student": student_name or None,
            "match_score": score,
            "detected_at": datetime.now().isoformat(),
            "assignment_id": assignment_id,
            "assignment_name": next((a["name"] for a in assignments if a["id"] == assignment_id), "其他"),
            "status": status,
            "subject_group": subject_result.get("subject_group", ""),
            "classification": {"stage": status, "subject_group": subject_result.get("subject_group", ""),
                               "subject_score": subject_result.get("score", 0), "assignment_id": assignment_id,
                               "assignment_score": assignment_result.get("score", 0),
                               "candidates": assignment_result.get("candidates", []),
                               "evidence": subject_result.get("evidence", []) + assignment_result.get("evidence", []),
                               "source_kind": "public_backfill" if filename_student else "manual"},
        }
        # 匹配到学生 → 复制到 organized_dir 供仪表盘展示
        if student_name and assignment_id and status in ("existing", "matched", "manual_matched"):
            dest = organize_file(file_info, student_name, assignment_id, sync_experiment=not filename_student)
            record["organized_to"] = dest
        # 附加上下文：来自哪个目录结构
        if not student_name and (subject_dir or exp_dir):
            hint = []
            if subject_dir: hint.append(subject_dir)
            if exp_dir: hint.append(exp_dir)
            record["context_hint"] = " ".join(hint)
        submissions.setdefault(bucket_id, []).append(record)

    deduped_records = dedupe_submission_records(submissions, cfg)
    save_submissions(submissions)
    if deduped_records:
        print(f"[ScanExisting] removed {deduped_records} duplicate submission records")
    print(f"[ScanExisting] scanned={scanned}, matched={matched}, deleted={deleted}, skipped_no_student={skipped_no_student}")
    return {
        "scanned": scanned,
        "matched": matched,
        "deleted": deleted,
        "deduped_records": deduped_records,
        "skipped_no_student": skipped_no_student,
    }

# ---------------------------------------------------------------------------
# Docx → PDF 批量转换
# ---------------------------------------------------------------------------

def convert_docx_to_pdf(docx_paths):
    """把 docx 文件批量转为 PDF（线程池+超时保护），返回转换结果列表"""
    import concurrent.futures

    def _convert_one(src, dst):
        with _WordComContext() as ctx:
            doc = ctx.open(src)
            doc.ExportAsFixedFormat(str(dst), 17)
        return dst.exists() and dst.stat().st_size > 0

    results = []
    tasks = []

    for src_str in docx_paths:
        src = Path(src_str)
        if not src.exists():
            results.append({"path": src_str, "status": "not_found"})
            continue
        if src.suffix.lower() != ".docx":
            results.append({"path": src_str, "status": "skip"})
            continue
        dst = src.with_suffix(".pdf")
        results.append({"path": src_str, "pdf_path": str(dst), "status": "pending"})
        tasks.append((src, dst, len(results) - 1))

    if not tasks:
        return results

    with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
        futures = {executor.submit(_convert_one, src, dst): idx for src, dst, idx in tasks}
        for f in concurrent.futures.as_completed(futures):
            idx = futures[f]
            try:
                ok = f.result(timeout=45)
                results[idx]["status"] = "ok" if ok else "error"
                if not ok:
                    results[idx]["msg"] = "PDF 输出为空"
            except concurrent.futures.TimeoutError:
                results[idx]["status"] = "error"
                results[idx]["msg"] = "转换超时（Word 无响应）"
            except Exception as e:
                results[idx]["status"] = "error"
                results[idx]["msg"] = str(e)[:200]

    return results

# ---------------------------------------------------------------------------
# 文件文本提取（用于分析和预览）
# ---------------------------------------------------------------------------

def read_docx_text(file_path):
    """用 python-docx 提取 .docx 全文"""
    if not HAS_DOCX:
        return None, "python-docx 未安装，无法读取 .docx 内容"
    try:
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs), None
    except Exception as e:
        return None, f"读取 .docx 失败: {str(e)[:200]}"

def read_pdf_text(file_path):
    """用 PyPDF2 提取 .pdf 全文"""
    if not HAS_PDF:
        return None, "PyPDF2 未安装，无法读取 .pdf 内容"
    try:
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages[:50]:  # 最多读 50 页
            t = page.extract_text()
            if t:
                pages.append(t)
        return "\n".join(pages), None
    except Exception as e:
        return None, f"读取 .pdf 失败: {str(e)[:200]}"

class _WordComContext:
    """Word COM 上下文管理器，统一处理 COM 初始化/清理"""
    def __enter__(self):
        import pythoncom
        pythoncom.CoInitialize()
        import win32com.client
        self.word = win32com.client.Dispatch("Word.Application")
        self.word.Visible = False
        self.word.DisplayAlerts = 0
        self.doc = None
        return self

    def open(self, file_path):
        self.doc = self.word.Documents.Open(str(file_path), ReadOnly=True)
        return self.doc

    def __exit__(self, *args):
        if self.doc:
            try: self.doc.Close(SaveChanges=False)
            except: pass
        try: self.word.Quit()
        except: pass
        import pythoncom
        pythoncom.CoUninitialize()
        return False  # 不吞异常


def read_doc_text_via_word(file_path):
    """通过 Word COM 提取 .doc 文件文本（最可靠）"""
    try:
        with _WordComContext() as ctx:
            doc = ctx.open(file_path)
            return doc.Content.Text, None
    except Exception as e:
        return None, f"Word COM 提取失败: {str(e)[:200]}"

def read_doc_text(file_path):
    """尝试读取 .doc（旧版 Word 二进制格式）"""
    # 策略1：有些 .doc 实际上就是 .docx 格式，先用 python-docx 尝试
    if HAS_DOCX:
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if paragraphs:
                return "\n".join(paragraphs), None
        except Exception:
            pass

    # 策略2：Word COM 精确提取（用于分析）
    text, err = read_doc_text_via_word(file_path)
    if text and text.strip():
        return text, None

    # 策略3：二进制明文提取（回退）
    try:
        with open(file_path, "rb") as f:
            raw = f.read()
        try:
            text = raw.decode("gbk", errors="ignore")
        except Exception:
            text = raw.decode("utf-8", errors="ignore")
        cleaned = []
        for ch in text:
            if ch.isprintable() or ch in "\n\r\t":
                cleaned.append(ch)
            else:
                cleaned.append(" ")
        result = "".join(cleaned)
        result = re.sub(r" {3,}", "  ", result)
        result = re.sub(r"\n{3,}", "\n\n", result)
        lines = [l.strip() for l in result.split("\n") if re.search(r"[\u4e00-\u9fff\w]", l)]
        if lines:
            return "\n".join(lines), None
        return None, err or ".doc 文件无法提取有效文本"
    except Exception as e:
        return None, f"读取 .doc 失败: {str(e)[:200]}"

def _read_doc_fast(file_path):
    """.doc 文件快速文本提取（跳过 Word COM，用于预览场景）"""
    # 先尝试 python-docx（有些 .doc 实际是 docx 格式）
    if HAS_DOCX:
        try:
            doc = DocxDocument(file_path)
            lines = [p.text for p in doc.paragraphs if p.text.strip()]
            if lines:
                return "\n".join(lines), None
        except Exception:
            pass
    # 二进制明文提取（回退方案）
    try:
        with open(file_path, "rb") as f:
            raw = f.read()
        for enc in ["gbk", "utf-8", "gb2312"]:
            try:
                text = raw.decode(enc, errors="ignore")
                # 过滤乱码：去掉过长无意义的行
                cleaned = [line for line in text.split("\n") if 3 < len(line.strip()) < 2000]
                if cleaned:
                    return "\n".join(cleaned[:1000]), None
            except Exception:
                continue
    except Exception:
        pass
    return None, "无法快速提取 .doc 文本，请使用 👁 预览按钮（转PDF）查看完整排版"

def read_file_text(file_path):
    """根据文件后缀读取文本"""
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".docx":
        return read_docx_text(str(path))
    elif ext == ".doc":
        return read_doc_text(str(path))
    elif ext == ".pdf":
        return read_pdf_text(str(path))
    elif ext == ".txt":
        try:
            for enc in ["utf-8", "gbk", "gb2312"]:
                try:
                    with open(path, "r", encoding=enc) as f:
                        return f.read(), None
                except UnicodeDecodeError:
                    continue
            return None, "无法识别 .txt 文件编码"
        except Exception as e:
            return None, f"读取 .txt 失败: {str(e)[:200]}"
    else:
        return None, f"不支持的文件格式: {ext}"

# ---------------------------------------------------------------------------
# 文件内容智能分析
# ---------------------------------------------------------------------------

# _SUBJECT_SHORT_MAP 已在文件前面定义（学科同义词表的反向索引）
# 这里补一份正向别名表（subject_group → 所有同义词），供匹配和路径校验使用
def _get_subject_synonyms(subject_group):
    """返回某个 subject_group 的所有同义词列表（包含自身）"""
    if not subject_group:
        return []
    syns = _SUBJECT_SYNONYMS.get(subject_group)
    if syns is not None:
        return syns
    # 兜底：如果 subject_group 不在表里，返回自身
    return [subject_group]

def _build_subject_mapping(assignments):
    """从 assignments 构建科目名称匹配表（兼容旧 API）
    返回 dict: 任意写法 → 规范的 subject_group
    """
    mapping = dict(_SUBJECT_SHORT_MAP)
    for a in assignments:
        sg = a.get("subject_group", "").strip()
        if not sg:
            continue
        # 把规范名也加入映射（self-mapping）
        if sg not in mapping:
            mapping[sg] = sg
        # 把同义词全部映射到规范名
        for syn in _SUBJECT_SYNONYMS.get(sg, [sg]):
            if syn not in mapping:
                mapping[syn] = sg
    return mapping

def _extract_experiment_number(text):
    """从文本提取实验编号，返回 '第X次' 格式"""
    return parse_experiment_number(text) or None

def _extract_content_type(text):
    """从文本提取内容类型：报告/作业/实验"""
    text_lower = text.lower()
    if "报告" in text_lower:
        return "报告"
    if "作业" in text_lower:
        return "作业"
    if "实验" in text_lower:
        return "实验"
    return ""

def analyze_file_content(file_path):
    """
    读取文件全文，提取：姓名、学号、科目、实验编号、内容类型
    返回: { ok, matches: [{name, student_id, subject, experiment, content_type, confidence}] }
    """
    path = Path(file_path)
    if not path.exists():
        return {"ok": False, "msg": "文件不存在", "matches": []}

    # 1. 读取文本
    text, err = read_file_text(str(path))

    # .doc 文件如果二进制提取结果太差（中文很少），强制用 Word COM 再提取一次
    ext = path.suffix.lower()
    if ext == ".doc":
        if err or not text or len(re.findall(r"[\u4e00-\u9fff]", text or "")) < 5:
            text2, _ = read_doc_text_via_word(str(path))
            if text2 and text2.strip():
                text = text2
                err = None

    if err:
        return {"ok": False, "msg": err, "matches": []}
    if not text or not text.strip():
        return {"ok": False, "msg": "文件内容为空或无法提取文本", "matches": []}

    # 2. 加载数据
    students = load_students()
    cfg = load_config()
    assignments = cfg.get("assignments", [])
    subject_map = _build_subject_mapping(assignments)

    # 3. 提取实验编号和内容类型（先提取，因为可能与科目/姓名重叠）
    experiment = _extract_experiment_number(text)
    content_type = _extract_content_type(text)

    # 4. 匹配学生（按姓名和学号）
    student_matches = []
    for s in students:
        name = s.get("name", "")
        sid = s.get("student_id", "")
        if not name:
            continue
        score = 0
        if name in text:
            score += len(name) * 10  # 名字越长越精确
        if sid and sid in text:
            score += 20  # 学号匹配权重很高
        if score > 0:
            student_matches.append({"name": name, "student_id": sid, "score": score})
    student_matches.sort(key=lambda x: x["score"], reverse=True)

    # 5. 匹配科目
    subject_matches = []
    full_text = text
    for full_name, short_name in subject_map.items():
        # 检查全称或简称是否在文本中出现
        if full_name in full_text or short_name in full_text:
            # 计算权重：全称匹配更高
            w = 10 if full_name in full_text else 5
            subject_matches.append({"subject": short_name, "full": full_name, "weight": w})
    subject_matches.sort(key=lambda x: x["weight"], reverse=True)

    # 6. 生成匹配结果组合
    matches = []
    best_student = student_matches[0] if student_matches else None
    best_subject = subject_matches[0] if subject_matches else None

    if best_student or best_subject or experiment:
        match = {
            "name": best_student["name"] if best_student else None,
            "student_id": best_student["student_id"] if best_student else None,
            "subject": best_subject["subject"] if best_subject else None,
            "experiment": experiment,
            "content_type": content_type,
            "confidence": "high" if (best_student and best_subject) else ("medium" if (best_student or best_subject) else "low"),
        }
        matches.append(match)

    # 如果有多个候选人，也列出备选
    if len(student_matches) > 1:
        for alt in student_matches[1:3]:  # 最多 2 个备选
            alt_match = {
                "name": alt["name"],
                "student_id": alt["student_id"],
                "subject": best_subject["subject"] if best_subject else None,
                "experiment": experiment,
                "content_type": content_type,
                "confidence": "low",
            }
            matches.append(alt_match)

    if not matches:
        # 哪怕没有找到明确匹配，也返回提取到的实验和类型信息
        if experiment or content_type:
            matches.append({
                "name": None, "student_id": None, "subject": None,
                "experiment": experiment, "content_type": content_type,
                "confidence": "low",
            })

    return {"ok": True, "matches": matches, "text_snippet": text[:500]}

# ---------------------------------------------------------------------------
# Multipart 表单解析（文件上传）
# ---------------------------------------------------------------------------

def parse_multipart(body: bytes, boundary: str):
    """解析 multipart/form-data，返回 [{name, filename, content_type, data}]"""
    boundary_bytes = boundary.encode("utf-8")
    parts = []
    # 分割各部分
    raw_parts = body.split(b"--" + boundary_bytes)
    for part in raw_parts:
        if not part or part == b"--" or part == b"--\r\n":
            continue
        # 分离头部和体
        header_end = part.find(b"\r\n\r\n")
        if header_end == -1:
            continue
        headers_block = part[:header_end].decode("utf-8", errors="replace")
        file_data = part[header_end + 4:]
        # 去掉尾部 \r\n
        if file_data.endswith(b"\r\n"):
            file_data = file_data[:-2]
        # 解析 Content-Disposition
        disp = {}
        for line in headers_block.split("\r\n"):
            if line.startswith("Content-Disposition:"):
                for item in line[len("Content-Disposition:"):].split(";"):
                    item = item.strip()
                    if "=" in item:
                        k, v = item.split("=", 1)
                        disp[k.strip()] = v.strip().strip('"')
            elif line.startswith("Content-Type:"):
                disp["content_type"] = line[len("Content-Type:"):].strip()
        parts.append({
            "name": disp.get("name", ""),
            "filename": disp.get("filename", ""),
            "content_type": disp.get("content_type", "application/octet-stream"),
            "data": file_data,
        })
    return parts

# ---------------------------------------------------------------------------
# 核心处理流程
# ---------------------------------------------------------------------------

def process_new_file(file_info, students, assignments, submissions):
    """处理一个新文件的完整流程（带去重：同一文件路径只保留最新记录）"""
    filename = file_info["name"]
    file_path = file_info["path"]
    cfg = load_config_raw()
    student_name, score = match_file_to_student(filename, students)
    subject_result = classify_file_subject(file_info, assignments, cfg, source_kind="wechat")
    subject_group = subject_result.get("subject_group", "")
    if subject_result.get("status") == "subject_matched":
        assignment_result = classify_assignment_in_subject(filename, subject_group, assignments)
    else:
        assignment_result = {"status": subject_result.get("status", "unmatched"), "assignment_id": "",
                             "score": 0, "candidates": [], "evidence": subject_result.get("evidence", [])}
    assignment_id = assignment_result.get("assignment_id", "")
    status = assignment_result.get("status", "unmatched")
    bucket_id = assignment_id if assignment_id else PENDING_ARCHIVE_BUCKET
    classification = {
        "stage": status, "subject_group": subject_group,
        "subject_score": subject_result.get("score", 0),
        "confidence": subject_result.get("confidence", subject_result.get("score", 0) / 100),
        "source": subject_result.get("source", "legacy_rules"),
        "subject_candidates": subject_result.get("subject_candidates", []),
        "assignment_id": assignment_id, "assignment_score": assignment_result.get("score", 0),
        "candidates": assignment_result.get("candidates", []),
        "evidence": list(subject_result.get("evidence", [])) + list(assignment_result.get("evidence", [])),
        "source_kind": "wechat",
    }

    assignment_name = next((a["name"] for a in assignments if a["id"] == assignment_id), "其他")
    if not assignment_id:
        assignment_name = "待确认"

    def apply_match_metadata(record, replace_file=False, refresh_time=True):
        if replace_file:
            record["file"] = file_info
        record["student"] = student_name
        record["match_score"] = score
        if refresh_time:
            record["detected_at"] = datetime.now().isoformat()
        record["assignment_id"] = assignment_id
        record["assignment_name"] = assignment_name
        record["subject_group"] = subject_group
        record["status"] = status if student_name else "unmatched"
        record["classification"] = classification
        return record

    # --- 去重检查：如果此文件路径已有记录，跳过 ---
    for aid, records in submissions.items():
        if not isinstance(records, list):
            continue
        for i, r in enumerate(records):
            if r.get("file", {}).get("path") == file_path:
                # 已有记录，更新状态而非新增
                apply_match_metadata(r)
                # 如果 assignment_id 变了，需要移动记录
                if aid != bucket_id:
                    submissions.setdefault(bucket_id, []).append(r)
                    del records[i]
                if student_name and _can_archive_record(r) and not r.get("organized_to"):
                    dest = organize_file(file_info, student_name, assignment_id)
                    r["organized_to"] = dest
                return r

    # --- 统一提交指纹去重：防止微信源文件和公示/实验目录副本重复提交 ---
    incoming_key = _submission_dedupe_key_from_parts(
        student_name,
        assignment_id or PENDING_ARCHIVE_BUCKET,
        filename,
        file_info.get("size", 0),
    )
    incoming_probe = {
        "file": file_info,
        "student": student_name,
        "assignment_id": assignment_id,
        "assignment_name": assignment_name,
        "subject_group": subject_group,
        "status": status if student_name else "unmatched",
        "detected_at": datetime.now().isoformat(),
        "classification": classification,
    }
    for aid, records2 in submissions.items():
        if not isinstance(records2, list):
            continue
        for i, r2 in enumerate(records2):
            if _submission_record_dedupe_key(r2, aid) == incoming_key:
                # 同学生、同作业、同规范化文件名、同大小已存在：更新而非新增。
                replace_file = _prefer_submission_record(incoming_probe, r2, cfg)
                apply_match_metadata(r2, replace_file=replace_file, refresh_time=replace_file)
                if aid != bucket_id:
                    submissions.setdefault(bucket_id, []).append(r2)
                    del records2[i]
                if student_name and _can_archive_record(r2) and not r2.get("organized_to"):
                    dest = organize_file(file_info, student_name, assignment_id)
                    r2["organized_to"] = dest
                return r2

    record = {
        "file": file_info,
        "student": student_name,
        "match_score": score,
        "detected_at": datetime.now().isoformat(),
        "assignment_id": assignment_id,
        "assignment_name": assignment_name,
        "subject_group": subject_group,
        "status": status if student_name else "unmatched",
        "classification": classification,
    }

    if student_name and _can_archive_record(record):
        dest = organize_file(file_info, student_name, assignment_id)
        record["organized_to"] = dest

    # 记录到 submissions
    submissions.setdefault(bucket_id, []).append(record)

    return record

# ---------------------------------------------------------------------------
# 文件扫描主循环（在独立线程中运行）
# ---------------------------------------------------------------------------

class FileWatcher:
    def __init__(self):
        self._stop_event = threading.Event()
        self._spinner_idx = 0

        # 恢复持久化状态
        state = _load_watcher_state()
        self.known_files = set(state["known_files"].keys())
        self.file_mtimes = dict(state["known_files"])
        if self.known_files:
            print(f"[Watcher] Restored {len(self.known_files)} known files from data/watcher_state.json")
        else:
            print(f"[Watcher] No prior state, starting fresh")

        self.running = False
        self.thread = None

        # 启动时检查 watch 目录是否存在
        for d in get_effective_watch_dirs():
            wpath = Path(d)
            if not wpath.exists():
                print(f"{C.RED}[WARN] Watch dir not found: {d}{C.RESET}")
                print(f"  → Files dropped here will NOT be auto-collected")
                print(f"  → Common causes: WeChat logged out / directory renamed / drive unmounted")
            elif not os.access(wpath, os.R_OK):
                print(f"{C.RED}[WARN] Watch dir permission denied: {d}{C.RESET}")

        # 快速初始化：只记录已有文件到 known_files，不处理
        self._quick_init()

    # --- 终端动画 ---

    def _tick_spinner(self):
        """原地刷新动画行（\r 回到行首，flush 强制输出）。"""
        frames = _SPINNER_FRAMES
        encoding = (getattr(sys.stdout, "encoding", None) or "").lower()
        if encoding and "utf" not in encoding:
            frames = _ASCII_SPINNER_FRAMES
        frame = frames[self._spinner_idx % len(frames)]
        self._spinner_idx += 1
        line = f"\r{C.GRAY}{frame} Watching... ({len(self.known_files)} known){C.RESET}"
        try:
            sys.stdout.write(line)
        except UnicodeEncodeError:
            fallback = _ASCII_SPINNER_FRAMES[self._spinner_idx % len(_ASCII_SPINNER_FRAMES)]
            sys.stdout.write(f"\r{C.GRAY}{fallback} Watching... ({len(self.known_files)} known){C.RESET}")
        sys.stdout.flush()

    def _log_new_file(self, f, record):
        """新文件多行日志：先结束动画行，再打详情。"""
        sys.stdout.write("\n")
        sys.stdout.flush()
        path = Path(f["path"])
        duration = record.get("_process_time", 0) if record else 0
        student = record.get("student", "?") if record else "?"
        match_type = "exact match" if (record and record.get("match_score", -1) >= 100) else ("fuzzy match" if record else "unknown")
        print(f"{C.GREEN}[NEW]{C.RESET} {path.name}")
        if record:
            print(f"  → student: {student} ({match_type})")
            print(f"  → assignment: {record.get('assignment_name', '?')}")
            if record.get("organized_to"):
                print(f"  → archived: {record['organized_to']}")
            elif record.get("file", {}).get("path"):
                print(f"  → saved: {record['file']['path']}")
            print(f"  → duration: {duration:.3f}s")
        else:
            print(f"  → status: unmatched (suggests manual review)")
            print(f"  → duration: {duration:.3f}s")

    def _save_state_if_dirty(self):
        """仅在有变动时将当前 file_mtimes 原子写回持久化文件。"""
        state = {"known_files": self.file_mtimes, "updated_at": datetime.now().isoformat()}
        _save_watcher_state(state)

    # --- 生命周期 ---

    def _quick_init(self):
        """快速初始化：扫描文件列表到 known_files，匹配关键词的文件跳过（留给 watch loop 处理）"""
        cfg = load_config_raw()
        watch_dirs = get_effective_watch_dirs(cfg)
        keywords = cfg.get("file_keywords", ["作业", "报告", "论文"])
        doc_exts = configured_file_types(cfg)
        count_known = 0
        count_deferred = 0
        count_skipped = 0
        for wdir in watch_dirs:
            wpath = Path(wdir)
            if not wpath.exists():
                continue
            for f in wpath.rglob("*"):
                if f.is_file() and f.suffix.lower() in doc_exts:
                    if f.name.startswith("~") or f.name.startswith("."):
                        continue
                    fstr = str(f)
                    if fstr in self.known_files:
                        count_skipped += 1
                        continue
                    # 匹配关键词的文件不加入 known_files，留给 watch loop 做完整处理
                    if keywords and any(kw in f.name for kw in keywords):
                        count_deferred += 1
                        continue
                    mtime = f.stat().st_mtime
                    self.known_files.add(fstr)
                    self.file_mtimes[fstr] = mtime
                    count_known += 1
        print(f"[Quick Init] Found {count_known + count_skipped} docx files, {count_deferred} keyword files deferred to watch loop")
        if count_skipped:
            print(f"[Quick Init] Skipped {count_skipped} already-known files")
        if count_known > 0:
            self._save_state_if_dirty()

    def start(self):
        self.running = True
        self.thread = threading.Thread(target=self._watch_loop, daemon=True)
        self.thread.start()
        print("[Watcher] Started")

    def stop(self):
        self.running = False
        self._stop_event.set()
        sys.stdout.write("\n")
        sys.stdout.flush()
        print("[Watcher] Stopped")

    def _watch_loop(self):
        cfg = load_config()
        interval = cfg.get("poll_interval", 5)
        print(f"[WatchLoop] Started, interval={interval}s")
        dirty = False
        while self.running:
            try:
                cfg = load_config()
                if cfg.get("watch_enabled", True):
                    students = load_students()
                    submissions = load_submissions()
                    assignments = cfg.get("assignments", [])
                    _t = time.time()
                    new_files = scan_new_files(self.known_files)

                    for f in new_files:
                        try:
                            path = f["path"]
                            mtime = f.get("mtime", 0)

                            # mtime 未变 → 跳过（持久化恢复后的防重复）
                            if path in self.known_files and self.file_mtimes.get(path) == mtime:
                                continue

                            _ts = time.time()
                            record = process_new_file(f, students, assignments, submissions)
                            if record:
                                record["_process_time"] = time.time() - _ts

                            self.known_files.add(path)
                            self.file_mtimes[path] = mtime
                            dirty = True

                            self._log_new_file(f, record)
                        except Exception as e:
                            sys.stdout.write("\n")
                            sys.stdout.flush()
                            print(f"{C.RED}[ERROR]{C.RESET} {Path(f['path']).name}: {e}")
                            traceback.print_exc()
                        self._tick_spinner()

                    if dirty and new_files:
                        save_submissions(submissions)
                        self._save_state_if_dirty()
                        dirty = False

                if not self._stop_event.wait(interval):
                    self._tick_spinner()
            except Exception as e:
                sys.stdout.write("\n")
                sys.stdout.flush()
                print(f"{C.RED}[Watcher Error]{C.RESET} {e}")
                traceback.print_exc()
                self._stop_event.wait(interval)
                self._tick_spinner()

watcher = None  # 延迟初始化，避免模块加载时阻塞

# ---------------------------------------------------------------------------
# HTTP API Server
# ---------------------------------------------------------------------------

class APIHandler(SimpleHTTPRequestHandler):
    """Flask-free 的轻量 API Handler，内嵌静态文件服务"""

    def log_message(self, format, *args):
        print(f"[HTTP] {self.command} {self.path}")

    def _request_is_local(self):
        return _is_loopback_ip(self.client_address[0])

    def _has_lan_session(self, cfg=None):
        cfg = cfg or load_config_raw()
        expected = str(cfg.get("lan_access_token") or "")
        if not expected:
            return False
        try:
            cookie = SimpleCookie(self.headers.get("Cookie", ""))
            supplied = cookie.get(LAN_SESSION_COOKIE)
            return bool(supplied and hmac.compare_digest(supplied.value, expected))
        except Exception:
            return False

    def _authorize_request(self, path):
        cfg = load_config_raw()
        if self._request_is_local() or not cfg.get("lan_access_enabled", False):
            return True
        if path in ("/lan-login", "/api/lan-auth"):
            return True
        if self._has_lan_session(cfg):
            return True
        if path.startswith("/api/"):
            self._json({"ok": False, "msg": "需要局域网访问口令"}, status=401)
        else:
            self.send_response(302)
            self.send_header("Location", "/lan-login")
            self.send_header("Cache-Control", "no-store")
            self.end_headers()
        return False

    def _origin_is_allowed(self):
        origin = self.headers.get("Origin", "").strip()
        if not origin:
            return True
        try:
            return urlparse(origin).netloc.lower() == self.headers.get("Host", "").lower()
        except ValueError:
            return False

    def _serve_lan_login(self):
        body = """<!doctype html><html lang='zh-CN'><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>局域网访问验证</title><style>body{font-family:Microsoft YaHei,sans-serif;background:#f5f7fb;color:#172033;display:grid;place-items:center;min-height:100vh;margin:0}.box{width:min(380px,calc(100vw - 40px));background:#fff;border:1px solid #dbe3ef;border-radius:8px;padding:28px;box-shadow:0 18px 50px #26334d1a}h1{font-size:22px;margin:0 0 10px}p{color:#64748b;line-height:1.65}input,button{box-sizing:border-box;width:100%;height:42px;border-radius:6px}input{border:1px solid #cbd5e1;padding:0 12px;margin:10px 0}button{border:0;background:#2563eb;color:#fff;font-weight:700;cursor:pointer}.msg{min-height:22px;color:#dc2626;font-size:13px;margin-top:10px}</style></head><body><form class='box' id='form'><h1>局域网访问验证</h1><p>请输入运行仪表盘电脑上显示的访问口令。仅在可信私人网络中使用。</p><input id='code' autocomplete='current-password' placeholder='访问口令' required><button>验证并进入</button><div class='msg' id='msg'></div></form><script>form.onsubmit=async(e)=>{e.preventDefault();msg.textContent='正在验证...';try{const r=await fetch('/api/lan-auth',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({token:code.value})});const d=await r.json();if(!r.ok||!d.ok)throw new Error(d.msg||'口令错误');location.href='/';}catch(err){msg.textContent=err.message||'验证失败';}}</script></body></html>""".encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Cache-Control", "no-store")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self):
        try:
            if not self._authorize_request(urlparse(self.path).path):
                return
            self._do_GET_impl()
        except (BrokenPipeError, ConnectionResetError, ConnectionAbortedError):
            pass  # 客户端断开，静默忽略
        except Exception as e:
            print(f"[ERROR] GET {self.path}: {e}")
            traceback.print_exc()
            try:
                self.send_error(500)
                self.end_headers()
            except Exception:
                pass

    def _do_GET_impl(self):
        parsed = urlparse(self.path)
        path = parsed.path
        qs = parse_qs(parsed.query)

        # ---- API routes ----
        if path == "/lan-login":
            self._serve_lan_login()

        elif path == "/api/network-access":
            port = self.server.server_address[1]
            self._json(network_access_payload(load_config_raw(), port, include_token=True,
                                              is_local=self._request_is_local()))

        elif path == "/api/status":
            self._json({
                "watching": watcher.running,
                "known_files": len(watcher.known_files),
                "watch_dirs": get_effective_watch_dirs(),
            })

        elif path == "/api/health":
            self._json(data_health())

        elif path == "/api/students":
            self._json(load_students())

        elif path == "/api/submissions":
            self._json(load_submissions())

        elif path == "/api/config":
            self._json(load_config())

        elif path == "/api/ai/brain":
            self._json(ai_brain_payload())

        elif path in ("/api/ai/rules", "/api/ai/rules/export"):
            self._json({"ok": True, "rule_pack": load_ai_rules()})

        elif path == "/api/theme":
            self._json(theme_payload())

        elif path == "/api/scan-now":
            cfg = load_config()
            students = load_students()
            assignments = cfg.get("assignments", [])
            submissions = load_submissions()
            found = []
            errors = []
            for f in scan_new_files(watcher.known_files):
                try:
                    record = process_new_file(f, students, assignments, submissions)
                    watcher.known_files.add(f["path"])
                    if f.get("mtime"):
                        watcher.file_mtimes[f["path"]] = f["mtime"]
                    found.append(record)
                except Exception as e:
                    err_msg = f"{f['name']}: {e}"
                    print(f"[ERROR] scan-now {err_msg}")
                    traceback.print_exc()
                    errors.append({"file": f["name"], "path": f["path"], "error": str(e)})
                    watcher.known_files.add(f["path"])
            if found:
                save_submissions(submissions)
                watcher._save_state_if_dirty()
            self._json({"scanned": len(found), "new": found, "errors": errors})
            # 后台预热 Word 预览缓存（已禁用：COM 弹窗问题）
            # if found:
            #     warm_preview_cache_async()

        elif path == "/api/scan-existing":
            cfg = load_config()
            # 支持 target=experiment: 从公示/实验目录回填到已收作业
            # GET 走 query string（apiGet），POST 走 body（apiPost），两者都支持
            # TODO(方案D-Step2): 此路由当前只在 GET handler 注册；POST handler 注册后此 try 块才真正命中 body 分支
            _qs_target = (qs.get("target", [""])[0] if qs else "")
            try:
                _body_target = data.get("target", "")
            except (NameError, UnboundLocalError):
                _body_target = ""
            target = _qs_target or _body_target
            if target == "experiment":
                if not cfg.get("experiment_enabled", False):
                    self._json({"scanned": 0, "matched": 0, "deleted": 0, "skipped_no_student": 0, "dirs": 0,
                                "ok": False, "msg": "公示文件夹未启用"})
                    return
                if not EXPERIMENT_BASE:
                    self._json({"scanned": 0, "matched": 0, "deleted": 0, "skipped_no_student": 0, "dirs": 0,
                                "msg": "experiment_dir 未配置"})
                    return
                exp_path = str(EXPERIMENT_BASE)
                result = scan_existing_directory(exp_path, filename_student=True)
                self._json(result)
                return

            scan_dirs = cfg.get("scan_dirs", [])
            if not scan_dirs:
                self._json({"scanned": 0, "matched": 0, "deleted": 0, "deduped_records": 0, "skipped_no_student": 0, "dirs": 0})
                return
            total = {"scanned": 0, "matched": 0, "deleted": 0, "deduped_records": 0, "skipped_no_student": 0, "dirs": 0}
            for d in scan_dirs:
                result = scan_existing_directory(d)
                if "error" in result:
                    continue
                total["scanned"] += result.get("scanned", 0)
                total["matched"] += result.get("matched", 0)
                total["deleted"] += result.get("deleted", 0)
                total["deduped_records"] += result.get("deduped_records", 0)
                total["skipped_no_student"] += result.get("skipped_no_student", 0)
                total["dirs"] += 1
            self._json(total)
            # 后台预热 Word 预览缓存（已禁用：COM 弹窗问题）
            # warm_preview_cache_async()

        elif path == "/api/dashboard":
            self._dashboard_data(qs.get("assignment", [None])[0])

        elif path == "/api/scan-dirs":
            self._json(scan_dirs_payload())

        elif path == "/api/assignments":
            self._assignments_list()

        elif path.startswith("/api/assignment/"):
            aid = path.split("/api/assignment/")[1]
            self._assignment_detail(aid)

        elif path == "/api/open-file":
            fp = qs.get("path", [""])[0]
            safe_fp = is_allowed_file_path(fp)
            if safe_fp:
                try:
                    if sys.platform == "win32":
                        os.startfile(str(safe_fp))
                    else:
                        subprocess.Popen(["xdg-open", str(safe_fp)])
                    self._json({"ok": True})
                except Exception as e:
                    self._json({"ok": False, "msg": str(e)})
            else:
                self._json({"ok": False, "msg": "文件不存在或无权访问"})

        elif path == "/api/open-folder":
            fp = qs.get("path", [""])[0]
            safe_fp = is_allowed_file_path(fp)
            if safe_fp:
                parent = str(safe_fp.parent)
                try:
                    if sys.platform == "win32":
                        os.startfile(parent)
                    else:
                        subprocess.Popen(["xdg-open", parent])
                    self._json({"ok": True})
                except Exception as e:
                    self._json({"ok": False, "msg": str(e)})
            else:
                self._json({"ok": False, "msg": "路径为空"})

        elif path == "/api/download":
            fp = qs.get("path", [""])[0]
            self._send_download(fp)

        elif path == "/api/download-zip":
            fps = qs.get("paths", [""])[0]
            paths = [p for p in fps.split(",") if p.strip()]
            self._send_zip(paths)

        elif path == "/api/pack-subject":
            sg = qs.get("subject_group", [""])[0]
            self._pack_subject(sg)

        elif path == "/api/templates":
            # 获取模板列表，支持 ?assignment_id=xxx 过滤（+ 科目/实验模糊匹配）
            cfg = load_config()
            templates = cfg.get("templates", [])
            aid_filter = qs.get("assignment_id", [None])[0]
            if aid_filter:
                # 精确按 assignment_id 匹配
                matched = [t for t in templates if t.get("assignment_id") == aid_filter]
                if not matched:
                    # 没精确匹配到，尝试按科目+实验模糊匹配
                    assignments = cfg.get("assignments", [])
                    target = next((a for a in assignments if a.get("id") == aid_filter), None)
                    if target:
                        sg = target.get("subject_group", "")
                        exp = target.get("experiment", "")
                        matched = [t for t in templates
                                   if (t.get("subject") == sg and (not exp or t.get("experiment") == exp))
                                   or t.get("assignment_id") == aid_filter]
                    templates = matched
                else:
                    templates = matched
            self._json({"templates": templates})

        elif path == "/api/preview":
            fp = qs.get("path", [""])[0]
            src = is_allowed_file_path(fp)
            if not src:
                self._json({"ok": False, "msg": "文件不存在或无权访问"})
                return
            # 大文件保护：超过 10MB 的文本文件不预览
            try:
                if src.stat().st_size > 10 * 1024 * 1024:
                    self._json({"ok": False, "msg": "文件过大，无法预览（> 10MB）。请直接打开文件查看。", "text": ""})
                    return
            except:
                pass
            # .doc 文件跳过 Word COM（太慢），只用快速方法
            if src.suffix.lower() == ".doc":
                text, err = _read_doc_fast(str(src))
            else:
                text, err = read_file_text(str(src))
            if err:
                self._json({"ok": False, "msg": err, "text": ""})
            else:
                # 截断过长文本（500KB 以内）
                t = text or ""
                if len(t) > 500000:
                    t = t[:500000] + "\n\n...[文本过长，已截断，完整内容请直接打开文件]..."
                self._json({"ok": True, "text": t, "path": str(src), "name": src.name})

        elif path == "/api/serve-file":
            # 直接提供文件内容（浏览器原生渲染 PDF/图片等）
            fp = qs.get("path", [""])[0]
            p = is_allowed_file_path(fp)
            if not p:
                self.send_error(403)
                self.end_headers()
                return
            if not p.exists():
                self.send_error(404)
                self.end_headers()
                return
            ext = p.suffix.lower()
            content_types = {
                ".pdf": "application/pdf",
                ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                ".gif": "image/gif", ".webp": "image/webp",
                ".txt": "text/plain; charset=utf-8",
                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ".doc": "application/msword",
            }
            ct = content_types.get(ext, "application/octet-stream")
            try:
                file_size = p.stat().st_size
                # 小文件（< 2MB）直接读内存，大文件流式传输
                if file_size < 2 * 1024 * 1024:
                    data = p.read_bytes()
                    self.send_response(200)
                    self.send_header("Content-Type", ct)
                    self.send_header("Content-Length", str(file_size))
                    self.send_header("Cache-Control", "public, max-age=60")
                    self.end_headers()
                    self.wfile.write(data)
                else:
                    self.send_response(200)
                    self.send_header("Content-Type", ct)
                    self.send_header("Content-Length", str(file_size))
                    self.send_header("Cache-Control", "public, max-age=60")
                    self.end_headers()
                    with open(p, "rb") as f:
                        shutil.copyfileobj(f, self.wfile, 256 * 1024)
            except Exception as e:
                self.send_error(500)
                self.end_headers()

        elif path == "/api/preview-status":
            job_id = qs.get("job_id", [""])[0]
            with _preview_jobs_lock:
                job = dict(_preview_jobs.get(job_id, {"status": "error", "error": "预览任务不存在"}))
            if job.get("status") == "ready":
                job["preview_url"] = "/api/preview-doc?path=" + quote(str(job.get("path", "")))
            self._json({"ok": job.get("status") != "error", **job})

        elif path == "/api/preview-doc":
            # word(.docx/.doc) → PDF 即时转换并直传（保留图表和排版）
            fp = qs.get("path", [""])[0]
            src = is_allowed_file_path(fp)
            if not src:
                self.send_error(403)
                self.end_headers()
                return
            ext = src.suffix.lower()
            if ext not in (".docx", ".doc"):
                self.send_error(400)
                self.end_headers()
                return

            # 缓存：key = 路径hash + 文件修改时间，缓存到 data/preview_cache/
            cache_pdf = _preview_cache_path(src)

            # 旧缓存清理（删除同一源文件的过期缓存）
            if PREVIEW_CACHE_DIR.exists():
                for old in PREVIEW_CACHE_DIR.glob(f"{hashlib.md5(str(src).encode()).hexdigest()[:12]}_*.pdf"):
                    if old != cache_pdf:
                        try: old.unlink()
                        except: pass

            if cache_pdf.exists():
                # 命中缓存 → 流式直传
                file_size = cache_pdf.stat().st_size
                self.send_response(200)
                self.send_header("Content-Type", "application/pdf")
                self.send_header("Content-Length", str(file_size))
                self.send_header("Cache-Control", "public, max-age=300")
                self.end_headers()
                with open(cache_pdf, "rb") as f:
                    shutil.copyfileobj(f, self.wfile, 256 * 1024)
                return

            # 转换（使用 _WordComContext 静默模式，不弹窗）
            # Convert outside the request thread so the dashboard stays responsive.
            job_id, job = _queue_preview_job(src)
            if job.get("status") == "ready":
                cache_pdf = _preview_cache_path(src)
                file_size = cache_pdf.stat().st_size
                self.send_response(200)
                self.send_header("Content-Type", "application/pdf")
                self.send_header("Content-Length", str(file_size))
                self.send_header("Cache-Control", "public, max-age=300")
                self.end_headers()
                with open(cache_pdf, "rb") as f:
                    shutil.copyfileobj(f, self.wfile, 256 * 1024)
                return
            payload = json.dumps({"ok": False, "status": job.get("status", "queued"),
                                  "job_id": job_id, "path": str(src),
                                  "msg": "正在后台生成预览"}, ensure_ascii=False).encode("utf-8")
            self.send_response(202)
            self.send_header("Content-Type", "application/json; charset=utf-8")
            self.send_header("Content-Length", str(len(payload)))
            self.end_headers()
            self.wfile.write(payload)
            return

            tmp_pdf = CONVERT_TEMP_DIR / (src.stem + "_preview.pdf")
            ok = False

            if ext in (".docx", ".doc"):
                try:
                    with _WordComContext() as ctx:
                        doc = ctx.open(src)
                        doc.ExportAsFixedFormat(str(tmp_pdf), 17)  # 17 = wdExportFormatPDF
                        ok = tmp_pdf.exists() and tmp_pdf.stat().st_size > 0
                except Exception as e:
                    print(f"[PreviewDoc] Word COM 失败: {e}")

            if ok:
                # 保存到缓存
                try: shutil.copy2(str(tmp_pdf), str(cache_pdf))
                except: pass
                # 流式直传
                file_size = tmp_pdf.stat().st_size
                self.send_response(200)
                self.send_header("Content-Type", "application/pdf")
                self.send_header("Content-Length", str(file_size))
                self.send_header("Cache-Control", "public, max-age=300")
                self.end_headers()
                with open(tmp_pdf, "rb") as f:
                    shutil.copyfileobj(f, self.wfile, 256 * 1024)
                try: tmp_pdf.unlink()
                except: pass
                return

            # 转换都失败了
            if ext == ".doc":
                html = self._error_html("无法预览 .doc 文件",
                    "需要本机安装 Microsoft Word 和 pywin32 库才能转换预览。<br>"
                    "请运行: <code>pip install pywin32</code><br>"
                    "然后执行: <code>python -m pywin32_postinstall -install</code><br><br>"
                    "或者将文件另存为 .docx 格式后重试。")
            else:
                text, _ = read_file_text(str(src))
                html = self._text_html(text or "无法提取内容")
            body = html.encode("utf-8")
            self.send_response(200)
            self.send_header("Content-Type", "text/html; charset=utf-8")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)

        elif path == "/api/changelog":
            # 返回 CHANGELOG.md 内容
            changelog_path = BASE_DIR / "CHANGELOG.md"
            if changelog_path.exists():
                content = changelog_path.read_text(encoding="utf-8")
                self._json({"ok": True, "content": content})
            else:
                self._json({"ok": True, "content": "# 暂无更新日志"})

        elif path == "/api/announcements":
            # 返回公告文件内容
            ann_path = BASE_DIR / "announcement.json"
            if ann_path.exists():
                try:
                    data = json.loads(ann_path.read_text("utf-8"))
                    self._json({
                        "ok": True,
                        "version": data.get("version", "0.0.0"),
                        "announcements": data.get("announcements", []),
                    })
                except Exception:
                    self._json({"ok": True, "version": "0.0.0", "announcements": []})
            else:
                self._json({"ok": True, "version": "0.0.0", "announcements": []})

        elif path == "/api/server-status":
            self._json({
                "status": "running",
                "version": current_app_version(),
                "uptime": int(time.time() - _app_start_time),
                "pid": os.getpid(),
            })

        elif path == "/api/update/check":
            self._check_github_update()

        # ---- 静态文件 ----
        elif path == "/":
            cfg = load_config_raw()
            if cfg.get("default_frontend") == "modern":
                self._serve_html("dashboard_modern.html")
            else:
                self._serve_html("dashboard.html")

        elif path == "/dashboard":
            self._serve_html("dashboard.html")
        elif path == "/modern":
            cfg = load_config_raw()
            if cfg.get("default_frontend") != "modern":
                cfg["default_frontend"] = "modern"
                save_config(cfg)
            self._serve_html("dashboard_modern.html")
        else:
            # 尝试静态文件
            file_path = BASE_DIR / path.lstrip("/")
            if file_path.exists() and file_path.is_file():
                self._serve_static(file_path)
            else:
                self.send_error(404)
                self.end_headers()

    def do_POST(self):
        try:
            path = urlparse(self.path).path
            if not self._authorize_request(path):
                return
            if not self._origin_is_allowed():
                self._json({"ok": False, "msg": "拒绝跨站请求"}, status=403)
                return
            self._do_POST_impl()
        except (BrokenPipeError, ConnectionResetError, ConnectionAbortedError):
            pass  # 客户端断开，静默忽略
        except Exception as e:
            print(f"[ERROR] POST {self.path}: {e}")
            traceback.print_exc()
            try:
                self.send_error(500)
                self.end_headers()
            except Exception:
                pass

    def _do_POST_impl(self):
        parsed = urlparse(self.path)
        path = parsed.path
        content_length = int(self.headers.get("Content-Length", 0))
        content_type = self.headers.get("Content-Type", "")

        # 检查是否为 multipart 上传
        if "multipart/form-data" in content_type:
            body = self.rfile.read(content_length) if content_length else b""
            # 提取 boundary
            boundary = None
            for item in content_type.split(";"):
                item = item.strip()
                if item.startswith("boundary="):
                    boundary = item[len("boundary="):].strip('"')
                    break
            if boundary:
                parts = parse_multipart(body, boundary)
                if path == "/api/convert-upload":
                    self._convert_upload(parts)
                    return
                if path == "/api/upload-template":
                    self._upload_template(parts)
                    return
                if path == "/api/update":
                    self._handle_update(parts)
                    return
            self._json({"ok": False, "msg": "不支持的上传请求"})
            return

        body = self.rfile.read(content_length) if content_length else b"{}"
        try:
            data = json.loads(body)
        except:
            data = {}

        if path == "/api/lan-auth":
            cfg = load_config_raw()
            supplied = str(data.get("token") or "")
            expected = str(cfg.get("lan_access_token") or "")
            if not cfg.get("lan_access_enabled", False) or not expected or not hmac.compare_digest(supplied, expected):
                self._json({"ok": False, "msg": "访问口令错误"}, status=401)
                return
            self._json({"ok": True}, extra_headers={
                "Set-Cookie": f"{LAN_SESSION_COOKIE}={expected}; Path=/; HttpOnly; SameSite=Strict"
            })

        elif path == "/api/network-access/configure":
            if not self._request_is_local():
                self._json({"ok": False, "msg": "只能在运行服务的电脑上修改访问模式"}, status=403)
                return
            cfg = load_config_raw()
            enabled = bool(data.get("enabled", False))
            if enabled and (data.get("regenerate_token") or not cfg.get("lan_access_token")):
                cfg["lan_access_token"] = secrets.token_urlsafe(12)
            cfg["lan_access_enabled"] = enabled
            save_config(cfg)
            port = self.server.server_address[1]
            payload = network_access_payload(cfg, port, include_token=True, is_local=True)
            payload.update({"ok": True, "restarting": True})
            self._json(payload)
            threading.Thread(target=_restart_after_delay, name="network-mode-restart").start()

        elif path == "/api/convert-upload":
            # JSON 模式（不支持，提示使用 multipart）
            self._json({"ok": False, "msg": "请使用表单上传文件"})

        elif path == "/api/build-update-package":
            self._build_update_package(data)

        elif path == "/api/theme/save":
            cfg = load_config_raw()
            cfg["ui_theme"] = normalize_ui_theme(data)
            save_config(cfg)
            self._json({"ok": True, "theme": theme_payload(cfg)})

        elif path == "/api/ai/settings":
            cfg = load_config_raw()
            current = ai_settings(cfg)
            mode = str(data.get("mode", current.get("mode", "rules"))).strip()
            if mode not in ("off", "rules"):
                self._json({"ok": False, "msg": "当前版本只支持关闭或规则模式"})
                return
            try:
                sensitivity = min(0.95, max(0.50, float(data.get("sensitivity", current["sensitivity"]))))
            except (TypeError, ValueError):
                self._json({"ok": False, "msg": "灵敏度必须是 0.50 到 0.95 之间的数字"})
                return
            preset = str(data.get("sensitivity_preset", "custom")).strip()
            if preset not in ("conservative", "balanced", "aggressive", "custom"):
                preset = "custom"
            cfg["ai_classifier"] = {
                "mode": mode,
                "sensitivity": round(sensitivity, 2),
                "sensitivity_preset": preset,
                "active_semester": str(data.get("active_semester", current.get("active_semester", ""))).strip()[:100],
            }
            save_config(cfg)
            self._json({"ok": True, "settings": ai_settings(cfg)})

        elif path == "/api/ai/rules/save":
            if not HAS_AI_CLASSIFIER:
                self._json({"ok": False, "msg": "分类大脑模块不可用"})
                return
            try:
                normalized = ai_classifier.normalize_rule_pack(data.get("rule_pack", data))
                if normalized.get("collisions"):
                    aliases = "、".join(item["alias"] for item in normalized["collisions"][:5])
                    self._json({"ok": False, "msg": f"课程别名冲突：{aliases}", "collisions": normalized["collisions"]})
                    return
                rules = save_ai_rules(normalized["rule_pack"])
                self._json({"ok": True, "rule_pack": rules, "summary": normalized["summary"]})
            except Exception as exc:
                self._json({"ok": False, "msg": str(exc)})

        elif path == "/api/ai/rules/import":
            if not HAS_AI_CLASSIFIER:
                self._json({"ok": False, "msg": "分类大脑模块不可用"})
                return
            payload = data.get("rule_pack", data.get("payload", {}))
            if isinstance(payload, str):
                try:
                    payload = json.loads(payload)
                except json.JSONDecodeError as exc:
                    self._json({"ok": False, "msg": f"JSON 解析失败：{exc.msg}"})
                    return
            try:
                normalized = ai_classifier.normalize_rule_pack(payload, data.get("allowed_subjects"))
                if data.get("preview", True):
                    self._json({"ok": True, **normalized})
                    return
                if normalized.get("collisions"):
                    self._json({"ok": False, "msg": "存在课程别名冲突，请调整后再导入",
                                "collisions": normalized["collisions"]})
                    return
                incoming = normalized["rule_pack"]
                rules = ai_classifier.merge_rule_packs(load_ai_rules(), incoming) if data.get("mode", "merge") == "merge" else incoming
                rules = save_ai_rules(rules)
                self._json({"ok": True, "rule_pack": rules, "summary": normalized["summary"]})
            except Exception as exc:
                self._json({"ok": False, "msg": str(exc)})

        elif path == "/api/ai/prompt":
            if not HAS_AI_CLASSIFIER:
                self._json({"ok": False, "msg": "分类大脑模块不可用"})
                return
            courses = data.get("courses", [])
            if isinstance(courses, str):
                courses = [item.strip() for item in re.split(r"[\n,，;；]+", courses) if item.strip()]
            prompt = ai_classifier.build_professional_pack_prompt(
                data.get("profile", {}), courses, data.get("known_aliases", {})
            )
            self._json({"ok": True, "prompt": prompt, "courses": courses})

        elif path == "/api/ai/extract-keywords":
            if not HAS_AI_CLASSIFIER:
                self._json({"ok": False, "msg": "分类大脑模块不可用"})
                return
            subject = str(data.get("subject", "")).strip()
            filenames = data.get("filenames", [])
            if not subject or not isinstance(filenames, list) or not filenames:
                self._json({"ok": False, "msg": "请填写科目并选择至少一个文件"})
                return
            candidates = ai_classifier.extract_keyword_candidates(subject, filenames[:200], load_students())
            self._json({"ok": True, "subject": subject, "candidates": candidates})

        elif path == "/api/assignment/classify":
            if not HAS_AI_CLASSIFIER:
                self._json({"ok": False, "msg": "分类大脑模块不可用"})
                return
            cfg = load_config_raw()
            filename = str(data.get("file_name") or data.get("filename") or "").strip()
            file_path_str = str(data.get("file_path") or "").strip()
            if file_path_str:
                for records in load_submissions().values():
                    for record in records if isinstance(records, list) else []:
                        if str((record.get("file") or {}).get("path") or "") == file_path_str:
                            filename = str((record.get("file") or {}).get("name") or filename)
                            break
                    if filename:
                        break
            if not filename:
                self._json({"ok": False, "msg": "缺少文件名"})
                return
            assignments = cfg.get("assignments", [])
            result = ai_classifier.classify_subject(
                filename,
                assignments=[item for item in assignments if str(item.get("subject_group") or item.get("subject") or "") not in _GENERIC_SUBJECT_GROUPS],
                rules=load_ai_rules(),
                feedback=cfg.get("match_feedback", {}),
                subject_synonyms={name: aliases for name, aliases in _SUBJECT_SYNONYMS.items() if name not in _GENERIC_SUBJECT_GROUPS},
                students=load_students(),
                sensitivity=ai_settings(cfg).get("sensitivity", 0.70),
            )
            assignment_result = {"status": "assignment_pending", "assignment_id": "", "score": 0,
                                 "candidates": [], "evidence": []}
            if result.get("status") == "subject_matched":
                assignment_result = classify_assignment_in_subject(filename, result.get("subject_group", ""), assignments)
            payload = {
                **result,
                "file_name": filename,
                "assignment_id": assignment_result.get("assignment_id", ""),
                "assignment_candidates": assignment_result.get("candidates", []),
                "assignment_score": assignment_result.get("score", 0),
                "stage": assignment_result.get("status") if result.get("status") == "subject_matched" else result.get("stage", "pending_archive"),
                "evidence": list(result.get("evidence", [])) + list(assignment_result.get("evidence", [])),
            }
            self._json({"ok": True, "result": payload})

        elif path == "/api/students/save":
            save_students(data)
            self._json({"ok": True})

        elif path == "/api/students/add":
            students = load_students()
            students.append(data)
            save_students(students)
            self._json({"ok": True, "students": students})

        elif path == "/api/students/delete":
            name = data.get("name")
            students = load_students()
            students = [s for s in students if s.get("name") != name]
            save_students(students)
            self._json({"ok": True, "students": students})

        elif path == "/api/submissions/assign":
            # 手动将未匹配文件分配给某个学生
            file_path_str = data.get("file_path")
            student_name = data.get("student")
            submissions = load_submissions()
            for aid, records in submissions.items():
                for r in records:
                    if r["file"]["path"] == file_path_str:
                        r["student"] = student_name
                        r["status"] = "manual_matched" if r.get("assignment_id") else r.get("status", "assignment_pending")
                        if _can_archive_record(r):
                            dest = organize_file(r["file"], student_name, aid)
                            r["organized_to"] = dest
                        break
            save_submissions(submissions)
            self._json({"ok": True})

        elif path == "/api/submissions/set-subject":
            file_path_str = str(data.get("file_path") or "")
            subject_group = str(data.get("subject_group") or "").strip()
            cfg = load_config_raw()
            assignments = cfg.get("assignments", [])
            subject_group = _canonical_subject_name(subject_group, assignments)
            if not file_path_str or not subject_group or not any(a.get("subject_group") == subject_group for a in assignments):
                self._json({"ok": False, "msg": "科目或文件无效"})
                return
            submissions = load_submissions()
            found = None
            old_bucket = ""
            for aid, records in submissions.items():
                for r in records if isinstance(records, list) else []:
                    if (r.get("file") or {}).get("path") == file_path_str:
                        found, old_bucket = r, aid
                        break
                if found:
                    break
            if not found:
                self._json({"ok": False, "msg": "未找到待归档文件"})
                return
            old_subject = found.get("subject_group", "")
            result = classify_assignment_in_subject((found.get("file") or {}).get("name", ""), subject_group, assignments)
            found["subject_group"] = subject_group
            found["assignment_id"] = ""
            found["assignment_name"] = "待确认"
            found["status"] = "assignment_pending"
            found["classification"] = {
                "stage": "assignment_pending", "subject_group": subject_group,
                "subject_score": 100, "assignment_id": "", "assignment_score": result.get("score", 0),
                "candidates": result.get("candidates", []),
                "evidence": ["用户手动确认科目"] + result.get("evidence", []), "source_kind": "manual",
            }
            if old_bucket != PENDING_ARCHIVE_BUCKET:
                submissions.setdefault(PENDING_ARCHIVE_BUCKET, []).append(found)
                submissions[old_bucket].remove(found)
            feedback = cfg.setdefault("match_feedback", {}).setdefault("subject_corrections", [])
            token = Path(file_path_str).stem[:80]
            existing = next((item for item in feedback if item.get("token") == token and
                             item.get("from_subject") == old_subject and item.get("to_subject") == subject_group), None)
            if existing:
                existing["count"] = int(existing.get("count", 1)) + 1
                existing["updated_at"] = datetime.now().isoformat()
            else:
                feedback.append({"token": token, "from_subject": old_subject,
                                 "to_subject": subject_group, "source_kind": "manual", "count": 1,
                                 "updated_at": datetime.now().isoformat()})
            cfg["match_feedback"]["subject_corrections"] = feedback[-500:]
            save_submissions(submissions)
            save_config(cfg)
            self._json({"ok": True, "classification": found["classification"]})

        elif path == "/api/submissions/assign-assignment":
            file_path_str = str(data.get("file_path") or "")
            assignment_id = str(data.get("assignment_id") or "")
            cfg = load_config_raw()
            assignment = next((a for a in cfg.get("assignments", []) if a.get("id") == assignment_id), None)
            submissions = load_submissions()
            found = None
            old_bucket = ""
            for aid, records in submissions.items():
                for r in records if isinstance(records, list) else []:
                    if (r.get("file") or {}).get("path") == file_path_str:
                        found, old_bucket = r, aid
                        break
                if found:
                    break
            if not found or not assignment:
                self._json({"ok": False, "msg": "文件或作业不存在"})
                return
            if found.get("subject_group") != assignment.get("subject_group"):
                self._json({"ok": False, "msg": "所选作业不属于当前科目"})
                return
            if not found.get("student"):
                self._json({"ok": False, "msg": "请先为文件选择学生，再确认作业"})
                return
            found.update({"assignment_id": assignment_id, "assignment_name": assignment.get("name", ""),
                          "status": "manual_matched"})
            found["classification"] = {**(found.get("classification") or {}), "stage": "manual_matched",
                                         "assignment_id": assignment_id, "evidence": ["用户手动确认作业"]}
            if old_bucket != assignment_id:
                submissions.setdefault(assignment_id, []).append(found)
                submissions[old_bucket].remove(found)
            found["organized_to"] = organize_file(found["file"], found["student"], assignment_id)
            save_submissions(submissions)
            self._json({"ok": True, "organized_to": found["organized_to"]})

        elif path == "/api/submissions/reject-assignment":
            assignment_id = data.get("assignment_id", "")
            file_path_str = data.get("file_path", "")
            if not assignment_id or not file_path_str:
                self._json({"ok": False, "msg": "缺少作业或文件信息"})
                return

            cfg = load_config_raw()
            rejected_terms = []
            seen_terms = set()
            for term in data.get("rejected_terms", []) or []:
                text = str(term or "").strip()
                key = text.lower()
                if len(text) >= 2 and key not in seen_terms:
                    rejected_terms.append(text)
                    seen_terms.add(key)
            submissions = load_submissions()
            records = submissions.get(assignment_id, [])
            if not isinstance(records, list):
                records = []

            target = None
            kept = []
            for r in records:
                if not target and (r.get("file") or {}).get("path") == file_path_str:
                    target = r
                    continue
                kept.append(r)

            if not target:
                self._json({"ok": False, "msg": "未找到这条提交记录"})
                return

            submissions[assignment_id] = kept
            assignment = next((a for a in cfg.get("assignments", []) if a.get("id") == assignment_id), {})
            assignment_label = " · ".join([
                str(assignment.get("subject_group") or assignment.get("subject") or "").strip(),
                str(assignment.get("experiment") or assignment.get("name") or "").strip(),
            ]).strip(" ·") or assignment_id
            if assignment and rejected_terms:
                students = load_students()
                student_terms = {
                    str(value).strip().lower()
                    for s in students
                    for value in (s.get("name"), s.get("student_id"), s.get("pinyin"))
                    if str(value or "").strip()
                }
                negative = list(assignment.get("negative_keywords", []))
                negative_keys = {str(term).lower() for term in negative}
                for term in rejected_terms:
                    term_key = term.lower()
                    if term_key in student_terms or re.match(r"^电国\d*", term):
                        continue
                    if term_key not in negative_keys:
                        negative.append(term)
                        negative_keys.add(term_key)
                assignment["negative_keywords"] = negative[-100:]

            rejected_record = dict(target)
            rejected_record["student"] = None
            rejected_record["status"] = "unmatched"
            rejected_record["assignment_id"] = "other"
            rejected_record["assignment_name"] = "其他"
            rejected_record["match_score"] = 0
            rejected_record["context_hint"] = f"已从「{assignment_label}」移出"
            rejected_record["rejected_from"] = assignment_id
            rejected_record["rejected_at"] = datetime.now().isoformat()
            rejected_record.pop("organized_to", None)

            other_records = submissions.setdefault("other", [])
            other_records[:] = [
                r for r in other_records
                if (r.get("file") or {}).get("path") != file_path_str
            ]
            other_records.append(rejected_record)
            save_submissions(submissions)

            feedback = cfg.setdefault("match_feedback", {})
            rejected = feedback.setdefault("rejected", [])
            rejected.append({
                "file_name": (target.get("file") or {}).get("name", ""),
                "file_path": file_path_str,
                "assignment_id": assignment_id,
                "assignment_label": assignment_label,
                "student": target.get("student"),
                "rejected_terms": rejected_terms,
                "created_at": datetime.now().isoformat(),
            })
            feedback["rejected"] = rejected[-500:]
            cfg["match_feedback"] = feedback
            save_config(cfg)

            self._json({"ok": True, "msg": "已移出本作业", "assignment_id": assignment_id})

        elif path == "/api/submissions/delete":
            assignment_id = data.get("assignment_id", "")
            file_path_str = data.get("file_path", "")
            student_name = str(data.get("student", "") or "")
            detected_at = str(data.get("detected_at", "") or "")
            reason = str(data.get("reason", "") or "手动删除").strip() or "手动删除"
            if not assignment_id or not file_path_str:
                self._json({"ok": False, "msg": "缺少作业或文件信息"})
                return

            submissions = load_submissions()
            records = submissions.get(assignment_id, [])
            if not isinstance(records, list):
                records = []

            removed = None
            kept = []
            for r in records:
                file_info = r.get("file") or {}
                same_file = file_info.get("path") == file_path_str
                same_student = not student_name or r.get("student") == student_name
                same_time = not detected_at or r.get("detected_at") == detected_at
                if removed is None and same_file and same_student and same_time:
                    removed = r
                    continue
                kept.append(r)

            if removed is None:
                fallback_removed = None
                fallback_kept = []
                for r in records:
                    file_info = r.get("file") or {}
                    if fallback_removed is None and file_info.get("path") == file_path_str:
                        fallback_removed = r
                        continue
                    fallback_kept.append(r)
                if fallback_removed is not None:
                    removed = fallback_removed
                    kept = fallback_kept

            if removed is None:
                self._json({"ok": False, "msg": "未找到这条提交记录"})
                return

            submissions[assignment_id] = kept
            save_submissions(submissions)
            cfg = load_config_raw()
            _append_submission_delete_log(cfg, removed, assignment_id, reason)
            save_config(cfg)
            self._json({"ok": True, "msg": "已删除提交记录", "assignment_id": assignment_id, "reason": reason})

        elif path == "/api/unmatched/ignore":
            file_path_str = data.get("file_path")
            if file_path_str:
                cfg = load_config_raw()
                ignored = cfg.get("ignored_files", [])
                if file_path_str not in ignored:
                    ignored.append(file_path_str)
                    cfg["ignored_files"] = ignored
                    save_config(cfg)
            self._json({"ok": True})

        elif path == "/api/config/save":
            old_cfg = load_config_raw()
            data = _clean_config_paths(data)
            save_config(data)
            # 配置变更回显
            for key in ("watch_dir", "scan_dir", "class_folder"):
                old_val = old_cfg.get(key)
                new_val = data.get(key)
                if old_val != new_val and (old_val or new_val):
                    color = C.GREEN
                    msg = f"[CONFIG] {key}: {old_val} → {new_val}"
                    if new_val:
                        p = Path(new_val)
                        if not p.exists():
                            color = C.RED
                            msg += f"\n          [ERROR] New path not found: {new_val}"
                    print(f"{color}{msg}{C.RESET}")
            # scan_dirs 变更回显
            old_scan = old_cfg.get("scan_dirs", [])
            new_scan = data.get("scan_dirs", [])
            if set(old_scan) != set(new_scan):
                added = set(new_scan) - set(old_scan)
                removed = set(old_scan) - set(new_scan)
                if added:
                    print(f"{C.GREEN}[CONFIG] scan_dirs added: {added}{C.RESET}")
                if removed:
                    print(f"{C.YELLOW}[CONFIG] scan_dirs removed: {removed}{C.RESET}")
            self._json({"ok": True})

        elif path == "/api/scan-dirs/add":
            new_dir = (data.get("path") or "").strip()
            if not new_dir:
                self._json({"ok": False, "msg": "路径不能为空"})
                return
            cfg = load_config_raw()
            scan_dirs = cfg.get("scan_dirs", [])
            if new_dir not in scan_dirs:
                scan_dirs.append(new_dir)
                cfg["scan_dirs"] = _clean_path_list(scan_dirs)
                save_config(cfg)
                color = C.GREEN if Path(new_dir).exists() else C.YELLOW
                print(f"{color}[CONFIG] scan_dir added: {new_dir}{C.RESET}")
                if not Path(new_dir).exists():
                    print(f"{C.YELLOW}[WARN] Added path does not exist yet: {new_dir}{C.RESET}")
            scan_dirs = _clean_path_list(scan_dirs)
            self._json({"ok": True, "scan_dirs": scan_dirs_payload(cfg)})

        elif path == "/api/scan-dirs/remove":
            remove_dir = (data.get("path") or "").strip()
            if not remove_dir:
                self._json({"ok": False, "msg": "路径不能为空"})
                return
            cfg = load_config_raw()
            scan_dirs = cfg.get("scan_dirs", [])
            if remove_dir in scan_dirs:
                print(f"{C.YELLOW}[CONFIG] scan_dir removed: {remove_dir}{C.RESET}")
            scan_dirs = _clean_path_list([d for d in scan_dirs if d != remove_dir])
            cfg["scan_dirs"] = scan_dirs
            save_config(cfg)
            self._json({"ok": True, "scan_dirs": scan_dirs_payload(cfg)})

        elif path == "/api/experiment-dir/delete":
            cfg = load_config_raw()
            ok, msg = _safe_delete_experiment_dir(cfg)
            self._json({"ok": ok, "msg": msg, "experiment_dir": cfg.get("experiment_dir", "")})

        elif path == "/api/assignment/add":
            cfg = load_config()
            cfg.setdefault("assignments", [])
            name = str(data.get("name") or "").strip()
            subject = str(data.get("subject_group") or data.get("subject") or "").strip()
            experiment = str(data.get("experiment") or "").strip()
            if not name or not subject:
                self._json({"ok": False, "msg": "缺少作业名称或科目"})
                return
            subject = _canonical_subject_name(subject, cfg["assignments"])
            data["name"] = name
            data["subject_group"] = subject
            data["subject"] = subject
            data["experiment"] = experiment
            # manual 作业用 m\d+ 前缀，避免与 auto a\d+ 撞车
            data["id"] = _next_manual_assignment_id(cfg["assignments"])
            data.setdefault("active", True)
            data.setdefault("due", "")
            data.setdefault("notes", "")
            data.setdefault("completed", False)
            # 从 subject_group 自动派生 keywords
            sg = data.get("subject_group", "")
            if not data.get("keywords") and sg:
                for canon_sg, syns in _SUBJECT_SYNONYMS.items():
                    if sg in syns or canon_sg == sg:
                        data["keywords"] = list(syns)
                        break
            data.setdefault("keywords", [])
            keywords = []
            for keyword in data.get("keywords", []):
                keyword = str(keyword or "").strip()
                if keyword and keyword not in keywords:
                    keywords.append(keyword)
            for keyword in (subject, experiment):
                if keyword and keyword not in keywords:
                    keywords.append(keyword)
            data["keywords"] = keywords
            ignored = set(cfg.get("ignored_assignments", []))
            ignored.difference_update(_assignment_ignore_keys(data))
            cfg["ignored_assignments"] = sorted(ignored)
            cfg["assignments"].append(data)
            save_config(cfg)
            # 同步在公示/实验目录创建空目录
            _sync_experiment_dir(data.get("subject_group", ""), data.get("experiment", ""))
            refreshed = load_config()
            self._json({"ok": True, "assignment": data, "assignments": refreshed.get("assignments", [])})

        elif path == "/api/assignment/delete":
            if not any(data.get(k) for k in ("id", "subject_group", "subject", "experiment", "name")):
                self._json({"ok": False, "msg": "缺少作业信息"})
                return
            cfg = load_config()
            assignments = cfg.get("assignments", [])
            targets = [a for a in assignments if _assignment_matches_delete_payload(a, data)]
            if not targets:
                self._json({"ok": False, "msg": "未找到要删除的作业，请刷新后再试"})
                return
            ignored = set(cfg.get("ignored_assignments", []))
            for target in targets:
                ignored.update(_assignment_ignore_keys(target))
            ignored.update(_assignment_ignore_keys(data))
            cfg["ignored_assignments"] = sorted(ignored)
            target_ids = {a.get("id") for a in targets}
            target_pairs = {
                (_norm_assignment_text(a.get("subject_group") or a.get("subject")), _norm_assignment_text(a.get("experiment")))
                for a in targets
            }
            cfg["assignments"] = [
                a for a in assignments
                if a.get("id") not in target_ids and (
                    _norm_assignment_text(a.get("subject_group") or a.get("subject")),
                    _norm_assignment_text(a.get("experiment")),
                ) not in target_pairs
            ]
            save_config(cfg)
            self._json({"ok": True, "deleted": len(assignments) - len(cfg["assignments"]), "assignments": cfg["assignments"]})

        elif path == "/api/subject/delete":
            subject_name = data.get("subject", "")
            if not subject_name:
                self._json({"ok": False, "msg": "缺少科目名称"})
                return
            cfg = load_config()
            subject_name = _canonical_subject_name(subject_name, cfg.get("assignments", []))
            before = len(cfg.get("assignments", []))
            cfg["assignments"] = [a for a in cfg.get("assignments", []) if a.get("subject_group", "") != subject_name]
            after = len(cfg["assignments"])
            # 记录已删除科目，防止 load_config 从目录重建时复活
            ignored = cfg.setdefault("ignored_subjects", [])
            if subject_name not in ignored:
                ignored.append(subject_name)
                cfg["ignored_subjects"] = ignored
            save_config(cfg)
            self._json({"ok": True, "deleted": before - after, "assignments": cfg["assignments"]})

        elif path == "/api/assignment/update":
            aid = data.get("id")
            cfg = load_config()
            for a in cfg.get("assignments", []):
                if a.get("id") == aid:
                    for k, v in data.items():
                        if k != "id":
                            a[k] = v
                    break
            save_config(cfg)
            self._json({"ok": True})

        elif path == "/api/reset":
            watcher.known_files.clear()
            watcher._quick_init()
            self._json({"ok": True, "known_files": len(watcher.known_files)})

        elif path == "/api/convert-docx":
            paths = data.get("paths", [])
            results = convert_docx_to_pdf(paths)
            self._json({"results": results})

        elif path == "/api/analyze-file":
            fp = data.get("file_path", "")
            if not fp:
                self._json({"ok": False, "msg": "缺少 file_path"})
                return
            result = analyze_file_content(fp)
            self._json(result)

        elif path == "/api/confirm-classify":
            fp = data.get("file_path", "")
            student_name = data.get("student_name", "")
            assignment_id = data.get("assignment_id", "")
            new_basename = data.get("new_basename", "")  # 如 "班级名+学生名+科目+第N次+实验.docx"
            if not fp or not Path(fp).exists():
                self._json({"ok": False, "msg": "文件不存在"})
                return

            cfg = load_config()
            dest_base = Path(cfg.get("organized_dir", str(ORGANIZED_DIR)))
            src = Path(fp)
            ext = src.suffix

            # 确定目标名称
            if new_basename:
                dest_name = new_basename if new_basename.endswith(ext) else new_basename + ext
            else:
                dest_name = src.name

            # 确定目标路径（按 科目/第X次/学生姓名 结构）
            subject, experiment = get_subject_experiment_for_file(dest_name, assignment_id)
            if experiment:
                sub_dir = dest_base / subject / experiment / student_name
            else:
                sub_dir = dest_base / subject / student_name
            sub_dir.mkdir(parents=True, exist_ok=True)

            dest = sub_dir / dest_name
            # 避免覆盖
            if dest.exists():
                stem = Path(dest_name).stem
                dest = sub_dir / f"{stem}_{datetime.now().strftime('%H%M%S')}{ext}"

            try:
                shutil.copy2(str(src), str(dest))
                dest_str = str(dest)
                print(f"[ConfirmClassify] {fp} -> {dest_str}")

                # 双写到公示/实验目录（学生文件平铺，文件名加学生名前缀，已含前缀时去重）
                fname = Path(dest_name).name
                exp_dest_name = _experiment_dest_name(fname, student_name)
                ok = _copy_to_experiment(src, subject, experiment, exp_dest_name)
                if not ok:
                    print(f"[WARN] 公示目录写入失败（不阻断）: {Path(dest_name).name} -> {EXPERIMENT_BASE}")

                # 更新 submissions
                submissions = load_submissions()
                found = False
                # 按 (subject_group, experiment) 二元组查找目标 assignment，解决 aid 撞车
                target_aid = assignment_id
                if target_aid:
                    cfg = load_config_raw()
                    for a in cfg.get("assignments", []):
                        if a.get("subject_group") == subject and a.get("experiment") == experiment:
                            target_aid = a["id"]
                            break
                else:
                    target_aid = "manual"

                for aid_key, records in submissions.items():
                    if not isinstance(records, list):
                        continue
                    for r in records:
                        if r.get("file", {}).get("path") == fp:
                            r["student"] = student_name
                            r["status"] = "confirmed"
                            r["organized_to"] = dest_str
                            found = True
                            break
                    if found:
                        # 如果找到的 record 所在 aid 与 target_aid 不同，直接移动
                        if found and target_aid and aid_key != target_aid:
                            to_move = [r for r in records if r.get("file", {}).get("path") == fp]
                            for r in to_move:
                                records.remove(r)
                                submissions.setdefault(target_aid, []).append(r)
                        break

                if not found:
                    # 创建新记录
                    submissions.setdefault(target_aid, []).append({
                        "file": {"name": src.name, "path": fp, "size": file_size(src)},
                        "student": student_name,
                        "status": "confirmed",
                        "organized_to": dest_str,
                        "detected_at": datetime.now().isoformat(),
                    })

                save_submissions(submissions)
                self._json({"ok": True, "dest": dest_str})
            except Exception as e:
                self._json({"ok": False, "msg": f"复制失败: {str(e)[:200]}"})

        elif path == "/api/mark-template":
            fp = data.get("file_path", "")
            assignment_id = data.get("assignment_id", "")
            if not fp or not Path(fp).exists():
                self._json({"ok": False, "msg": "文件不存在"})
                return

            cfg = load_config()
            students = load_students()
            dest_base = Path(cfg.get("organized_dir", str(ORGANIZED_DIR)))
            template_dir = dest_base / "模板"
            template_dir.mkdir(parents=True, exist_ok=True)

            src = Path(fp)
            ext = src.suffix

            # 尝试分析文件内容以生成模板名称
            analysis = analyze_file_content(str(src))
            best = analysis.get("matches", [{}])[0] if analysis.get("ok") else {}

            # 根据分析结果自动匹配 assignment_id
            subject_name = best.get("subject") or ""
            exp_name = best.get("experiment") or ""
            if not assignment_id and subject_name:
                assignments = cfg.get("assignments", [])
                for a in assignments:
                    if a.get("subject_group") == subject_name:
                        # 优先同时匹配实验次数的
                        if exp_name and a.get("experiment") == exp_name:
                            assignment_id = a["id"]
                            break
                        # 退而求其次：同科目
                        elif not assignment_id:
                            assignment_id = a["id"]
            template_name_parts = []
            if subject_name:
                template_name_parts.append(subject_name)
            if exp_name:
                template_name_parts.append(exp_name)
            template_name_parts.append("模板")
            base = "".join(template_name_parts)
            dest_name = f"{base}{ext}"
            dest = template_dir / dest_name

            # 避免覆盖
            counter = 1
            while dest.exists():
                dest = template_dir / f"{base}_{counter}{ext}"
                counter += 1

            try:
                shutil.copy2(str(src), str(dest))
                dest_str = str(dest)
                print(f"[MarkTemplate] {fp} -> {dest_str}")

                # 记录到 config
                templates = cfg.get("templates", [])
                templates.append({
                    "name": dest_name,
                    "path": dest_str,
                    "assignment_id": assignment_id or "",
                    "subject": subject_name,
                    "experiment": exp_name,
                    "created_at": datetime.now().isoformat(),
                })
                cfg["templates"] = templates
                save_config(cfg)
                self._json({"ok": True, "dest": dest_str, "name": dest_name})
            except Exception as e:
                self._json({"ok": False, "msg": f"复制模板失败: {str(e)[:200]}"})

        elif path == "/api/server/restart":
            self._json({"ok": True, "msg": "正在重启..."})
            # 重启线程必须保持非 daemon：shutdown() 会让主 serve_forever 返回，
            # daemon 线程会随旧进程结束而被中断，导致新进程尚未拉起。
            threading.Thread(target=_restart_server, name="server-restart").start()
        elif path == "/api/server/shutdown":
            self._json({"ok": True, "msg": "正在关闭..."})
            threading.Thread(target=_shutdown_server, daemon=True).start()

        else:
            self.send_error(404)
            self.end_headers()

    def _auto_cleanup_completed(self, cfg):
        """自动删除标记完成超过1天的作业"""
        from datetime import timedelta
        assignments = cfg.get("assignments", [])
        if not assignments:
            return False
        modified = False
        remaining = []
        deleted_ids = []
        for a in assignments:
            completed_at_str = a.get("completed_at", "")
            if a.get("completed", False) and completed_at_str:
                try:
                    completed_at = datetime.fromisoformat(str(completed_at_str).replace("Z", "+00:00"))
                    now = datetime.now(completed_at.tzinfo) if completed_at.tzinfo else datetime.now()
                    if (now - completed_at) > timedelta(days=1):
                        print(f"[AutoClean] 删除已完成科目作业: {a.get('subject_group','')} {a.get('experiment','')} {a.get('name','')}（完成于 {completed_at_str}）")
                        deleted_ids.append(a.get("id"))
                        modified = True
                        continue  # 跳过此作业，不保留
                except (TypeError, ValueError):
                    pass
            remaining.append(a)
        if modified:
            cfg["assignments"] = remaining
            save_config(cfg)
            # 同步清理 submissions.json 中的孤儿数据
            if deleted_ids:
                submissions = load_submissions()
                for aid in deleted_ids:
                    if aid in submissions:
                        del submissions[aid]
                        print(f"[AutoClean] 已清理作业 {aid} 的提交记录")
                save_submissions(submissions)
        return modified

    def _dashboard_data(self, assignment_id=None):
        """组装仪表盘所需的汇总数据"""
        _t0 = time.time()
        cfg = load_config()
        if self._auto_cleanup_completed(cfg):
            cfg = load_config()
        print(f"[Dashboard] load_config: {time.time()-_t0:.3f}s")
        students = load_students()
        submissions = load_submissions()
        assignments = cfg.get("assignments", [])

        # 如果指定了作业ID，用那个；否则用第一个活跃的
        if assignment_id:
            active_assignment = next((a for a in assignments if a["id"] == assignment_id), assignments[0] if assignments else {"id": "other", "name": "默认"})
        else:
            active = [a for a in assignments if a.get("active", True)]
            active_assignment = active[0] if active else {"id": "other", "name": "默认"}

        aid = active_assignment["id"]
        raw_records = submissions.get(aid, [])
        if not isinstance(raw_records, list):
            raw_records = []
        # 展示有归纳路径或有学生名的记录（organized_to 优先，student 兜底）
        organized_dir = cfg.get("organized_dir", "")
        records = [r for r in raw_records if isinstance(r, dict) and (r.get("organized_to") or r.get("student"))]

        # 按学生汇总
        student_map = {}
        for s in students:
            student_map[s["name"]] = {
                "name": s["name"],
                "student_id": s.get("student_id", ""),
                "submitted": False,
                "files": [],
                "submit_time": None,
            }

        # 收集所有提交中的未匹配文件（跨所有科目，去重）
        ignored_files = cfg.get("ignored_files", [])
        unmatched = []
        pending_archive = []
        seen_unmatched = set()
        for aid_key, aid_records in submissions.items():
            if not isinstance(aid_records, list):
                continue
            for r in aid_records:
                if not isinstance(r, dict):
                    continue
                fp = r["file"].get("path", "")
                if not fp or fp in seen_unmatched:
                    continue
                classification = r.get("classification") or {}
                stage = classification.get("stage") or r.get("status", "")
                if stage in ("subject_matched", "assignment_pending", "subject_conflict", "subject_suggested", "unknown_subject"):
                    seen_unmatched.add(fp)
                    pending_archive.append({
                        "file_name": r["file"].get("name", ""), "file_path": fp,
                        "student": r.get("student", ""), "detected_at": r.get("detected_at", ""),
                        "size": r["file"].get("size", 0), "classification": classification,
                        "context_hint": "科目已识别但尚未确认具体作业，文件未归档也未同步到公示文件夹。",
                    })
                    continue
                student = r.get("student")
                if not student or student not in student_map:
                    # 跳过已忽略的文件
                    if fp in ignored_files:
                        seen_unmatched.add(fp)
                        continue
                    # 取文件 mtime（修改时间），文件不存在则为空
                    file_mtime = ""
                    try:
                        file_mtime = datetime.fromtimestamp(Path(fp).stat().st_mtime).isoformat()
                    except (OSError, FileNotFoundError):
                        pass
                    seen_unmatched.add(fp)
                    unmatched.append({
                        "file_name": r["file"]["name"],
                        "file_path": fp,
                        "detected_at": r.get("detected_at", ""),
                        "mtime": file_mtime,
                        "size": r["file"]["size"],
                        "context_hint": r.get("context_hint", ""),
                    })

        seen_active_files = set()
        for r in records:
            if r.get("student") and r.get("student") in student_map:
                key = _submission_record_dedupe_key(r, aid)
                if key in seen_active_files:
                    continue
                seen_active_files.add(key)
                sm = student_map[r["student"]]
                sm["submitted"] = True
                sm["files"].append({
                    "name": r["file"]["name"],
                    "size": r["file"]["size"],
                    "path": r["file"]["path"],
                    "time": r.get("detected_at", ""),
                    "organized_to": r.get("organized_to"),
                })
                if r.get("detected_at"):
                    if not sm["submit_time"] or r["detected_at"] < sm["submit_time"]:
                        sm["submit_time"] = r["detected_at"]

        submitted_count = sum(1 for s in student_map.values() if s["submitted"])
        not_submitted = [s for s in student_map.values() if not s["submitted"]]
        submitted = [s for s in student_map.values() if s["submitted"]]

        # 计算每个学生在所有作业上的提交状态（精确到 科目/实验次数 目录）
        assignment_status = {}
        for s in students:
            sn = s["name"]
            assignment_status[sn] = {}
            for a in assignments:
                aid2 = a["id"]
                a_sg = a.get("subject_group", "")
                a_exp = a.get("experiment", "")
                recs = submissions.get(aid2, [])
                has_submitted = False
                for r in recs:
                    if r.get("student") == sn:
                        org_to = r.get("organized_to")
                        if not org_to or is_file_in_correct_path(org_to, a_sg, a_exp):
                            has_submitted = True
                            break
                assignment_status[sn][aid2] = has_submitted

        # 每个学生在所有作业下的文件清单（仅归纳后的，且按 科目/实验次数 精确归类）
        all_files = {}
        for s in students:
            sn = s["name"]
            all_files[sn] = {}
            seen = set()  # 跨作业去重
            for a in assignments:
                aid2 = a["id"]
                recs = submissions.get(aid2, [])
                org_files = []
                a_sg = a.get("subject_group", "")
                a_exp = a.get("experiment", "")
                for r in recs:
                    if r.get("student") == sn:
                        # 精确检查：有 organized_to 时校验路径，无时用 student 兜底
                        org_to = r.get("organized_to")
                        if org_to and not is_file_in_correct_path(org_to, a_sg, a_exp):
                            continue
                        key = _submission_record_dedupe_key(r, aid2)
                        if key not in seen:
                            seen.add(key)
                            org_files.append({
                                "name": r["file"]["name"],
                                "path": r["file"].get("path", ""),
                                "size": r["file"].get("size", 0),
                                "organized_to": r.get("organized_to"),
                            })
                if org_files:
                    all_files[sn][aid2] = org_files

        # 科目列表：按常用顺序排列（而非字母序）
        subject_priority = {
            "课程作业": 1, "课程报告": 2, "项目作业": 3, "课程论文": 4,
            "实验报告": 5, "课程设计": 6, "小组作业": 7,
            "数字电子技术": 20, "程序设计": 21, "信号与系统": 22,
            "自动控制原理": 23, "单片机": 24, "电机学": 25, "电气导论": 26,
        }
        subjects = sorted(
            set(a.get("subject_group","") for a in assignments if a.get("subject_group","") not in ("其他","")),
            key=lambda x: subject_priority.get(x, 50)
        )
        
        self._json({
            "assignment": active_assignment,
            "assignments": assignments,
            "subjects": subjects,
            "assignment_status": assignment_status,
            "all_files": all_files,
            "students": students,
            "total": len(students),
            "submitted_count": submitted_count,
            "not_submitted_count": len(not_submitted),
            "not_submitted": not_submitted,
            "submitted": submitted,
            "unmatched_files": unmatched,
            "pending_archive_files": pending_archive,
            "recent_files": records[-20:] if records else [],
            "last_scan": datetime.now().isoformat(),
        })

    def _assignments_list(self):
        """返回所有作业及其提交统计"""
        cfg = load_config()
        if self._auto_cleanup_completed(cfg):
            cfg = load_config()
        students = load_students()
        submissions = load_submissions()
        assignments = cfg.get("assignments", [])
        # 先按 id 去重（保留先出现的，防止 config.json 手工重复）
        seen_ids = {}
        deduped = []
        for a in assignments:
            aid = a.get("id", "")
            if aid and aid not in seen_ids:
                seen_ids[aid] = True
                deduped.append(a)
        assignments = deduped
        result = []
        for a in assignments:
            records = submissions.get(a["id"], [])
            if not isinstance(records, list):
                records = []
            # 统计已提交：用统一的 _record_matches_assignment 入口（与详情页口径一致）
            # 该入口处理三层 fallback：aid 桶 → 路径匹配（含同义词） → filename 匹配
            submitted = set()
            for r in records:
                if not isinstance(r, dict) or not r.get("student"):
                    continue
                if not _record_matches_assignment(r, a):
                    continue
                submitted.add(r["student"])
            # 计算截止时间倒计时
            countdown = None
            due_str = a.get("due", "")
            if due_str:
                try:
                    due_date = datetime.strptime(due_str, "%Y-%m-%d").date()
                    delta = (due_date - date.today()).days
                    countdown = delta
                except:
                    pass

            result.append({
                "id": a["id"],
                "name": a["name"],
                "subject_group": a.get("subject_group", a.get("subject", "")),
                "experiment": a.get("experiment", ""),
                "due": due_str,
                "notes": a.get("notes", ""),
                "countdown": countdown,
                "active": a.get("active", True),
                "completed": a.get("completed", False),
                "keywords": a.get("keywords", []),
                "total": len(students),
                "submitted": len(submitted),
                "missing": len(students) - len(submitted),
                "rate": round(len(submitted) / len(students) * 100) if students else 0,
            })
        self._json(result)

    def _assignment_detail(self, aid):
        """返回单个作业的详细提交情况"""
        cfg = load_config()
        students = load_students()
        submissions = load_submissions()
        assignment = next((a for a in cfg.get("assignments", []) if a["id"] == aid), None)
        if not assignment:
            self._json({"error": "not found"})
            return

        records = submissions.get(aid, [])
        if not isinstance(records, list):
            records = []
        submitted_map = {}
        seen_files = set()  # 去重
        
        # 获取当前作业的科目和实验次数，用于精确过滤文件目录
        current_sg = assignment.get("subject_group", "")
        current_exp = assignment.get("experiment", "")
        
        for r in records:
            if not isinstance(r, dict):
                continue
            sn = r.get("student", "")
            if not sn:
                continue
            # 统一入口：aid 桶 + 路径匹配 + filename 匹配（与列表口径完全一致）
            # 保留班级路径校验作为"全收作业"目录的硬约束（防跨班级串）
            org_to = r.get("organized_to") or ""
            if org_to:
                class_name, class_folder, _, _ = resolve_class_config()
                if str(class_folder) not in org_to and class_name not in org_to:
                    continue
            # 用统一入口判断是否属于本作业（含 fallback 重新匹配）
            if not _record_matches_assignment(r, assignment):
                continue

            key = _submission_record_dedupe_key(r, aid)
            if key in seen_files:
                continue
            seen_files.add(key)
            if sn not in submitted_map:
                submitted_map[sn] = {"name": sn, "files": [], "time": None}
            submitted_map[sn]["files"].append({
                "name": r["file"]["name"],
                "size": r["file"]["size"],
                "path": r["file"]["path"],
                "organized_to": r.get("organized_to"),
                "time": r.get("detected_at", ""),
            })
            st = r.get("detected_at", "")
            if not submitted_map[sn]["time"] or st < submitted_map[sn]["time"]:
                submitted_map[sn]["time"] = st

        # 已提交
        submitted = []
        for s in students:
            if s["name"] in submitted_map:
                info = submitted_map[s["name"]]
                submitted.append({**s, **info})

        # 未提交
        not_submitted = [s for s in students if s["name"] not in submitted_map]

        self._json({
            "assignment": assignment,
            "total": len(students),
            "submitted_count": len(submitted),
            "not_submitted_count": len(not_submitted),
            "submitted": submitted,
            "not_submitted": not_submitted,
            "records": records,
        })

    def _json(self, data, status=200, extra_headers=None):
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
        self.send_header("Pragma", "no-cache")
        self.send_header("Content-Length", str(len(body)))
        for key, value in (extra_headers or {}).items():
            self.send_header(key, value)
        self.end_headers()
        self.wfile.write(body)

    def _error_html(self, title, detail):
        return f"""<html><head><meta charset='utf-8'><style>
            body{{font-family:"Microsoft YaHei",sans-serif;display:flex;align-items:center;justify-content:center;height:100vh;margin:0;background:#0f1117;color:#e1e4ed}}
            .box{{text-align:center;padding:40px;max-width:500px}}
            h2{{color:#ef4444;margin-bottom:16px}}
            p{{color:#8b8fa3;font-size:14px;line-height:1.8}}
        </style></head><body><div class='box'><h2>❌ {title}</h2><p>{detail}</p></div></body></html>"""

    def _text_html(self, text):
        safe = text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        return f"<html><head><meta charset='utf-8'><style>body{{font-family:monospace;white-space:pre-wrap;padding:20px;background:#0f1117;color:#e1e4ed;line-height:1.8}}</style></head><body>{safe}</body></html>"

    def _convert_upload(self, parts):
        """处理 docx 文件上传并转换为 PDF"""
        import concurrent.futures
        results = []
        for part in parts:
            filename = part.get("filename", "")
            if not filename:
                results.append({"filename": "(unknown)", "status": "error", "msg": "缺少文件名"})
                continue
            if not filename.lower().endswith(".docx"):
                results.append({"filename": filename, "status": "skip", "msg": "仅支持 .docx 文件"})
                continue

            try:
                # 保存上传的 docx 到临时目录
                src_path = CONVERT_TEMP_DIR / filename
                src_path.write_bytes(part["data"])

                # 转换（静默模式）
                pdf_path = src_path.with_suffix(".pdf")
                with _WordComContext() as ctx:
                    doc = ctx.open(src_path)
                    doc.ExportAsFixedFormat(str(pdf_path), 17)

                ok = pdf_path.exists() and pdf_path.stat().st_size > 0
                results.append({
                    "filename": filename,
                    "status": "ok" if ok else "error",
                    "pdf_path": str(pdf_path) if ok else "",
                    "pdf_url": f"/api/download?path={str(pdf_path)}" if ok else "",
                    "msg": "" if ok else "PDF 输出为空",
                })
            except Exception as e:
                results.append({"filename": filename, "status": "error", "msg": str(e)[:200]})

        # 统计
        ok_count = sum(1 for r in results if r["status"] == "ok")
        skip_count = sum(1 for r in results if r["status"] == "skip")
        err_count = sum(1 for r in results if r["status"] == "error")

        self._json({
            "ok": True,
            "results": results,
            "summary": {"total": len(results), "ok": ok_count, "skip": skip_count, "error": err_count},
        })

    def _upload_template(self, parts):
        """处理模板文件上传"""
        cfg = load_config()
        dest_base = Path(cfg.get("organized_dir", str(ORGANIZED_DIR)))
        template_dir = dest_base / "模板"
        template_dir.mkdir(parents=True, exist_ok=True)

        assignment_id = ""
        file_part = None
        for part in parts:
            name = part.get("name", "")
            if name == "assignment_id":
                assignment_id = part.get("data", b"").decode("utf-8", errors="replace")
            elif name == "file":
                file_part = part

        if not file_part:
            self._json({"ok": False, "msg": "缺少文件"})
            return

        filename = file_part.get("filename", "template.docx")
        ext = Path(filename).suffix.lower()
        if ext not in (".docx", ".doc"):
            self._json({"ok": False, "msg": "仅支持 .docx 和 .doc 文件"})
            return

        # 保存到模板目录
        dest = template_dir / filename
        counter = 1
        stem, suffix = Path(filename).stem, Path(filename).suffix
        while dest.exists():
            dest = template_dir / f"{stem}_{counter}{suffix}"
            counter += 1

        dest.write_bytes(file_part.get("data", b""))

        # 分析文件内容获取科目和实验
        analysis = analyze_file_content(str(dest))
        best = analysis.get("matches", [{}])[0] if analysis.get("ok") else {}
        subject_name = best.get("subject") or ""
        exp_name = best.get("experiment") or ""

        # 自动匹配 assignment_id
        if not assignment_id and subject_name:
            assignments = cfg.get("assignments", [])
            for a in assignments:
                if a.get("subject_group") == subject_name:
                    if exp_name and a.get("experiment") == exp_name:
                        assignment_id = a["id"]
                        break
                    elif not assignment_id:
                        assignment_id = a["id"]

        # 记录到 config
        templates = cfg.get("templates", [])
        templates.append({
            "name": dest.name,
            "path": str(dest),
            "assignment_id": assignment_id or "",
            "subject": subject_name,
            "experiment": exp_name,
            "created_at": datetime.now().isoformat(),
        })
        cfg["templates"] = templates
        save_config(cfg)

        print(f"[UploadTemplate] {dest.name} (subject={subject_name}, exp={exp_name}, aid={assignment_id})")
        self._json({"ok": True, "name": dest.name, "path": str(dest)})

    def _build_update_package(self, data):
        """管理员一键生成更新包：当前代码 + 公告 → ZIP 下载"""
        import zipfile
        from urllib.parse import quote

        version = str(data.get("version") or current_app_version() or "1.0.0").strip()
        announcements = data.get("announcements") or []
        if not isinstance(announcements, list):
            announcements = []

        if not version:
            self._json({"ok": False, "msg": "缺少版本号"})
            return

        safe_version = re.sub(r"[^0-9A-Za-z._-]+", "_", version).strip("._-") or "update"
        package_files = [
            "server.py",
            "ai_classifier.py",
            "dashboard.html",
            "dashboard_modern.html",
            "restart_helper.py",
            "pack.py",
            "repair_update.py",
            "repair_update.bat",
            "启动作业追踪器.bat",
            "更新修复工具.bat",
        ]

        manifest = {
            "app": "Assignment_Dashboard",
            "version": version,
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "files": package_files,
            "has_changelog": False,
            "has_announcement": len(announcements) > 0,
        }
        announcement_payload = {
            "version": version,
            "announcements": announcements,
        }

        try:
            out = io.BytesIO()
            with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zf:
                for name in package_files:
                    fp = BASE_DIR / name
                    if fp.exists() and fp.is_file():
                        zf.write(fp, name)
                if announcements:
                    zf.writestr("announcement.json", json.dumps(announcement_payload, ensure_ascii=False, indent=2).encode("utf-8"))
                zf.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False, indent=2).encode("utf-8"))

            zip_data = out.getvalue()
            filename = f"dashboard_update_v{safe_version}.zip"
            encoded_filename = quote(filename)
            self.send_response(200)
            self.send_header("Content-Type", "application/zip")
            self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{encoded_filename}")
            self.send_header("Content-Length", str(len(zip_data)))
            self.end_headers()
            self.wfile.write(zip_data)
            print(f"[BuildUpdate] {filename}: {len(zip_data)/1024:.1f} KB, announcements={len(announcements)}")
        except Exception as e:
            traceback.print_exc()
            self._json({"ok": False, "msg": f"生成更新包失败: {str(e)[:200]}"})

    def _check_github_update(self):
        """查询 GitHub Releases 中最新的公开更新包。"""
        api_url = f"https://api.github.com/repos/{UPDATE_REPOSITORY}/releases/latest"
        request = urllib.request.Request(
            api_url,
            headers={
                "Accept": "application/vnd.github+json",
                "User-Agent": "Assignment_Dashboard-Updater",
            },
        )
        try:
            with urllib.request.urlopen(request, timeout=8) as response:
                release = json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as e:
            if e.code == 404:
                self._json({"ok": False, "msg": "GitHub 仓库不存在，或尚未发布任何 Release"}, status=404)
            else:
                self._json({"ok": False, "msg": f"GitHub 暂时无法访问（HTTP {e.code}）"}, status=502)
            return
        except (urllib.error.URLError, TimeoutError, json.JSONDecodeError) as e:
            self._json({"ok": False, "msg": f"检查更新失败：{str(e)[:160]}"}, status=502)
            return

        assets = release.get("assets") or []
        update_assets = [
            asset for asset in assets
            if str(asset.get("name", "")).lower().endswith(".zip")
            and str(asset.get("name", "")).lower().startswith("dashboard_update_v")
        ]
        if not update_assets:
            self._json({"ok": False, "msg": "最新 Release 没有找到 dashboard_update_v*.zip 更新包"}, status=404)
            return

        asset = update_assets[0]
        latest_version = str(release.get("tag_name") or "").strip().lstrip("vV")
        self._json({
            "ok": True,
            "repository": UPDATE_REPOSITORY,
            "current_version": current_app_version(),
            "latest_version": latest_version,
            "has_update": _version_key(latest_version) > _version_key(current_app_version()),
            "release_name": release.get("name") or release.get("tag_name") or "最新版本",
            "release_url": release.get("html_url", ""),
            "published_at": release.get("published_at", ""),
            "body": release.get("body", "") or "",
            "asset": {
                "name": asset.get("name", ""),
                "size": asset.get("size", 0),
                "download_url": asset.get("browser_download_url", ""),
            },
        })

    def _handle_update(self, parts):
        """处理系统更新：接收 ZIP 包 → 校验 → 备份 → 解压替换 → 重启"""
        import zipfile
        import hashlib

        # 1. 提取 ZIP 数据
        zip_data = None
        zip_filename = "update.zip"
        for part in parts:
            if part.get("name") == "update_zip":
                zip_data = part.get("data", b"")
                zip_filename = part.get("filename", "update.zip")
                break

        if not zip_data:
            self._json({"ok": False, "msg": "未收到更新包文件"})
            return

        # 2. 校验 ZIP 有效性
        try:
            with zipfile.ZipFile(io.BytesIO(zip_data), "r") as zf:
                bad = zf.testzip()
                if bad:
                    self._json({"ok": False, "msg": f"更新包损坏: {bad}"})
                    return
                file_list = zf.namelist()
        except zipfile.BadZipFile:
            self._json({"ok": False, "msg": "无效的 ZIP 文件"})
            return

        # 3. 验证关键文件存在
        required_files = ["server.py", "dashboard.html"]
        missing = [f for f in required_files if f not in file_list]
        if missing:
            self._json({"ok": False, "msg": f"更新包缺少关键文件: {', '.join(missing)}"})
            return

        # 4. 提取 CHANGELOG.md 内容
        changelog_content = ""
        if "CHANGELOG.md" in file_list:
            try:
                with zipfile.ZipFile(io.BytesIO(zip_data), "r") as zf:
                    changelog_content = zf.read("CHANGELOG.md").decode("utf-8", errors="replace")
            except Exception:
                changelog_content = ""

        # 4b. 检测并保存 announcement.json
        has_announcement = False
        announcement_content = None
        if "announcement.json" in file_list:
            try:
                with zipfile.ZipFile(io.BytesIO(zip_data), "r") as zf:
                    ann_data = zf.read("announcement.json")
                    json.loads(ann_data.decode("utf-8", errors="replace"))
                    announcement_content = ann_data
                    has_announcement = True
                    print("[Update] 检测到公告文件: announcement.json")
            except Exception as e:
                self._json({"ok": False, "msg": f"公告文件格式错误: {str(e)[:200]}"})
                return

        # 5. 创建备份（带时间戳）
        backup_dir = BASE_DIR / "backups"
        backup_dir.mkdir(exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_name = f"backup_{timestamp}.zip"
        backup_path = backup_dir / backup_name

        try:
            # 备份当前关键文件
            backup_entries = []
            for item in ["server.py", "ai_classifier.py", "dashboard.html", "dashboard_modern.html", "pack.py", "repair_update.py", "repair_update.bat", "CHANGELOG.md", "announcement.json", "manifest.json", "启动作业追踪器.bat", "更新修复工具.bat"]:
                fp = BASE_DIR / item
                if fp.exists():
                    backup_entries.append((str(fp), item))
            # 也备份 data 目录
            data_dir = BASE_DIR / "data"
            if data_dir.exists():
                for f in data_dir.rglob("*"):
                    if f.is_file():
                        arcname = str(f.relative_to(BASE_DIR))
                        backup_entries.append((str(f), arcname))

            backup_zip_data = _create_zip(backup_entries)
            backup_path.write_bytes(backup_zip_data)
            print(f"[Update] 备份已创建: {backup_path} ({len(backup_zip_data)/1024:.1f} KB)")
        except Exception as e:
            self._json({"ok": False, "msg": f"备份失败: {str(e)[:200]}"})
            return

        # 6. 解压替换文件
        updated_files = []
        try:
            with zipfile.ZipFile(io.BytesIO(zip_data), "r") as zf:
                for member in zf.namelist():
                    # 保护用户数据：不覆盖 data/*.json 文件
                    if member.startswith("data/") and member.endswith(".json"):
                        print(f"[Update] 跳过用户数据文件: {member}")
                        continue
                    # 提取到项目根目录
                    target_path = BASE_DIR / member
                    # 安全检查：确保目标在 BASE_DIR 内
                    try:
                        target_path.resolve().relative_to(BASE_DIR.resolve())
                    except ValueError:
                        print(f"[Update] 安全跳过: {member} (路径越界)")
                        continue
                    # 创建父目录
                    target_path.parent.mkdir(parents=True, exist_ok=True)
                    # 提取文件
                    content = zf.read(member)
                    target_path.write_bytes(content)
                    updated_files.append(member)
                    print(f"[Update] 已更新: {member}")
        except Exception as e:
            # 回滚：恢复备份
            print(f"[Update] 解压失败，开始回滚: {e}")
            self._restore_backup(backup_path)
            self._json({"ok": False, "msg": f"更新失败，已自动回滚: {str(e)[:200]}"})
            return

        # 7. 清理旧备份（保留最近 5 个）
        try:
            backups = sorted(backup_dir.glob("backup_*.zip"), key=lambda p: p.stat().st_mtime, reverse=True)
            for old in backups[5:]:
                old.unlink()
                print(f"[Update] 清理旧备份: {old.name}")
        except Exception:
            pass

        # 8. 触发重启（延迟 1 秒，确保响应已发送）
        def _restart_delayed():
            time.sleep(0.8)
            print("[Update] 正在重启服务...")
            try:
                _restart_server()
            except Exception as e:
                print(f"[Update] 重启失败: {e}")

        # 更新后的重启也需要保活到新进程就绪，不能随旧服务退出而中断。
        threading.Thread(target=_restart_delayed, name="update-restart").start()

        self._json({
            "ok": True,
            "msg": "更新成功，服务即将重启，请刷新浏览器" + ("（本次更新含公告）" if has_announcement else ""),
            "updated_files": updated_files,
            "backup": str(backup_path),
            "changelog": changelog_content,
            "has_announcement": has_announcement,
        })

    def _restore_backup(self, backup_path):
        """从备份恢复文件"""
        import zipfile
        try:
            if not backup_path.exists():
                print("[Update] 备份文件不存在，无法回滚")
                return
            with zipfile.ZipFile(backup_path, "r") as zf:
                zf.extractall(BASE_DIR)
            print("[Update] 回滚成功")
        except Exception as e:
            print(f"[Update] 回滚失败: {e}")

    def _send_download(self, file_path):
        """发送文件供用户下载（浏览器弹出保存对话框）"""
        fp = _safe_resolve_path(file_path)
        if not fp or not fp.exists() or not fp.is_file():
            self._json({"ok": False, "msg": "文件不存在"})
            return
        convert_root = _safe_resolve_path(CONVERT_TEMP_DIR)
        if not convert_root or not _is_path_inside(fp, convert_root):
            self._json({"ok": False, "msg": "无权访问该文件"})
            return
        try:
            content = fp.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", "application/pdf")
            self.send_header("Content-Disposition", f'attachment; filename="{fp.name}"')
            self.send_header("Content-Length", str(len(content)))
            self.end_headers()
            self.wfile.write(content)
        except Exception as e:
            self._json({"ok": False, "msg": str(e)})

    def _send_zip(self, paths):
        """打包多个PDF为ZIP并供用户下载"""
        valid_paths = []
        convert_root = _safe_resolve_path(CONVERT_TEMP_DIR)
        for p in paths:
            fp = _safe_resolve_path(p.strip())
            if fp and convert_root and fp.exists() and fp.is_file() and _is_path_inside(fp, convert_root):
                valid_paths.append(fp)

        if not valid_paths:
            self._json({"ok": False, "msg": "没有可下载的文件"})
            return

        entries = [(str(fp), fp.name) for fp in valid_paths]
        zip_data = _create_zip(entries)

        self.send_response(200)
        self.send_header("Content-Type", "application/zip")
        self.send_header("Content-Disposition", 'attachment; filename="converted_pdfs.zip"')
        self.send_header("Content-Length", str(len(zip_data)))
        self.end_headers()
        self.wfile.write(zip_data)

    def _pack_subject(self, subject_group):
        """一键打包：将某个科目所有实验的已提交文件按 科目/第X次/学生名 结构打包为 ZIP"""
        import traceback

        if not subject_group:
            self._json({"ok": False, "msg": "缺少 subject_group 参数"})
            return

        try:
            cfg = load_config()
            assignments = cfg.get("assignments", [])
            submissions = load_submissions()
            students = load_students()
            student_names = {s["name"] for s in students}

            # 找出该科目下的所有作业
            subject_assignments = [a for a in assignments if a.get("subject_group") == subject_group]
            if not subject_assignments:
                self._json({"ok": False, "msg": f"未找到科目 '{subject_group}' 的作业"})
                return

            # 按实验次数排序
            EXP_ORDER = {"第一次": 1, "第二次": 2, "第三次": 3, "第四次": 4, "第五次": 5, "第六次": 6}
            subject_assignments.sort(key=lambda a: EXP_ORDER.get(a.get("experiment", ""), 99))

            # 收集文件：{ (实验次数, 学生名): file_path } 每学生每实验只保留一个文件
            zip_entries = {}
            file_count = 0

            for a in subject_assignments:
                exp_name = a.get("experiment", a.get("name", ""))
                records = submissions.get(a["id"], [])
                if not isinstance(records, list):
                    continue

                for r in records:
                    if not isinstance(r, dict):
                        continue
                    org_to = r.get("organized_to")
                    if not org_to:
                        continue
                    org_path = is_allowed_file_path(org_to)
                    if not org_path:
                        continue

                    student = r.get("student", "未知")
                    file_name = org_path.name

                    # 去重：同学生+同实验只保留一个文件，优先保留不带 (N) 后缀的
                    dedup_key = (exp_name, student)
                    if dedup_key in zip_entries:
                        existing = zip_entries[dedup_key]
                        existing_name = Path(existing[0]).name
                        # 当前文件不含 (1)(2) 但已有文件含 → 替换
                        import re as _re2
                        if not _re2.search(r'\(\d+\)', file_name) and _re2.search(r'\(\d+\)', existing_name):
                            zip_entries[dedup_key] = (str(org_path), f"{subject_group}/{exp_name}/{student}_{file_name}")
                        continue

                    arcname = f"{subject_group}/{exp_name}/{student}_{file_name}"
                    zip_entries[dedup_key] = (str(org_path), arcname)
                    file_count += 1

            if not zip_entries:
                self._json({"ok": False, "msg": f"科目 '{subject_group}' 下没有已提交的文件"})
                return

            # 使用通用 _create_zip 打包
            entries = [(file_path, arcname) for (file_path, arcname) in sorted(zip_entries.values(), key=lambda x: x[1])]
            zip_data = _create_zip(entries)

            # 文件名（URL编码避免中文 header 问题）
            from urllib.parse import quote
            safe_name = subject_group.replace("/", "_").replace("\\", "_")
            filename = f"{safe_name}.zip"
            encoded_filename = quote(filename)

            self.send_response(200)
            self.send_header("Content-Type", "application/zip")
            self.send_header("Content-Disposition", f'attachment; filename*=UTF-8\'\'{encoded_filename}')
            self.send_header("Content-Length", str(len(zip_data)))
            self.end_headers()
            self.wfile.write(zip_data)
            print(f"[PackSubject] {subject_group}: {file_count} files, {len(zip_data)/1024/1024:.1f} MB")
        except Exception as e:
            traceback.print_exc()
            self._json({"ok": False, "msg": f"打包失败: {str(e)[:200]}"})

    def _serve_html(self, filename):
        file_path = BASE_DIR / filename
        if file_path.exists():
            self._serve_static(file_path)
        else:
            self.send_error(404)
            self.end_headers()

    def _serve_static(self, file_path):
        suffix_map = {
            ".html": "text/html; charset=utf-8",
            ".css": "text/css",
            ".js": "application/javascript",
            ".json": "application/json",
            ".png": "image/png",
            ".svg": "image/svg+xml",
            ".ico": "image/x-icon",
        }
        content_type = suffix_map.get(file_path.suffix.lower(), "application/octet-stream")
        try:
            with open(file_path, "rb") as f:
                content = f.read()
            self.send_response(200)
            self.send_header("Content-Type", content_type)
            self.send_header("Content-Length", str(len(content)))
            if file_path.suffix.lower() in (".html", ".js", ".css"):
                self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
                self.send_header("Pragma", "no-cache")
            self.end_headers()
            self.wfile.write(content)
        except:
            self.send_error(404)
            self.end_headers()

    def do_OPTIONS(self):
        self.send_response(204)
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()


# ---------------------------------------------------------------------------
# 启动
# ---------------------------------------------------------------------------

def main():
    global watcher, _http_server, _app_start_time
    _app_start_time = time.time()
    parser = argparse.ArgumentParser(description="作业提交追踪器")
    parser.add_argument("--port", type=int, default=18765, help="HTTP 服务端口")
    parser.add_argument("--no-watch", action="store_true", help="不启动文件监控")
    args = parser.parse_args()

    print("=" * 50)
    print(f"  作业提交追踪器  v{current_app_version()}")
    print("=" * 50)

    if not acquire_server_lock(args.port):
        return

    startup_cfg = load_config_raw()
    bind_host = "0.0.0.0" if startup_cfg.get("lan_access_enabled", False) else "127.0.0.1"
    try:
        server = ThreadingHTTPServer((bind_host, args.port), APIHandler)
        _http_server = server
    except OSError as e:
        release_server_lock()
        print(f"\n[ERROR] 端口 {args.port} 无法启动：{e}")
        print(f"可能已经有一个作业追踪器在运行，请先打开 http://localhost:{args.port} 检查。")
        print("如果确认是旧进程残留，可以关闭旧命令行窗口后再启动。\n")
        return

    # 初始化文件监控
    watcher = FileWatcher()
    if not args.no_watch:
        watcher.start()

    # 启动 HTTP
    print(f"\n  Dashboard:  http://localhost:{args.port}")
    print(f"  学生管理:  http://localhost:{args.port}/dashboard\n")
    print(f"  访问模式:  {'局域网' if bind_host == '0.0.0.0' else '仅本机'} ({bind_host})")
    if bind_host == "0.0.0.0":
        for url in network_access_payload(startup_cfg, args.port).get("lan_urls", []):
            print(f"  局域网地址: {url}")

    # --- 启动路径检查 ---
    cfg = load_config()
    scan_dirs = cfg.get("scan_dirs", [])
    organized_dir = cfg.get("organized_dir", str(ORGANIZED_DIR))

    print("监控目录:")
    for d in get_effective_watch_dirs(cfg):
        exists = _path_exists(d)
        mark = "" if exists else "  ⚠ 目录不存在"
        print(f"  {d}{mark}")

    if scan_dirs:
        print("\n扫描目录:")
        for d in scan_dirs:
            exists = _path_exists(d)
            mark = "" if exists else "  ⚠ 目录不存在"
            print(f"  {d}{mark}")

    print(f"\n已收作业: {organized_dir}")
    if not Path(organized_dir).exists():
        print("  ⚠ 已收作业目录不存在，将自动创建")

    exp_dir_str = cfg.get("experiment_dir", str(EXPERIMENT_BASE)) if EXPERIMENT_BASE else ""
    if exp_dir_str:
        exp_exists = Path(exp_dir_str).exists() if exp_dir_str else False
        exp_mark = "" if exp_exists else "  ⚠ 目录不存在"
        print(f"\n公示文件夹: {exp_dir_str}{exp_mark}")

    print("按 Ctrl+C 停止\n")

    # 启动后预热 Word 预览缓存（已禁用：COM 弹窗问题）
    # warm_preview_cache_async()

    # 启动自动回填：扫描公示目录，未入库文件自动入库
    _auto_backfill_from_experiment()

    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\n[Shutdown] Stopping...")
        watcher.stop()
        server.shutdown()
    finally:
        release_server_lock()

# ---------------------------------------------------------------------------
# 服务器控制：重启 & 关闭（仅在后台线程中调用）
# ---------------------------------------------------------------------------

def _shutdown_server():
    """关闭服务器：停监控 → 关 HTTP → 释放锁 → 退出进程"""
    try:
        if watcher:
            watcher.stop()
            if watcher.thread:
                watcher.thread.join(timeout=2)
        if _http_server:
            _http_server.shutdown()
        time.sleep(0.3)
    except Exception as e:
        print(f"[Shutdown] 异常（可忽略）: {e}")
    finally:
        release_server_lock()
        os._exit(0)


def _restart_after_delay():
    time.sleep(0.8)
    _restart_server()

def _restart_server():
    """Restart through a detached helper so Windows can finish releasing the port."""
    # 取出端口：优先从 _http_server 拿实际 bind 的端口
    port = None
    if _http_server is not None:
        try:
            port = _http_server.server_address[1]
        except Exception:
            pass
    if port is None:
        # 回退：从命令行参数解析
        for i, arg in enumerate(sys.argv):
            if arg == "--port" and i + 1 < len(sys.argv):
                try:
                    port = int(sys.argv[i + 1])
                except ValueError:
                    pass
                break
    if port is None:
        port = 18765  # 与 main() 默认值保持一致
    try:
        helper_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "restart_helper.py")
        if not os.path.exists(helper_path):
            raise RuntimeError("restart_helper.py is missing; please reinstall the update package")

        # Hand off before stopping this process. The helper waits for the port
        # to disappear and retries child startup if Windows is still releasing it.
        release_server_lock()
        helper_kwargs = {"cwd": os.path.dirname(os.path.abspath(__file__))}
        if sys.platform == "win32":
            helper_kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW
        helper_proc = subprocess.Popen(
            [sys.executable, helper_path, "--port", str(port), "--", sys.executable,
             os.path.abspath(__file__)] + sys.argv[1:],
            **helper_kwargs
        )
        print(f"[Restart] 接力进程已启动，PID: {helper_proc.pid}")

        if watcher:
            watcher.stop()
            if watcher.thread:
                watcher.thread.join(timeout=2)
        time.sleep(0.3)
        if _http_server:
            _http_server.shutdown()
            # shutdown() 只停止 serve_forever，不会释放监听 socket。
            # 若不显式 close，新进程会因端口仍被旧进程占用而直接退出。
            _http_server.server_close()
        next_cfg = load_config_raw()
        next_host = "0.0.0.0" if next_cfg.get("lan_access_enabled", False) else "127.0.0.1"
        print(f"[Restart] 旧服务已停止，等待接力进程在 {next_host}:{port} 拉起新服务")
    except Exception as e:
        print(f"[Restart] 异常: {e}")
        traceback.print_exc()
        return


if __name__ == "__main__":
    main()
